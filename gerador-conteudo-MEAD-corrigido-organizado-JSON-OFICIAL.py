import json
import os
import re
import hashlib
import time
import unicodedata
import tempfile
import threading
import requests
import fitz
import tkinter as tk
from datetime import datetime
from tkinter import filedialog
from tkinter import messagebox
from tkinter import ttk

from docx import Document
from bs4 import BeautifulSoup
from ddgs import DDGS



ARQUIVO_BANCO = r"C:\Python\gerador-conteudo\banco\conteudo-site.json"

ARQUIVO_MEAD = r"C:\Python\gerador-conteudo\mead\mead.json"

ARQUIVO_PROGRESSO = r"C:\Python\gerador-conteudo\banco\progresso.json"

janela = None

IA_PROCESSANDO = False
PROCESSAMENTO_ATIVO = False
etapa_atual = 0
inicio_geracao = 0.0
tempos_etapas = []
etapas_total = 4
PAGINAS_EM_PROCESSAMENTO = set()

from pathlib import Path

PASTA_DADOS = Path(
    r"C:\Python\gerador-conteudo\dados-brutos"
)

CATEGORIAS = [
    "definicao",
    "beneficios",
    "vantagens",
    "materia_prima",
    "aplicacoes",
    "fabricacao",
    "manutencao",
    "ativos_narrativos",
    "duvidas_frequentes"
]


# ============================================================
# CARREGAR MEAD
# ============================================================

def carregar_mead():

    if not os.path.exists(ARQUIVO_MEAD):

        print("❌ Arquivo MEAD não encontrado:")
        print(ARQUIVO_MEAD)

        return {}


    try:

        with open(
            ARQUIVO_MEAD,
            "r",
            encoding="utf-8"
        ) as arquivo:

            mead = json.load(arquivo)


        if not isinstance(mead, dict):

            print("❌ MEAD inválido. Estrutura esperada: objeto JSON.")

            return {}


        print("\nMEAD carregado:")
        print(mead)


        return mead


    except json.JSONDecodeError as erro:

        print("❌ Erro no JSON do MEAD:")
        print(erro)

        return {}


    except Exception as erro:

        print("❌ Erro ao carregar MEAD:")
        print(erro)

        return {}



# ============================================================
# PREPARAR MEAD PARA IA
# ============================================================

def preparar_mead(mead):

    """
    Prepara somente as regras editoriais necessárias do MEAD
    para serem utilizadas pelo Ollama.

    IMPORTANTE:
    - Não envia o JSON MEAD inteiro.
    - Não envia patrimônio ou textos de pesquisa.
    - Não seleciona informações.
    - Não altera os fragmentos escolhidos pelo Python.
    - O Python continua responsável pela seleção.
    - O Ollama continua responsável pela redação.
    """

    if not isinstance(mead, dict):
        return ""

    try:

        identidade = mead.get("identidade", {})
        principio = mead.get("principio_operacional", {})
        preparacao = mead.get("preparacao", {})
        tags = preparacao.get("tags", {})
        construcao = mead.get("construcao_conteudo", {})
        estrutura = mead.get("estrutura_pagina_editorial", {})
        seo = mead.get("regras_seo", {})
        fidelidade = mead.get("fidelidade", {})
        contexto_blocos = mead.get("contexto_blocos", {})
        lista_segmentos = mead.get("lista_segmentos", {})

        mead_ia = {

            "metodo": identidade.get(
                "metodo",
                "MEAD"
            ),

            "objetivo": identidade.get(
                "objetivo",
                ""
            ),

            "principio_operacional": {

                "regra_principal": principio.get(
                    "regra_principal",
                    ""
                ),

                "regra_de_fidelidade": principio.get(
                    "regra_de_fidelidade",
                    ""
                ),

                "regra_de_protagonismo": principio.get(
                    "regra_de_protagonismo",
                    ""
                ),

                "regra_de_diversidade": principio.get(
                    "regra_de_diversidade",
                    ""
                ),

                "regra_de_separacao": principio.get(
                    "regra_de_separacao",
                    ""
                )
            },

            # ==================================================
            # PALAVRA-CHAVE
            # ==================================================

            "palavra_chave": preparacao.get(
                "palavra_chave",
                {}
            ),

            # ==================================================
            # REGRAS DE TAGS PARA O OLLAMA
            # ==================================================

            "tags": {

                "quantidade": tags.get(
                    "quantidade",
                    30
                ),

                "regra": tags.get(
                    "regra",
                    ""
                ),

                "principio_origem": tags.get(
                    "principio_origem",
                    {}
                ),

                "grupos": tags.get(
                    "grupos",
                    []
                ),

                "ordem_prioridade": tags.get(
                    "ordem_prioridade",
                    []
                ),

                "regra_localizacao": tags.get(
                    "regra_localizacao",
                    {}
                ),

                "regra_servicos": {

                    "usar": tags.get(
                        "regra_servicos",
                        {}
                    ).get(
                        "usar",
                        True
                    ),

                    "base": tags.get(
                        "regra_servicos",
                        {}
                    ).get(
                        "base",
                        ""
                    ),

                    "variacoes_permitidas": tags.get(
                        "regra_servicos",
                        {}
                    ).get(
                        "variacoes_permitidas",
                        []
                    ),

                    "regra": tags.get(
                        "regra_servicos",
                        {}
                    ).get(
                        "regra",
                        ""
                    )
                },

                "regra_produtos": {

                    "usar": tags.get(
                        "regra_produtos",
                        {}
                    ).get(
                        "usar",
                        True
                    ),

                    "base": tags.get(
                        "regra_produtos",
                        {}
                    ).get(
                        "base",
                        ""
                    ),

                    "variacoes_permitidas": tags.get(
                        "regra_produtos",
                        {}
                    ).get(
                        "variacoes_permitidas",
                        []
                    ),

                    "regra": tags.get(
                        "regra_produtos",
                        {}
                    ).get(
                        "regra",
                        ""
                    )
                },

                "regra_completude": tags.get(
                    "regra_completude",
                    {}
                ),

                "regra_substituicao": tags.get(
                    "regra_substituicao",
                    {}
                ),

                "restricoes": tags.get(
                    "restricoes",
                    []
                )
            },

            # ==================================================
            # CONTEXTO DOS CINCO BLOCOS
            # ==================================================

            "contexto_blocos": contexto_blocos,

            # ==================================================
            # CONSTRUÇÃO DO CONTEÚDO
            # ==================================================

            "construcao_conteudo": {

                "total_blocos": construcao.get(
                    "total_blocos",
                    5
                ),

                "paragrafos_por_bloco": construcao.get(
                    "paragrafos_por_bloco",
                    3
                ),

                "total_paragrafos": construcao.get(
                    "total_paragrafos",
                    15
                ),

                "palavras_por_paragrafo":
                    construcao.get(
                        "palavras_por_paragrafo",
                        {}
                    ),

                "regra_narrativa":
                    construcao.get(
                        "regra_narrativa",
                        {}
                    )
            },

            # ==================================================
            # ESTRUTURA DA PÁGINA
            # ==================================================

            "estrutura_pagina": {

                "blocos_texto": estrutura.get(
                    "blocos_texto",
                    5
                ),

                "paragrafos_por_bloco":
                    estrutura.get(
                        "paragrafos_por_bloco",
                        3
                    ),

                "total_paragrafos":
                    estrutura.get(
                        "total_paragrafos",
                        15
                    ),

                "palavras_por_paragrafo":
                    estrutura.get(
                        "palavras_por_paragrafo",
                        {}
                    )
            },


           

            # ==================================================
            # SEO
            # ==================================================

            "regras_seo": {

                "evitar": seo.get(
                    "evitar",
                    []
                ),

                "priorizar": seo.get(
                    "priorizar",
                    []
                )
            },

            # ==================================================
            # FIDELIDADE
            # ==================================================

            "fidelidade": {

                "permitido": fidelidade.get(
                    "permitido",
                    []
                ),

                "proibido": fidelidade.get(
                    "proibido",
                    []
                )
            }
        }

        return json.dumps(
            mead_ia,
            ensure_ascii=False,
            indent=2
        )

    except Exception as erro:

        print("❌ Erro ao preparar MEAD para IA:")
        print(erro)

        return ""


# ============================================================
# PREPARAR DADOS FIXOS DA PÁGINA
# ============================================================

def preparar_dados_pagina(
        tema,
        tags=None,
        segmentos_listas=None,
        dados_pagina=None
    ):

    """
    Organiza os dados já produzidos pelo Python
    para serem enviados posteriormente para
    montar_pagina_json().

    NÃO pesquisa.
    NÃO chama Ollama.
    NÃO gera conteúdo.
    NÃO altera fragmentos.
    NÃO grava o JSON.

    IMPORTANTE:
    - Não cria mais "informacoes_adicionais".
    - Não cria "categorias".
    - Os parágrafos da página devem ser recebidos
      posteriormente pelo fluxo de montagem dos blocos.
    """

    try:

        # ----------------------------------------------------
        # TEMA
        # ----------------------------------------------------

        tema = str(
            tema or ""
        ).strip()

        if not tema:

            print(
                "❌ preparar_dados_pagina(): "
                "tema vazio."
            )

            return None

        # ----------------------------------------------------
        # DADOS DA PÁGINA
        # ----------------------------------------------------

        if not isinstance(
            dados_pagina,
            dict
        ):

            dados_pagina = {}

        # ----------------------------------------------------
        # TAGS
        # ----------------------------------------------------

        if not isinstance(
            tags,
            list
        ):

            tags = []

        tags_finais = []

        for tag in tags:

            tag = str(
                tag or ""
            ).strip()

            if not tag:
                continue

            if tag in tags_finais:
                continue

            tags_finais.append(
                tag
            )

            if len(
                tags_finais
            ) >= 30:

                break

        dados_pagina[
            "tags"
        ] = tags_finais

        # ----------------------------------------------------
        # SEGMENTOS / LISTAS
        # ----------------------------------------------------

        if not isinstance(
            segmentos_listas,
            dict
        ):

            segmentos_listas = {}

        segmentos_oficiais = {}

        # ----------------------------------------------------
        # GARANTIR OS 12 SEGMENTOS
        # ----------------------------------------------------

        for numero in range(
            1,
            13
        ):

            chave = (
                f"segmento_{numero}"
            )

            lista = segmentos_listas.get(
                chave,
                []
            )

            if not isinstance(
                lista,
                list
            ):

                lista = []

            lista_final = []

            for item in lista:

                item = str(
                    item or ""
                ).strip()

                if not item:
                    continue

                lista_final.append(
                    item
                )

            segmentos_oficiais[
                chave
            ] = lista_final

        dados_pagina[
            "segmentos_listas"
        ] = segmentos_oficiais

        # ----------------------------------------------------
        # IMPORTANTE
        # ----------------------------------------------------
        #
        # NÃO criar:
        #
        # dados_pagina["informacoes_adicionais"]
        #
        # NÃO criar:
        #
        # dados_pagina["categorias"]
        #
        # Essas estruturas antigas não fazem mais parte
        # da estrutura editorial final da página.
        #
        # Os textos dos 15 parágrafos devem entrar nos
        # blocos da página posteriormente.
        # ----------------------------------------------------

        dados_pagina.pop(
            "informacoes_adicionais",
            None
        )

        dados_pagina.pop(
            "categorias",
            None
        )

        # ----------------------------------------------------
        # RETORNO
        # ----------------------------------------------------

        print()
        print(
            "=========================================="
        )
        print(
            "DADOS FIXOS DA PÁGINA PREPARADOS"
        )
        print(
            "=========================================="
        )
        print(
            f"TEMA: {tema}"
        )
        print(
            f"TAGS: {len(tags_finais)}"
        )

        total_itens = 0

        for numero in range(
            1,
            13
        ):

            chave = (
                f"segmento_{numero}"
            )

            quantidade = len(
                segmentos_oficiais[
                    chave
                ]
            )

            total_itens += quantidade

            print(
                f"{chave}: "
                f"{quantidade} itens"
            )

        print(
            f"TOTAL ITENS DAS LISTAS: "
            f"{total_itens}"
        )

        print(
            "INFORMACOES_ADICIONAIS: REMOVIDO"
        )

        print(
            "CATEGORIAS: REMOVIDO"
        )

        print(
            "=========================================="
        )

        return dados_pagina

    except Exception as erro:

        print()
        print(
            "❌ ERRO ao preparar dados da página:"
        )
        print(erro)

        return None


        
        
# ============================================================
# MONTAR PÁGINA OFICIAL PARA O JSON
# ============================================================

def montar_pagina_json(
        tema,
        grupo="",
        tipo="",
        tags=None,
        controle_repeticoes=None,
        mapa_mead=None,
        informacoes_relevantes=None,
        dados_pagina=None,
        informacoes_adicionais=None,
        categorias=None
    ):

    """
    Monta a estrutura oficial completa de uma página
    antes da gravação no conteudo-site.json.

    IMPORTANTE:

    - Não realiza pesquisa.
    - Não seleciona informações.
    - Não gera conteúdo.
    - Não altera os fragmentos recebidos.
    - Não grava diretamente no arquivo.
    - Apenas organiza os dados na estrutura oficial.

    ESTRUTURA:

        tema
            └── pagina
                ├── tema
                ├── arquivo_origem
                ├── h1
                ├── titulo
                ├── subtitulo
                ├── descricao
                ├── bloco_1
                │   └── informacoes_relevantes
                ├── bloco_2
                │   └── informacoes_relevantes
                ├── bloco_3
                │   └── informacoes_relevantes
                ├── bloco_4
                │   └── informacoes_relevantes
                ├── bloco_5
                │   └── informacoes_relevantes
                ├── segmentos_listas
                │   ├── segmento_1
                │   ├── ...
                │   └── segmento_12
                └── imagens

    NÃO EXISTE MAIS:

        pagina_completa["informacoes_relevantes"]

    As informações relevantes pertencem aos cinco blocos.
    """

    try:

        # ====================================================
        # NORMALIZAR TEMA
        # ====================================================

        tema = str(
            tema or ""
        ).strip()

        if not tema:

            raise ValueError(
                "Tema vazio ao montar página JSON."
            )

        # ====================================================
        # NORMALIZAR ENTRADAS
        # ====================================================

        if not isinstance(
            tags,
            list
        ):

            tags = []

        if not isinstance(
            controle_repeticoes,
            dict
        ):

            controle_repeticoes = {}

        if not isinstance(
            mapa_mead,
            dict
        ):

            mapa_mead = {}

        if not isinstance(
            informacoes_relevantes,
            dict
        ):

            informacoes_relevantes = {}

        if not isinstance(
            dados_pagina,
            dict
        ):

            dados_pagina = {}

        if not isinstance(
            informacoes_adicionais,
            dict
        ):

            informacoes_adicionais = {}

        if not isinstance(
            categorias,
            dict
        ):

            categorias = {}

        # ====================================================
        # CRIAR ESTRUTURA OFICIAL
        # ====================================================

        estrutura = criar_estrutura_json_pagina(
            tema
        )

        if not isinstance(
            estrutura,
            dict
        ):

            print(
                "❌ criar_estrutura_json_pagina() "
                "não retornou um dicionário."
            )

            return None

        if tema not in estrutura:

            print(
                "❌ Tema não encontrado na estrutura "
                "criada para o JSON."
            )

            return None

        pagina_completa = estrutura[
            tema
        ]

        if not isinstance(
            pagina_completa,
            dict
        ):

            print(
                "❌ Estrutura da página inválida."
            )

            return None

        pagina = pagina_completa.get(
            "pagina"
        )

        if not isinstance(
            pagina,
            dict
        ):

            print(
                "❌ Campo 'pagina' inexistente "
                "ou inválido."
            )

            return None

        # ====================================================
        # DADOS PRINCIPAIS
        # ====================================================

        pagina_completa[
            "tema"
        ] = tema

        pagina_completa[
            "grupo"
        ] = str(
            grupo or ""
        ).strip()

        pagina_completa[
            "tipo"
        ] = str(
            tipo or ""
        ).strip()

        # ====================================================
        # TAGS
        # ====================================================

        tags_finais = []

        for tag in tags:

            tag_limpa = str(
                tag or ""
            ).strip()

            if not tag_limpa:
                continue

            if tag_limpa in tags_finais:
                continue

            tags_finais.append(
                tag_limpa
            )

            if len(
                tags_finais
            ) >= 30:

                break

        pagina_completa[
            "tags"
        ] = tags_finais

        # ====================================================
        # CONTROLE DE REPETIÇÕES
        # ====================================================

        pagina_completa[
            "controle_repeticoes"
        ] = {

            "palavra_chave":
                str(
                    controle_repeticoes.get(
                        "palavra_chave",
                        tema
                    )
                    or tema
                ).strip(),

            "meta_repeticoes":
                controle_repeticoes.get(
                    "meta_repeticoes",
                    60
                ),

            "repeticoes_realizadas":
                controle_repeticoes.get(
                    "repeticoes_realizadas",
                    0
                ),

            "repeticoes_faltantes":
                controle_repeticoes.get(
                    "repeticoes_faltantes",
                    60
                )
        }

        # ====================================================
        # MAPA MEAD
        # ====================================================

        pagina_completa[
            "mapa_mead"
        ] = {

            "status":
                str(
                    mapa_mead.get(
                        "status",
                        ""
                    )
                    or ""
                ).strip(),

            "texto":
                str(
                    mapa_mead.get(
                        "texto",
                        ""
                    )
                    or ""
                ).strip()
        }

        # ====================================================
        # DADOS GERAIS DA PÁGINA
        # ====================================================

        campos_pagina = [

            "tema",
            "arquivo_origem",
            "h1",
            "titulo",
            "subtitulo",
            "descricao"

        ]

        for campo in campos_pagina:

            if campo in dados_pagina:

                pagina[
                    campo
                ] = str(
                    dados_pagina.get(
                        campo,
                        ""
                    )
                    or ""
                ).strip()

        # ====================================================
        # GARANTIR TEMA
        # ====================================================

        pagina[
            "tema"
        ] = tema

        # ====================================================
        # REGRA OFICIAL
        #
        # H1 = palavra-chave
        # TÍTULO = palavra-chave
        # ====================================================

        pagina[
            "h1"
        ] = tema

        pagina[
            "titulo"
        ] = tema

        # ====================================================
        # ARQUIVO DE ORIGEM
        # ====================================================

        pagina[
            "arquivo_origem"
        ] = str(
            dados_pagina.get(
                "arquivo_origem",
                pagina.get(
                    "arquivo_origem",
                    ""
                )
            )
            or ""
        ).strip()

        # ====================================================
        # SUBTÍTULO
        # ====================================================

        pagina[
            "subtitulo"
        ] = str(
            dados_pagina.get(
                "subtitulo",
                pagina.get(
                    "subtitulo",
                    ""
                )
            )
            or ""
        ).strip()

        # ====================================================
        # DESCRIÇÃO
        # ====================================================

        pagina[
            "descricao"
        ] = str(
            dados_pagina.get(
                "descricao",
                pagina.get(
                    "descricao",
                    ""
                )
            )
            or ""
        ).strip()

        # ====================================================
        # INFORMAÇÕES RELEVANTES POR BLOCO
        # ====================================================

        fragmentos_blocos = (
            informacoes_relevantes.get(
                "blocos",
                {}
            )
        )

        if not isinstance(
            fragmentos_blocos,
            dict
        ):

            fragmentos_blocos = {}

        # ====================================================
        # ATUALIZAR OS CINCO BLOCOS
        # ====================================================

        for numero in range(
            1,
            6
        ):

            chave_bloco = (
                f"bloco_{numero}"
            )

            dados_bloco = dados_pagina.get(
                chave_bloco,
                {}
            )

            if not isinstance(
                dados_bloco,
                dict
            ):

                dados_bloco = {}

            # ------------------------------------------------
            # FRAGMENTOS SELECIONADOS PELO PYTHON
            # ------------------------------------------------

            fragmentos_bloco = (
                fragmentos_blocos.get(
                    chave_bloco,
                    ""
                )
            )

            # ------------------------------------------------
            # Se vier como lista de fragmentos,
            # transformar em texto mantendo todos.
            # ------------------------------------------------

            if isinstance(
                fragmentos_bloco,
                list
            ):

                fragmentos_bloco = "\n\n".join(
                    str(
                        item or ""
                    ).strip()
                    for item in fragmentos_bloco
                    if str(
                        item or ""
                    ).strip()
                )

            else:

                fragmentos_bloco = str(
                    fragmentos_bloco or ""
                ).strip()

            # ------------------------------------------------
            # COPIAR DADOS EXISTENTES DO BLOCO
            # ------------------------------------------------

            dados_bloco_final = dict(
                dados_bloco
            )

            # ------------------------------------------------
            # INFORMACOES_RELEVANTES FICA DENTRO DO BLOCO
            # ------------------------------------------------

            dados_bloco_final[
                "informacoes_relevantes"
            ] = fragmentos_bloco

            # ------------------------------------------------
            # ATUALIZAR BLOCO
            # ------------------------------------------------

            atualizar_bloco(
                numero,
                dados_bloco_final
            )

        # ====================================================
        # SEGMENTOS DAS LISTAS
        # ====================================================

        segmentos_recebidos = (
            dados_pagina.get(
                "segmentos_listas",
                {}
            )
        )

        if not isinstance(
            segmentos_recebidos,
            dict
        ):

            segmentos_recebidos = {}

        segmentos_oficiais = {}

        for numero in range(
            1,
            13
        ):

            chave_segmento = (
                f"segmento_{numero}"
            )

            valor = segmentos_recebidos.get(
                chave_segmento,
                []
            )

            if not isinstance(
                valor,
                list
            ):

                valor = []

            segmentos_oficiais[
                chave_segmento
            ] = [

                str(
                    item or ""
                ).strip()

                for item in valor

                if str(
                    item or ""
                ).strip()
            ]

        pagina[
            "segmentos_listas"
        ] = segmentos_oficiais

        # ====================================================
        # POSICIONAMENTO DAS LISTAS
        # ====================================================

        posicionamento = (
            dados_pagina.get(
                "posicionamento_listas",
                {}
            )
        )

        if not isinstance(
            posicionamento,
            dict
        ):

            posicionamento = {}

        pagina[
            "posicionamento_listas"
        ][
            "bloco"
        ] = posicionamento.get(
            "bloco"
        )

        # ====================================================
        # IMAGENS
        # ====================================================

        imagens = (
            dados_pagina.get(
                "imagens",
                {}
            )
        )

        if not isinstance(
            imagens,
            dict
        ):

            imagens = {}

        for numero in range(
            1,
            7
        ):

            chave_imagem = (
                f"imagem_{numero}"
            )

            imagem = imagens.get(
                chave_imagem,
                {}
            )

            if not isinstance(
                imagem,
                dict
            ):

                imagem = {}

            pagina[
                "imagens"
            ][
                chave_imagem
            ] = {

                "url":
                    str(
                        imagem.get(
                            "url",
                            ""
                        )
                        or ""
                    ).strip(),

                "arquivo":
                    str(
                        imagem.get(
                            "arquivo",
                            ""
                        )
                        or ""
                    ).strip(),

                "alt":
                    str(
                        imagem.get(
                            "alt",
                            ""
                        )
                        or ""
                    ).strip(),

                "descricao":
                    str(
                        imagem.get(
                            "descricao",
                            ""
                        )
                        or ""
                    ).strip()
            }

        # ====================================================
        # ESTRUTURAS LEGADAS — NÃO GRAVAR
        # ====================================================
        
        pagina_completa.pop(
            "informacoes_adicionais",
            None
        )

        # ====================================================
        # CATEGORIAS — NÃO GRAVAR NO JSON
        # ====================================================
        
        pagina_completa.pop(
            "categorias",
            None
        )

        # ====================================================
        # CALCULAR CARACTERES DA PÁGINA
        # ====================================================

        textos = []

        for numero in range(
            1,
            6
        ):

            chave_bloco = (
                f"bloco_{numero}"
            )

            bloco = pagina.get(
                chave_bloco,
                {}
            )

            if not isinstance(
                bloco,
                dict
            ):

                continue

            # ------------------------------------------------
            # INFORMACOES RELEVANTES
            # ------------------------------------------------

            textos.append(
                str(
                    bloco.get(
                        "informacoes_relevantes",
                        ""
                    )
                    or ""
                )
            )

            # ------------------------------------------------
            # TITULO DO BLOCO
            # ------------------------------------------------

            textos.append(
                str(
                    bloco.get(
                        "titulo",
                        ""
                    )
                    or ""
                )
            )

            # ------------------------------------------------
            # PARAGRAFOS
            # ------------------------------------------------

            paragrafos = bloco.get(
                "paragrafos",
                []
            )

            if isinstance(
                paragrafos,
                list
            ):

                textos.extend(
                    str(
                        item or ""
                    )
                    for item in paragrafos[:3]
                )

        pagina[
            "caracteres"
        ] = sum(
            len(
                texto
            )
            for texto in textos
        )

        # ====================================================
        # STATUS
        # ====================================================

        pagina[
            "status"
        ] = "pronta_para_gravacao"

        # ====================================================
        # CHECK DE SEGURANÇA
        #
        # Nunca permitir informacoes_relevantes
        # no nível global da página.
        # ====================================================

        if (
            "informacoes_relevantes"
            in pagina_completa
        ):

            print(
                "⚠️ ERRO: campo global "
                "'informacoes_relevantes' "
                "foi recriado."
            )

            pagina_completa.pop(
                "informacoes_relevantes",
                None
            )

        # ====================================================
        # CHECK DOS FRAGMENTOS POR BLOCO
        # ====================================================

        print()
        print(
            "======================================"
        )
        print(
            "CHECK INFORMAÇÕES RELEVANTES"
        )
        print(
            "======================================"
        )

        total_fragmentos = 0

        for numero in range(
            1,
            6
        ):

            chave_bloco = (
                f"bloco_{numero}"
            )

            bloco = pagina.get(
                chave_bloco,
                {}
            )

            texto_bloco = ""

            if isinstance(
                bloco,
                dict
            ):

                texto_bloco = str(
                    bloco.get(
                        "informacoes_relevantes",
                        ""
                    )
                    or ""
                ).strip()

            caracteres_bloco = len(
                texto_bloco
            )

            if caracteres_bloco > 0:

                print(
                    f"🟢 {chave_bloco}: "
                    f"{caracteres_bloco} caracteres"
                )

                total_fragmentos += (
                    caracteres_bloco
                )

            else:

                print(
                    f"🔴 {chave_bloco}: "
                    "sem fragmentos"
                )

        print(
            "TOTAL FRAGMENTOS NOS BLOCOS: "
            f"{total_fragmentos}"
        )

        print(
            "======================================"
        )

        # ====================================================
        # CHECK DAS 12 LISTAS
        # ====================================================

        print()
        print(
            "======================================"
        )
        print(
            "CHECK SEGMENTOS DAS LISTAS"
        )
        print(
            "======================================"
        )

        total_itens_listas = 0

        for numero in range(
            1,
            13
        ):

            chave_segmento = (
                f"segmento_{numero}"
            )

            itens = segmentos_oficiais.get(
                chave_segmento,
                []
            )

            quantidade = len(
                itens
            )

            total_itens_listas += (
                quantidade
            )

            if quantidade > 0:

                print(
                    f"🟢 {chave_segmento}: "
                    f"{quantidade} itens"
                )

            else:

                print(
                    f"🔴 {chave_segmento}: "
                    "sem itens"
                )

        print(
            "TOTAL ITENS DAS 12 LISTAS: "
            f"{total_itens_listas}"
        )

        print(
            "======================================"
        )

        # ====================================================
        # RETORNO
        # ====================================================

        return pagina_completa

    except Exception as erro:

        print()
        print(
            "❌ Erro ao montar página oficial:"
        )

        print(
            erro
        )

        return None


# ========================================================
# 01. CARREGAR MEAD GLOBAL
# ========================================================

MEAD = carregar_mead()

MEAD_TEXTO = preparar_mead(MEAD)

print("\nMEAD preparado para IA:")
print(MEAD_TEXTO)


# ========================================================
# 02. GRUPOS TEMÁTICOS
# ========================================================

GRUPOS_TEMATICOS = {

    "protecao_contra_incendio": [

        "argamassa",
        "intumescente",
        "selagem",
        "corta fogo",
        "contra incendio",
        "lã de rocha",
        "la de rocha",
        "colar intumescente",
        "fita intumescente",
        "proteção passiva",
        "protecao passiva"

    ],

    "valvulas_industriais": [

        "válvula",
        "valvula",
        "hidraulica",
        "hidráulica",
        "esfera",
        "gaveta",
        "globo",
        "retenção",
        "retencao",
        "atuador"

    ],

    "componentes_mecanicos": [

        "terminal rotular",
        "terminal rotular esférico",
        "terminal rotular esferico",
        "rótula",
        "rotula",
        "rolamento",
        "mancal",
        "articulação",
        "articulacao",
        "junta esférica",
        "junta esferica"

    ]

}




# ============================================================
# IDENTIFICAR TIPO DO TEMA
# ============================================================

def identificar_tipo_tema(tema):

    tema = normalizar_texto(tema)


    if any(x in tema for x in [
        "empresa",
        "especialista",
        "profissional"
    ]):

        return "empresa"


    elif any(x in tema for x in [
        "servico",
        "aplicacao",
        "instalacao",
        "aplicar"
    ]):

        return "servico"


    elif any(x in tema for x in [
        "comprar",
        "fornecedor",
        "fabricante"
    ]):

        return "fornecedor"


    elif any(x in tema for x in [
        "orcamento",
        "cotacao",
        "preco",
        "valor"
    ]):

        return "orcamento"


    return "produto"




# ============================================================
# PESQUISA
# ============================================================

def pesquisar(termo, limite=10):

    resultados = []


    dominios_bloqueados = [

        "pinterest.",
        "facebook.",
        "instagram.",
        "youtube.com",
        "tiktok."

    ]


    try:

        with DDGS() as ddgs:


            busca = ddgs.text(
                termo,
                safesearch="off",
                max_results=limite
            )


            for item in busca:


                if not isinstance(item, dict):

                    continue


                url = item.get("href")


                if not url:

                    continue


                if any(
                    dominio in url.lower()
                    for dominio in dominios_bloqueados
                ):

                    continue


                resultados.append(url)



    except Exception as e:

        print()
        print("ERRO NA PESQUISA:")
        print(e)



    return resultados


# ============================================================
# FILTRAR URLS RUINS
# ============================================================

def filtrar_urls(urls):


    bloqueados = [

        "mercadolivre",
        "amazon",
        "shopee",
        "aliexpress",
        "alibaba",
        "ebay",
        "made-in-china",

        "facebook",
        "instagram",
        "youtube",
        "pinterest",
        "tiktok"

    ]


    resultado = []


    for url in urls:


        if not url:

            continue


        url_lower = url.lower()


        if any(
            item in url_lower
            for item in bloqueados
        ):

            continue


        resultado.append(url)



    return resultado




# ============================================================
# FILTRAR URL
# ============================================================

def url_valida(url):


    if not url:

        return False


    url_lower = url.lower()



    dominios_bloqueados = [

        "oceanofpdf.com",
        "pdfcoffee.com",
        "scribd.com",

        "pinterest.com",
        "facebook.com",
        "instagram.com",

        "tiktok.com",
        "youtube.com",
        "youtu.be"

    ]



    palavras_bloqueadas = [

        "ebook",
        "torrent",
        "download",
        "pdf-epub",
        "curso",
        "apostila",
        "manual-download"

    ]



    if any(
        dominio in url_lower
        for dominio in dominios_bloqueados
    ):

        return False



    if any(
        palavra in url_lower
        for palavra in palavras_bloqueadas
    ):

        return False



    return True
    
    


    

# ============================================================
# PESQUISA COMPLETA
# ============================================================

def pesquisar_completo(tema):

    urls = []


    # ========================================================
    # 01. PESQUISA PDF BRASIL
    # ========================================================

    print()
    print("==============================")
    print("PESQUISA PDF BRASIL")
    print("==============================")


    consultas_pdf_brasil = [

        f"{tema} pdf",

        f"{tema} catálogo pdf",

        f"{tema} catalogo pdf",

        f"{tema} ficha técnica pdf",

        f"{tema} ficha tecnica pdf",

        f"{tema} manual pdf"

    ]


    total_pdf_brasil = 0


    for consulta in consultas_pdf_brasil:

        resultado = pesquisar(

            consulta,

            limite=5

        )

        urls.extend(resultado)

        total_pdf_brasil += len(resultado)


    print(

        "PDF BRASIL:",

        total_pdf_brasil

    )


    # ========================================================
    # 02. PESQUISA PDF EXTERIOR
    # ========================================================

    print()
    print("==============================")
    print("PESQUISA PDF EXTERIOR")
    print("==============================")


    consultas_pdf_exterior = [

        f"{tema} datasheet pdf",

        f"{tema} catalog pdf",

        f"{tema} technical data pdf",

        f"{tema} specification pdf",

        f"{tema} specifications pdf",

        f"{tema} installation manual pdf"

    ]


    total_pdf_exterior = 0


    for consulta in consultas_pdf_exterior:

        resultado = pesquisar(

            consulta,

            limite=5

        )

        urls.extend(resultado)

        total_pdf_exterior += len(resultado)


        print(
        "PDF EXTERIOR:",
        total_pdf_exterior
    )


    # ========================================================
    # 03. PESQUISA POR CATEGORIAS
    # ========================================================

    print()
    print("==============================")
    print("PESQUISA POR CATEGORIAS")
    print("==============================")


    mapa_categorias = {

        "definicao": "definição",

        "beneficios": "benefícios",

        "vantagens": "vantagens",

        "materia_prima": "materiais",

        "aplicacoes": "aplicações",

        "fabricacao": "fabricação",

        "manutencao": "manutenção",

        "ativos_narrativos": "problemas casos",

        "duvidas_frequentes": "dúvidas frequentes"

    }


    total_categorias = 0


    for categoria in CATEGORIAS:

        termo = mapa_categorias.get(

            categoria,

            categoria.replace("_", " ")

        )

        consulta = f"{tema} {termo}"


        print()

        print(
            "CONSULTA CATEGORIA:",
            consulta
        )


        resultado = pesquisar(

            consulta,

            limite=5

        )


        urls.extend(resultado)


        total_categorias += len(resultado)


    print()

    print(
        "FONTES POR CATEGORIAS:",
        total_categorias
    )



    # ========================================================
    # 03.1 BRASIL - 10 FONTES
    # ========================================================

    print()
    print("==============================")
    print("PESQUISA BRASIL")
    print("==============================")


    brasil = pesquisar(

        f"{tema} Brasil",

        limite=10

    )


    urls.extend(

        brasil

    )


    print(

        "BRASIL:",

        len(brasil)

    )


    # =====================================
    # EXTERIOR - 10 FONTES
    # PESQUISA INTERNACIONAL INTELIGENTE
    # =====================================

    print()
    print("==============================")
    print("PESQUISA EXTERIOR")
    print("==============================")


    exterior = []


    # ========================================================
    # 04. DEFINIR CONSULTAS PELO GRUPO DO TEMA
    # ========================================================

    consultas_exterior = []


    tema_normalizado = normalizar_texto(

        tema

    )


    grupo_encontrado = None


    for grupo, palavras in GRUPOS_TEMATICOS.items():

        for palavra in palavras:

            if normalizar_texto(

                palavra

            ) in tema_normalizado:

                grupo_encontrado = grupo

                break

        if grupo_encontrado:

            break
    
    
    # ========================================================
    # 05. COMPONENTES MECÂNICOS
    # ========================================================

    if grupo_encontrado == "componentes_mecanicos":

        consultas_exterior = [

            f"{tema} manufacturer",

            f"{tema} technical catalog",

            f"{tema} engineering",

            f"{tema} rod end",

            f"{tema} spherical rod end",

            f"{tema} spherical plain bearing",

            f"{tema} technical",

            f"{tema} application"

        ]


    # ========================================================
    # 06. PROTEÇÃO CONTRA INCÊNDIO
    # ========================================================

    elif grupo_encontrado == "protecao_contra_incendio":

        consultas_exterior = [

            f"{tema} fire protection",

            f"{tema} manufacturer",

            f"{tema} technical catalog",

            f"{tema} datasheet",

            f"{tema} application",

            f"{tema} passive fire protection"

        ]


    # ========================================================
    # 07. VÁLVULAS
    # ========================================================

    elif grupo_encontrado == "valvulas_industriais":

        consultas_exterior = [

            f"{tema} industrial valve",

            f"{tema} manufacturer",

            f"{tema} technical catalog",

            f"{tema} datasheet",

            f"{tema} engineering",

            f"{tema} application"

        ]


    # ========================================================
    # 08. PADRÃO
    # ========================================================

    else:

        consultas_exterior = [

            f"{tema} manufacturer",

            f"{tema} technical",

            f"{tema} engineering",

            f"{tema} catalog",

            f"{tema} application"

        ]


    # ========================================================
    # 09. EXECUTAR BUSCAS INTERNACIONAIS
    # ========================================================

    for consulta in consultas_exterior:

        encontrados = pesquisar(

            consulta,

            limite=5

        )

        exterior.extend(

            encontrados

        )


    # ========================================================
    # 10. REMOVER DUPLICADOS DO EXTERIOR
    # ========================================================

    exterior = list(

        dict.fromkeys(exterior)

    )


    urls.extend(

        exterior[:10]

    )


    print(

        "EXTERIOR:",

        len(exterior[:10])

    )


    # ========================================================
    # 11. REMOVER DUPLICADOS GERAIS
    # ========================================================

    urls_unicas = []

    vistos = set()


    for url in urls:

        if not url_valida(url):

            print("URL BLOQUEADA:")
            print(url)

            continue


        if url not in vistos:

            vistos.add(url)

            urls_unicas.append(url)


    # ========================================================
    # 12. PRIORIZAR PDFs
    # ========================================================

    urls_unicas.sort(

        key=lambda url: (

            ".pdf" not in url.lower(),

            url.lower()

        )

    )


    print()
    print("==============================")
    print("TOTAL FONTES")
    print("==============================")

    print(

        len(urls_unicas)

    )


    urls_unicas = filtrar_urls(

        urls_unicas

    )


    print()
    print("URLS APROVADAS PARA COLETA:")

    for i, url in enumerate(

        urls_unicas,

        start=1

    ):

        print(f"{i}. {url}")

    return urls_unicas


# ============================================================
# IDENTIFICAR GRUPO
# ============================================================

def identificar_grupo_tema(tema):

    tema_normalizado = normalizar_texto(tema)

    for grupo, palavras in GRUPOS_TEMATICOS.items():

        for palavra in palavras:

            if normalizar_texto(palavra) in tema_normalizado:

                return grupo

    return "geral"
    
    

# ============================================================
# ENTENDIMENTO INICIAL DO PRODUTO
# ============================================================

def gerar_entendimento_produto(
    tema,
    grupo_principal=""
):

    print()
    print("==============================")
    print("GERANDO ENTENDIMENTO INICIAL")
    print("==============================")

    print("TEMA:", tema)
    print("GRUPO PRINCIPAL:", grupo_principal)

    if not tema:
        print("TEMA VAZIO")
        return ""

    if not grupo_principal:
        grupo_principal = (
            "Não informado. "
            "Identificar a categoria técnica "
            "mais provável a partir do tema."
        )

    prompt = f"""
Analise tecnicamente o tema para orientar uma pesquisa.

TEMA: {tema}
GRUPO: {grupo_principal}

Retorne somente 4 pontos curtos:

1. O que é.
2. Para que serve.
3. Como funciona.
4. O que deve ser pesquisado tecnicamente.

REGRAS:
- Seja extremamente objetivo.
- Máximo de 120 palavras.
- Não invente informações.
- Não use marcas ou modelos.
- Não informe números ou especificações.
- Não escreva conteúdo editorial.
"""

    try:

        inicio_ollama = time.time()

        resposta = requests.post(

            "http://localhost:11434/api/generate",

            json={

                "model": "qwen2.5:3b",

                "prompt": prompt,

                "stream": False,

                "options": {

                    "num_predict": 150,

                    "num_ctx": 2048,

                    "temperature": 0.0,

                    "think": False

                }

            },

            timeout=(30, 300)

        )

        tempo_ollama = (
            time.time()
            - inicio_ollama
        )

    except requests.exceptions.Timeout:

        print("TIMEOUT ENTENDIMENTO")
        return ""

    except requests.exceptions.ConnectionError as e:

        print("ERRO DE CONEXÃO OLLAMA")
        print(repr(e))
        return ""

    except Exception as e:

        print("ERRO ENTENDIMENTO")
        print(repr(e))
        return ""

    print()
    print("==============================")
    print("OLLAMA RESPONDEU")
    print("==============================")

    print("STATUS:", resposta.status_code)

    print(
        "TEMPO:",
        round(tempo_ollama, 1),
        "segundos"
    )

    if resposta.status_code != 200:
        return ""

    try:

        dados = resposta.json()

        entendimento = dados.get(
            "response",
            ""
        )

    except Exception as e:

        print(
            "ERRO AO LER RESPOSTA:",
            repr(e)
        )

        return ""

    entendimento = entendimento.strip()

    print()
    print("==============================")
    print("ENTENDIMENTO GERADO")
    print("==============================")

    print(
        "CARACTERES:",
        len(entendimento)
    )

    print()
    print(entendimento)

    print()
    print("==============================")
    print("FIM ENTENDIMENTO INICIAL")
    print("==============================")

    return entendimento    


# ============================================================
# CAPTURA TEXTO
# ============================================================


def limpar_texto_coletado(texto):

    cortes = [

        "Principais cidades e regiões",

        "Solicite um orçamento",

        "Entre em contato",

        "Produtos relacionados",

        "Confira Também",

        "Crime de violação de direito autoral",

        "Política de Privacidade",

        "WhatsApp",

        "Online Fale com a gente",

        "Todos os direitos reservados"
    ]

    for corte in cortes:

        if corte in texto:

            texto = texto.split(corte)[0]

    texto = re.sub(
        r"\s+",
        " ",
        texto
    )

    texto = texto.strip()

    return texto[:8000]
    

# ============================================================
# COLETAR PÁGINA
# ============================================================

def coletar_pagina(url):

    try:

        headers = {

            "User-Agent": (
                "Mozilla/5.0 "
                "(Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 "
                "(KHTML, like Gecko) "
                "Chrome/120 Safari/537.36"
            )

        }

        resposta = requests.get(

            url,
            headers=headers,
            timeout=30

        )

        if resposta.status_code != 200:

            return None

        content_type = resposta.headers.get(

            "Content-Type",
            ""

        ).lower()

        eh_pdf = (

            "pdf" in content_type

            or

            url.lower().split("?")[0].endswith(".pdf")

        )

    # ========================================================
    # 01. FUNÇÃO INTERNA PARA EXTRAIR PDF
    # ========================================================

        def extrair_pdf(bytes_pdf, origem):

            arquivo_temp = None

            try:


                arquivo_temp = tempfile.NamedTemporaryFile(

                    delete=False,
                    suffix=".pdf"

                )

                arquivo_temp.write(

                    bytes_pdf

                )

                arquivo_temp.close()

                documento = fitz.open(

                    arquivo_temp.name

                )

                texto = ""

                for pagina in documento:

                    texto += pagina.get_text("text")

                documento.close()

                os.remove(

                    arquivo_temp.name

                )

                texto = limpar_texto_coletado(

                    texto

                )

                if len(texto) < 300:

                    return None

                print()
                print("==============================")
                print("PDF UTILIZADO")
                print("==============================")
                print(origem)
                print("CARACTERES:", len(texto))

                return {
                
                    "url": origem,
                
                    "tipo": "pdf",
                
                    "qualidade": "alta",
                
                    "caracteres": len(texto),
                
                    "texto": texto
                
                }

            except Exception as erro:

                print("ERRO PDF:")
                print(erro)

                if arquivo_temp:

                    try:

                        os.remove(

                            arquivo_temp.name

                        )

                    except:

                        pass

                return None

    # ========================================================
    # 02. URL JÁ É PDF
    # ========================================================

        if eh_pdf:

            return extrair_pdf(

                resposta.content,
                url

            )

    # ========================================================
    # 03. HTML
    # ========================================================

        html = resposta.text

        soup = BeautifulSoup(

            html,
            "html.parser"

        )

    # ========================================================
    # 04. PROCURAR PDF TÉCNICO
    # ========================================================

        palavras_pdf = [

            "pdf",
            "catalog",
            "catalogo",
            "catálogo",
            "datasheet",
            "manual",
            "brochure",
            "technical",
            "specification",
            "especificacao",
            "especificação",
            "download"

        ]

        pdfs = []

        from urllib.parse import urljoin

        for link in soup.find_all(

            "a",
            href=True

        ):

            href = link["href"]

            texto_link = link.get_text(

                " ",
                strip=True

            ).lower()

            href_lower = href.lower()

            if (

                ".pdf" in href_lower

                or

                any(

                    palavra in href_lower

                    for palavra in palavras_pdf

                )

                or

                any(

                    palavra in texto_link

                    for palavra in palavras_pdf

                )

            ):

                pdfs.append(

                    urljoin(

                        url,
                        href

                    )

                )

        # remover duplicados

        pdfs = list(

            dict.fromkeys(pdfs)

        )

    # ========================================================
    # 05. TENTAR TODOS OS PDFs
    # ========================================================

        for pdf in pdfs:

            print()
            print("==============================")
            print("TESTANDO PDF ENCONTRADO")
            print("==============================")
            print(pdf)

            try:

                r = requests.get(

                    pdf,
                    headers=headers,
                    timeout=30

                )

                if r.status_code != 200:

                    continue

                resultado = extrair_pdf(

                    r.content,
                    pdf

                )

                if resultado:

                    return resultado

            except:

                pass

    # ========================================================
    # 06. REMOVER LIXO HTML
    # ========================================================

        for tag in soup([

            "script",
            "style",
            "noscript",
            "svg",
            "iframe",
            "header",
            "footer",
            "nav",
            "form"

        ]):

            tag.decompose()

        texto = soup.get_text(

            separator=" ",
            strip=True

        )

        texto = limpar_texto_coletado(

            texto

        )

        if len(texto) < 300:

            return None

        return {
        
            "url": url,
        
            "tipo": "html",
        
            "qualidade": "normal",
        
            "caracteres": len(texto),
        
            "texto": texto
        
        }

    except requests.exceptions.Timeout:

        print("TIMEOUT:")
        print(url)

        return None

    except requests.exceptions.RequestException as erro:

        print("FALHA ACESSO:")
        print(url)

        print(erro)

        return None

    except Exception as erro:

        print("ERRO COLETA:")
        print(url)

        print(erro)

        return None
        

# ============================================================
# LIMPAR REFERÊNCIAS COMERCIAIS
# ============================================================

def limpar_referencias(
    texto,
    tema="",
    tipo="html"
):


    if not texto:

        return ""



    tema_normalizado = normalizar_texto(
        tema
    )



    # ========================================================
    # 01. DEFINIR NÍVEL DE LIMPEZA
    # ========================================================

    eh_pdf = (

        tipo.lower() == "pdf"

    )



    limpeza_forte = not eh_pdf



    # =====================================
    # LIMPEZA ESPECÍFICA
    # PROTEÇÃO CONTRA INCÊNDIO
    # =====================================

    eh_construcao = any(

        termo in tema_normalizado

        for termo in [

            "firestop",
            "selagem",
            "corta fogo",
            "passagem corta fogo",
            "protecao passiva"

        ]

    )



    if eh_construcao and limpeza_forte:


        padroes = [

            r"\bCP\s*\d+\b",

            r"\bCFS\s*\d+\b",

            r"\bFFC\b",

            r"\bCKC\b",

            r"\bFS-?\d+\b",


            r"\bHilti\b",

            r"\bFirestop\b",

            r"\bPromat\b",

            r"\b3M\b",

            r"\bSika\b",

            r"\bFischer\b",

            r"\bRockwool\b",


            r"\bClasse\s+Ultimate\b",

            r"\bUltimate\b"

        ]



        for padrao in padroes:


            texto = re.sub(

                padrao,

                "",

                texto,

                flags=re.IGNORECASE

            )



    # ========================================================
    # 02. LIMPEZA COMERCIAL GERAL
    # ========================================================

    remover = [

        "entre em contato",

        "solicite orçamento",

        "fale conosco",

        "peça sua cotação",

        "comprar agora",

        "consulte disponibilidade",

        "melhor preço",

        "oferta exclusiva"

    ]



    for item in remover:


        texto = re.sub(

            item,

            "",

            texto,

            flags=re.IGNORECASE

        )



    # =====================================
    # CAMPOS COMERCIAIS
    # SOMENTE HTML
    # =====================================

    if limpeza_forte:


        texto = re.sub(

            r"(part number|codigo interno|código interno)\s*[:\-]?\s*\w+",

            "",

            texto,

            flags=re.IGNORECASE

        )



    # ========================================================
    # 03. REMOVER LINKS
    # ========================================================

    texto = re.sub(

        r"https?://\S+",

        "",

        texto

    )



    texto = re.sub(

        r"www\.\S+",

        "",

        texto

    )



    # ========================================================
    # 04. NORMALIZAR ESPAÇOS
    # ========================================================

    texto = re.sub(

        r"\s+",

        " ",

        texto

    )



    return texto.strip()


# ============================================================
# SALVAR BRUTO
# ============================================================

def salvar_bruto(
    tema,
    paginas
):


    os.makedirs(

        PASTA_DADOS,

        exist_ok=True

    )



    paginas_validas = []



    for pagina in paginas:


        if not isinstance(

            pagina,

            dict

        ):

            continue



        texto = pagina.get(

            "texto",

            ""

        )



        tipo = pagina.get(

            "tipo",

            "html"

        )



        # =====================================
        # VALIDAÇÃO
        # PDF TEM PRIORIDADE
        # =====================================


        tamanho_minimo = 200


        if tipo == "pdf":

            tamanho_minimo = 150



        if len(texto) < tamanho_minimo:

            continue



        paginas_validas.append(

            pagina

        )



    # ========================================================
    # 01. ORDENAR PDF PRIMEIRO
    # ========================================================

    paginas_validas.sort(

        key=lambda x:

        0 if x.get("tipo") == "pdf" else 1

    )



    # ========================================================
    # 02. NORMALIZAR NOME ARQUIVO
    # ========================================================

    nome = unicodedata.normalize(

        "NFD",

        tema.lower()

    )


    nome = "".join(

        c for c in nome

        if unicodedata.category(c) != "Mn"

    )


    nome = re.sub(

        r"[^a-z0-9]+",

        "_",

        nome

    )


    nome = nome.strip("_")



    arquivo = os.path.join(

        PASTA_DADOS,

        nome + ".json"

    )



    # ========================================================
    # 03. CONTADORES
    # ========================================================

    total_pdf = sum(

        1

        for p in paginas_validas

        if p.get("tipo") == "pdf"

    )


    total_html = sum(

        1

        for p in paginas_validas

        if p.get("tipo") == "html"

    )



    dados = {


        "tema": tema,


        "data_coleta": datetime.now().strftime(

            "%Y-%m-%d %H:%M:%S"

        ),


        "fontes": paginas_validas,


        "quantidade": len(

            paginas_validas

        ),


        "estatisticas": {

            "pdfs": total_pdf,

            "html": total_html

        }

    }



    try:


        with open(

            arquivo,

            "w",

            encoding="utf-8"

        ) as f:


            json.dump(

                dados,

                f,

                ensure_ascii=False,

                indent=4

            )



        print()

        print("==============================")

        print("SALVANDO DADOS BRUTOS")

        print("==============================")


        print(

            "TEMA:",

            tema

        )


        print(

            "FONTES:",

            len(paginas_validas)

        )


        print(

            "PDFS:",

            total_pdf

        )


        print(

            "HTML:",

            total_html

        )


        print()

        print("JSON SALVO:")

        print(arquivo)



    except Exception as erro:


        print()

        print("ERRO SALVANDO JSON:")

        print(erro)
        

# ============================================================
# REMOVER ACENTOS PARA NOMES DE ARQUIVO
# ============================================================

def remover_acentos(texto):

    texto = unicodedata.normalize(
        "NFD",
        texto
    )


    texto = "".join(

        c for c in texto

        if unicodedata.category(c) != "Mn"

    )


    return texto




# ============================================================
# NORMALIZAR NOME DE ARQUIVO
# ============================================================

def normalizar_nome_arquivo(texto):

    texto = remover_acentos(

        texto.lower()

    )


    texto = re.sub(

        r"[^a-z0-9]+",

        "_",

        texto

    )


    return texto.strip("_")




# ============================================================
# CARREGAR DADOS BRUTOS EXISTENTES
# ============================================================

def carregar_bruto(tema):


    nome = normalizar_nome_arquivo(

        tema

    )


    arquivo = Path(

        PASTA_DADOS

    ) / f"{nome}.json"



    if not arquivo.exists():


        print()

        print("==============================")

        print("BRUTO NÃO ENCONTRADO")

        print("==============================")

        print(arquivo)


        return []



    try:


        with open(

            arquivo,

            "r",

            encoding="utf-8"

        ) as f:


            dados = json.load(f)



        fontes = dados.get(

            "fontes",

            []

        )



        textos = []



        for item in fontes:


            if not isinstance(

                item,

                dict

            ):

                continue



            texto = item.get(

                "texto",

                ""

            )


            tipo = item.get(

                "tipo",

                "html"

            )



            tamanho_minimo = 200



            # PDF pode ter tabelas e pouco texto

            if tipo == "pdf":

                tamanho_minimo = 150



            if len(texto) < tamanho_minimo:

                continue



            textos.append({

                "texto": texto,

                "url": item.get(

                    "url",

                    ""

                ),

                "tipo": tipo

            })



    # ========================================================
    # 01. PDFS PRIMEIRO
    # ========================================================

        textos.sort(

            key=lambda x:

            0 if x.get("tipo") == "pdf" else 1

        )



        print()

        print("==============================")

        print("BRUTO ENCONTRADO")

        print("==============================")

        print(

            "ARQUIVO:",

            arquivo

        )

        print(

            "FONTES JSON:",

            len(fontes)

        )

        print(

            "TEXTOS CARREGADOS:",

            len(textos)

        )



        pdfs = sum(

            1

            for t in textos

            if t.get("tipo") == "pdf"

        )


        print(

            "PDFS:",

            pdfs

        )



        return textos



    except Exception as erro:


        print()

        print(

            "ERRO AO CARREGAR BRUTO:",

            erro

        )


        return []

    # ========================================================
    # 02. REGRAS PARA CONTEÚDO
    # ========================================================
{
    "versao": "4.0",

    "identidade": {

        "metodo": "MEAD",

        "objetivo":
        "Criar páginas estratégicas com autoridade, contexto técnico, diferenciação, profundidade semântica e narrativa humana."
    },


    "preparacao": {


        "palavra_chave": {


            "regra":

            "A palavra-chave principal informada pelo usuário é o protagonista absoluto da narrativa.",


            "prioridade":

            [

                "produto",

                "serviço",

                "fabricante",

                "distribuidor",

                "fornecedor"

            ],


            "proibido":

            [

                "substituir o protagonista",

                "transformar aplicação em produto",

                "transformar benefício em produto",

                "transformar tecnologia em produto",

                "transformar característica em produto"

            ]

        },



        "classificacao_intencao": {


            "produto":

            {

                "regra":

                "O protagonista permanece exatamente como o produto informado."

            },


            "servico":

            {

                "regra":

                "O protagonista permanece exatamente como o serviço informado."

            },


            "fabricante":

            {

                "regra":

                "O protagonista é o fabricante. Desenvolver capacidade técnica, engenharia, processos e estrutura."

            },


            "distribuidor":

            {

                "regra":

                "O protagonista é o distribuidor. Desenvolver fornecimento, atendimento e suporte."

            },


            "fornecedor":

            {

                "regra":

                "O protagonista é o fornecedor. Desenvolver disponibilidade, seleção e relacionamento."

            }

        }

    },



    "mapa_mead": {


        "protagonista":

        "Elemento principal exatamente alinhado ao tema informado.",


        "cenario":

        "Ambiente operacional completo onde o protagonista atua.",


        "problema":

        "Necessidade real que o protagonista resolve.",


        "solucao":

        "Como o protagonista atende esta necessidade.",


        "coadjuvantes":

        "Componentes, processos, materiais, tecnologias e conceitos que ajudam a explicar.",


        "aplicacoes":

        "Locais, equipamentos, sistemas ou situações reais de utilização.",


        "publico_alvo":

        "Quem compra, especifica, utiliza ou aplica.",


        "diferenciais":

        "Características técnicas reais que influenciam a escolha.",


        "evidencias_tecnicas":

        "Informações encontradas nas fontes que comprovam características, aplicações ou funcionamento.",


        "competencia_principal":

        "Maior demonstração de domínio técnico."

    },



    "densidade_estrategica": {


        "regra":

        "Cada campo deve possuir profundidade suficiente para orientar a criação do conteúdo.",


        "priorizar":

        [

            "contexto",

            "finalidade",

            "funcionamento",

            "consequências",

            "benefícios técnicos",

            "limitações",

            "relação entre elementos",

            "evidências técnicas"

        ]

    },



    "fontes_pesquisa": {


        "prioridade_fontes":

        [

            "datasheet técnico",

            "catálogo técnico PDF",

            "manual técnico PDF",

            "ficha técnica PDF",

            "documentação técnica fabricante",

            "artigos técnicos especializados",

            "páginas comerciais"

        ],



        "hierarquia_confiabilidade": {


            "nivel_1":

            [

                "PDF técnico",

                "catálogo fabricante",

                "manual engenharia"

            ],


            "nivel_2":

            [

                "site fabricante",

                "distribuidor especializado",

                "empresa técnica"

            ],


            "nivel_3":

            [

                "blogs técnicos",

                "portais industriais"

            ]

        },



        "usar_para":

        [

            "conhecimento técnico",

            "aplicações",

            "processos",

            "entendimento do mercado",

            "identificação de componentes",

            "variações do produto",

            "materiais utilizados"

        ],



        "nunca_usar_para":

        [

            "trocar protagonista",

            "alterar intenção comercial",

            "mudar segmento",

            "copiar textos",

            "inventar informações",

            "criar especificações sem fonte"

        ]

    },



    "prioridade_documentos": {


        "regra":

        "Documentos técnicos possuem prioridade sobre páginas comerciais quando existirem informações suficientes.",


        "preferir":

        [

            "datasheet",

            "manual",

            "catálogo",

            "ficha técnica",

            "desenho técnico"

        ],


        "extrair":

        [

            "materiais",

            "componentes",

            "aplicações",

            "modelos",

            "processos",

            "condições de uso"

        ]

    },



    "controle_tecnico": {


        "nunca_criar":

        [

            "certificações",

            "normas",

            "números",

            "percentuais",

            "testes",

            "garantias",

            "aprovações",

            "resultados quantitativos"

        ]

    },



    "validacao_informacao": {


        "regra":

        "Toda informação técnica deve possuir origem nas fontes coletadas ou ser uma conclusão lógica baseada nelas.",


        "bloquear":

        [

            "números técnicos sem fonte",

            "capacidade sem documento",

            "temperatura sem referência",

            "pressão sem referência",

            "normas sem comprovação",

            "certificações inexistentes"

        ]

    },



    "classificacao_fontes": {


        "pdf_tecnico":

        {

            "peso": 10

        },


        "fabricante":

        {

            "peso": 9

        },


        "distribuidor_tecnico":

        {

            "peso": 7

        },


        "artigo_tecnico":

        {

            "peso": 6

        },


        "pagina_comercial":

        {

            "peso": 4

        }

    },



    "gambiarra_narrativa": {


        "objetivo":

        "Criar textos humanos evitando repetição estrutural.",


        "permitir_inicio":

        [

            "empresa",

            "engenharia",

            "cliente",

            "processo",

            "cenário",

            "necessidade",

            "aplicação",

            "experiência",

            "benefício"

        ],


        "evitar":

        [

            "todos os parágrafos começarem pela palavra-chave",

            "todos os parágrafos começarem igual",

            "excesso de nossa empresa",

            "excesso de nossa equipe"

        ]

    },



    "ativos_narrativos": {


        "lista":

        [

            "experiencia",

            "conhecimento_tecnico",

            "engenharia",

            "qualidade",

            "seguranca",

            "suporte_tecnico",

            "atendimento_especializado",

            "confiabilidade",

            "responsabilidade"

        ],


        "regra":

        "Utilizar somente quando fizer sentido sem inventar histórico."

    },



    "regras_seo": {


        "evitar":

        [

            "texto genérico",

            "enchimento",

            "repetição",

            "promessas sem comprovação",

            "frases artificiais"

        ],


        "priorizar":

        [

            "autoridade",

            "clareza",

            "profundidade técnica",

            "SEO semântico",

            "naturalidade"

        ]

    },



    "auditoria_seven": {


        "executar":

        True,


        "verificar":

        [

            "protagonista correto",

            "segmento correto",

            "intenção comercial correta",

            "ausência de invenções",

            "cenário profundo",

            "problema real",

            "solução coerente",

            "diversidade narrativa",

            "repetição de tese",

            "repetição sintática",

            "naturalidade humana",

            "qualidade das fontes",

            "prioridade documental"

        ]

    }

}



# ============================================================
# OLLAMA - GERAR MAPA MEAD
# ============================================================

def gerar_mapa_mead(
    tema,
    textos
):

    print()
    print("==============================")
    print("DEBUG ENTRADA MAPA MEAD")
    print("==============================")

    print(
        "TIPO TEXTOS:",
        type(textos)
    )

    print(
        "QUANTIDADE:",
        len(textos)
    )

    if textos:

        print(
            "TIPO PRIMEIRO ITEM:",
            type(textos[0])
        )


    # ========================================================
    # 01. CONTEXTO MEAD
    # ========================================================

    contexto_mead_mapa = """
MAPA MEAD

O tema é o protagonista absoluto.

PROTAGONISTA:
Elemento principal exatamente alinhado ao tema.

INTENÇÃO_COMERCIAL:
O que o usuário busca resolver ou obter
ao procurar pelo tema.

DEFINIÇÃO:
O que é o protagonista.

COADJUVANTES:
Elementos técnicos que ajudam a explicar
o protagonista.

CENÁRIO:
Onde o protagonista é utilizado.

PROBLEMA:
Necessidade real relacionada ao protagonista.

SOLUÇÃO:
Como o protagonista atende essa necessidade.

LACUNAS:
Informações importantes que não foram encontradas
nas fontes.

REGRA CENTRAL:
Aplicações, componentes, instalação, manutenção,
benefícios e segmentos são secundários.
Nunca podem substituir o protagonista.

Nunca invente informações.
"""


    # ========================================================
    # 02. SELECIONAR FONTES REPRESENTATIVAS
    # ========================================================

    fontes_mapa = []


    for item in textos:

        if not isinstance(
            item,
            dict
        ):

            continue


        texto_item = item.get(
            "texto",
            ""
        )


        if not texto_item:

            continue


        url_item = item.get(
            "url",
            ""
        )


    # ========================================================
    # 03. AMOSTRA DA FONTE
    # ========================================================

        amostra = texto_item[:1000]


        fontes_mapa.append(

            f"FONTE {len(fontes_mapa) + 1}\n"
            f"URL: {url_item}\n"
            f"{amostra}"

        )


    # ========================================================
    # 04. LIMITE
    # ========================================================

        if len(fontes_mapa) >= 5:

            break


    contexto_fontes = "\n\n".join(
        fontes_mapa
    )


    # ========================================================
    # 05. LIMITE ABSOLUTO
    # ========================================================

    contexto_fontes = contexto_fontes[:5000]


    print()
    print("==============================")
    print("CONTROLE CONTEXTO MAPA")
    print("==============================")


    print(
        "MEAD MAPA:",
        len(contexto_mead_mapa)
    )


    print(
        "FONTES MAPA:",
        len(contexto_fontes)
    )


    # ========================================================
    # 06. PROMPT
    # ========================================================

    prompt = f"""
Crie um MAPA_MEAD técnico e objetivo para o tema:

{tema}

O tema deve permanecer como protagonista absoluto.

Extraia somente informações sustentadas pelas fontes.

Não escreva artigo.
Não escreva introdução.
Não escreva conclusão.
Não explique o processo.
Não invente informações.

Preencha os campos abaixo de forma curta,
objetiva e útil para uma segunda IA produzir
o conteúdo editorial.

Se um campo não possuir fundamento suficiente,
deixe-o vazio e registre a ausência em LACUNAS.

{contexto_mead_mapa}

FORMATO OBRIGATÓRIO:

**MAPA_MEAD**

**PROTAGONISTA:**
[tema]

**INTENÇÃO_COMERCIAL:**
[...]

**DEFINIÇÃO:**
[...]

**COADJUVANTES:**
- [...]
- [...]
- [...]

**CENÁRIO:**
[...]

**PROBLEMA:**
[...]

**SOLUÇÃO:**
[...]

**LACUNAS:**
[...]

FONTES:

{contexto_fontes}

Responda somente com o MAPA_MEAD.
"""


    # ========================================================
    # 07. DEBUG
    # ========================================================

    print()
    print("==============================")
    print("TEMA ENVIADO AO MAPA MEAD")
    print("==============================")


    print(
        tema
    )


    print()
    print("==============================")
    print("PROMPT MAPA MEAD")
    print("==============================")


    print(
        "PALAVRA-CHAVE:",
        tema
    )


    print(
        "TAMANHO PROMPT:",
        len(prompt)
    )


    print(
        "CONTEXTO MEAD:",
        len(contexto_mead_mapa)
    )


    print(
        "CONTEXTO FONTES:",
        len(contexto_fontes)
    )


    # ========================================================
    # 08. OLLAMA
    # ========================================================

    print()
    print("==============================")
    print("ENVIANDO MAPA MEAD PARA OLLAMA")
    print("==============================")


    num_predict = 300
    num_ctx = 8192


    print(
        "MODELO:",
        "qwen2.5:3b"
    )


    print(
        "TOKENS:",
        num_predict
    )


    print(
        "CONTEXTO:",
        num_ctx
    )


    print(
        "AGUARDANDO RESPOSTA..."
    )


    print()
    print("==============================")
    print("INICIANDO CHAMADA OLLAMA")
    print("==============================")


    inicio_ollama = time.time()


    try:

        resposta = requests.post(

            "http://localhost:11434/api/generate",

            json={

                "model":
                    "qwen2.5:3b",

                "prompt":
                    prompt,

                "stream":
                    False,

                "think":
                    False,

                "options": {

                    "num_predict":
                        num_predict,

                    "num_ctx":
                        num_ctx,

                    "temperature":
                        0.0,

                    "top_p":
                        0.9,

                    "repeat_penalty":
                        1.05

                }

            },

            timeout=(
                30,
                900
            )

        )


    except requests.exceptions.Timeout:

        print()
        print("==============================")
        print("TIMEOUT OLLAMA")
        print("==============================")


        print(
            "A IA demorou mais de 900 segundos."
        )


        print(
            "MAPA MEAD NÃO GERADO."
        )


        return ""


    except requests.exceptions.ConnectionError as e:

        print()
        print("==============================")
        print("ERRO DE CONEXÃO COM OLLAMA")
        print("==============================")


        print(
            repr(e)
        )


        return ""


    except Exception as e:

        print()
        print("==============================")
        print("ERRO NA CHAMADA OLLAMA")
        print("==============================")


        print(
            repr(e)
        )


        return ""


    fim_ollama = time.time()


    # ========================================================
    # 09. RESPOSTA HTTP
    # ========================================================

    print()
    print("==============================")
    print("OLLAMA RESPONDEU")
    print("==============================")


    print(
        "STATUS:",
        resposta.status_code
    )


    print(
        "TEMPO:",
        formatar_tempo(
            fim_ollama - inicio_ollama
        )
    )


    if resposta.status_code != 200:

        print()
        print("==============================")
        print("ERRO HTTP OLLAMA")
        print("==============================")


        print(
            resposta.text[:1000]
        )


        return ""


    # ========================================================
    # 10. EXTRAIR JSON
    # ========================================================

    try:

        dados = resposta.json()

    except Exception as e:

        print()
        print("==============================")
        print("ERRO AO LER JSON DO OLLAMA")
        print("==============================")


        print(
            repr(e)
        )


        print(
            resposta.text[:1000]
        )


        return ""


    mapa_mead = dados.get(
        "response",
        ""
    )


    # ========================================================
    # 11. VALIDAR RESPOSTA
    # ========================================================

    if not mapa_mead:

        print()
        print("==============================")
        print("OLLAMA RETORNOU VAZIO")
        print("==============================")


        return ""


    mapa_mead = str(
        mapa_mead
    ).strip()


    # ========================================================
    # 12. RESULTADO
    # ========================================================

    print()
    print("==============================")
    print("MAPA MEAD GERADO")
    print("==============================")


    print(
        "CARACTERES:",
        len(mapa_mead)
    )


    print()


    print(
        mapa_mead[:3000]
    )


    print()
    print("==============================")
    print("FIM MAPA MEAD")
    print("==============================")


    print(
        "TEMPO OLLAMA:",
        round(
            fim_ollama - inicio_ollama,
            2
        ),
        "segundos"
    )


    print(
        "STATUS OLLAMA:",
        resposta.status_code
    )


    return mapa_mead
    

# ============================================================
# VALIDAR E RECUPERAR MAPA MEAD
# ============================================================

def validar_e_recuperar_mapa_mead(
    tema,
    mapa_mead,
    textos
):

    # ========================================================
    # 01. NORMALIZAR MAPA RECEBIDO
    # ========================================================

    if isinstance(
        mapa_mead,
        dict
    ):

        mapa_mead = mapa_mead.get(
            "texto",
            ""
        )

    elif mapa_mead is not None:

        mapa_mead = str(
            mapa_mead
        )

    else:

        mapa_mead = ""


    mapa_mead = mapa_mead.strip()


    # ========================================================
    # 02. VALIDAR MAPA RECEBIDO
    # ========================================================

    if mapa_mead:

        print()
        print("==============================")
        print("VALIDANDO MAPA MEAD PARA CONTEÚDO")
        print("==============================")

        print(
            "TEMA:",
            tema
        )

        print(
            "TAMANHO:",
            len(mapa_mead)
        )


    # ========================================================
    # 03. VALIDAR PROTAGONISTA
    # ========================================================

        protagonista_valido = validar_protagonista_mead(
            mapa_mead,
            tema
        )


        print(
            "PROTAGONISTA:",
            protagonista_valido
        )


        if not protagonista_valido:

            print()
            print("==============================")
            print("MAPA MEAD REJEITADO")
            print("==============================")
            print(
                "MOTIVO: PROTAGONISTA INCOMPATÍVEL"
            )

        else:

    # ========================================================
    # 04. VALIDAR CONTEXTO TÉCNICO
    # ========================================================

            contexto_valido = validar_contexto_tecnico_mead(
                mapa_mead,
                tema
            )


            print(
                "CONTEXTO TÉCNICO:",
                contexto_valido
            )


            if contexto_valido:

                print()
                print("==============================")
                print("MAPA MEAD VALIDADO")
                print("==============================")

                return mapa_mead


    # ========================================================
    # 05. TENTAR RECUPERAR DO BANCO
    # ========================================================

    print()
    print("==============================")
    print("TENTANDO RECUPERAR MAPA MEAD")
    print("==============================")

    mapa_recuperado = obter_mapa_mead_tema(
        tema
    )


    if mapa_recuperado:

        mapa_recuperado = str(
            mapa_recuperado
        ).strip()


    # ========================================================
    # 06. VALIDAR PROTAGONISTA RECUPERADO
    # ========================================================

        protagonista_valido = validar_protagonista_mead(
            mapa_recuperado,
            tema
        )


        print(
            "PROTAGONISTA RECUPERADO:",
            protagonista_valido
        )


        if protagonista_valido:

    # ========================================================
    # 07. VALIDAR CONTEXTO RECUPERADO
    # ========================================================

            contexto_valido = validar_contexto_tecnico_mead(
                mapa_recuperado,
                tema
            )


            print(
                "CONTEXTO RECUPERADO:",
                contexto_valido
            )


            if contexto_valido:

                print()
                print("==============================")
                print("MAPA MEAD RECUPERADO")
                print("==============================")

                print(
                    "TAMANHO:",
                    len(mapa_recuperado)
                )

                return mapa_recuperado


    # ========================================================
    # 08. MAPA NÃO DISPONÍVEL
    # ========================================================

    print()
    print("==============================")
    print("MAPA MEAD NÃO DISPONÍVEL")
    print("==============================")

    print(
        "TEMA:",
        tema
    )

    return None
    
    # ========================================================
    # 09. CONFIGURAÇÃO DE FOCOS EDITORIAIS
    # ========================================================

FOCOS_EDITORIAIS = {

    "apresentacao_contexto": {
        "nome": "Apresentação e contexto",
        "ativo": True
    },

    "funcionamento": {
        "nome": "Funcionamento e domínio técnico",
        "ativo": True
    },

    "aplicacoes": {
        "nome": "Aplicações e necessidades",
        "ativo": True
    },

    "criterios": {
        "nome": "Critérios, diferenciação e confiança",
        "ativo": True
    },

    "comercial": {
        "nome": "Decisão e contexto comercial",
        "ativo": True
    },

    "informacao_tecnica": {
        "nome": "Informação técnica",
        "ativo": True
    },

    "instalacao_manutencao": {
        "nome": "Instalação e manutenção",
        "ativo": False
    },

    "beneficios": {
        "nome": "Benefícios",
        "ativo": True
    },

    "problemas_necessidades": {
        "nome": "Problemas e necessidades",
        "ativo": False
    },

    "processo_execucao": {
        "nome": "Processo de execução",
        "ativo": False
    },

    "experiencia_autoridade": {
        "nome": "Experiência e autoridade",
        "ativo": False
    },

    "seguranca": {
        "nome": "Segurança",
        "ativo": False
    },

    "atendimento_suporte": {
        "nome": "Atendimento e suporte",
        "ativo": False
    },

    "personalizacao_projeto": {
        "nome": "Personalização e projeto",
        "ativo": False
    },

    "pos_venda": {
        "nome": "Pós-venda",
        "ativo": False
    }

}



# ============================================================
# OBTER FOCOS EDITORIAIS
# ============================================================

def obter_focos_editoriais():

    focos = {}

    for chave, dados in FOCOS_EDITORIAIS.items():

        focos[chave] = bool(
            dados.get(
                "ativo",
                False
            )
        )

    return focos



# ============================================================
# FORMATAR FOCOS EDITORIAIS
# ============================================================

def formatar_focos_editoriais(
    focos=None
):

    # --------------------------------------------------------
    # USAR CONFIGURAÇÃO PADRÃO
    # --------------------------------------------------------

    if focos is None:

        focos = obter_focos_editoriais()


    linhas = []


    linhas.append(
        "FOCOS EDITORIAIS:"
    )

    linhas.append("")


    for chave, dados in FOCOS_EDITORIAIS.items():

        selecionado = bool(
            focos.get(
                chave,
                False
            )
        )


        simbolo = (
            "☑"
            if selecionado
            else "☐"
        )


        nome = dados.get(
            "nome",
            chave
        )


        linhas.append(
            f"{simbolo} {nome}"
        )


    return "\n".join(
        linhas
    )



# ============================================================
# DEBUG — MOSTRAR CONFIGURAÇÃO DOS FOCOS
# ============================================================

def mostrar_focos_editoriais():

    focos = obter_focos_editoriais()


    print()
    print(
        "========================================"
    )
    print(
        "FOCOS EDITORIAIS"
    )
    print(
        "========================================"
    )


    print(
        formatar_focos_editoriais(
            focos
        )
    )


    print(
        "========================================"
    )


    return focos    



# ============================================================
# CONTROLAR CONTEXTO PARA IA
# ============================================================

def controlar_contexto_ia(
    textos,
    limite_total=16000
):

    contexto = ""

    if not textos:
        return contexto


    # ========================================================
    # 01. PERCORRER FONTES
    # ========================================================

    for item in textos:

        if len(contexto) >= limite_total:
            break


        # ---------------------------------
        # TEXTO DA FONTE
        # ---------------------------------

        if isinstance(item, dict):

            texto = item.get(
                "texto",
                ""
            )

        elif isinstance(item, str):

            texto = item

        else:

            continue


        if not texto:
            continue


        # ---------------------------------
        # LIMITE RESTANTE
        # ---------------------------------

        restante = (
            limite_total
            - len(contexto)
        )


        if restante <= 0:
            break


        # ---------------------------------
        # ADICIONAR TEXTO
        # ---------------------------------

        trecho = texto[:restante]


        contexto += (
            "\n\n"
            + trecho
        )


    # ========================================================
    # 02. NORMALIZAÇÃO
    # ========================================================

    contexto = contexto.strip()


    # ========================================================
    # 03. DEBUG
    # ========================================================

    print()
    print("==============================")
    print("CONTROLE CONTEXTO IA")
    print("==============================")

    print(
        "LIMITE:",
        limite_total
    )

    print(
        "CARACTERES ENVIADOS:",
        len(contexto)
    )

    print("==============================")


    return contexto
    

# ============================================================
# CONTROLAR ABERTURAS DO PROTAGONISTA
# ============================================================

def controlar_aberturas_protagonista(texto, tema, limite=3):



    print()
    print("=" * 60)
    print("CONTROLE DE ABERTURAS DO PROTAGONISTA")
    print("=" * 60)

    print("TEMA:", tema)
    print("LIMITE DE ABERTURAS DIRETAS:", limite)

    if not texto or not texto.strip():

        print("TEXTO VAZIO")
        print("CONTROLE NÃO EXECUTADO")

        return texto

    # --------------------------------------------------------
    # NORMALIZAÇÃO
    # --------------------------------------------------------

    def normalizar(texto_local):

        texto_local = texto_local.lower().strip()

        texto_local = unicodedata.normalize(
            "NFD",
            texto_local
        )

        texto_local = "".join(
            caractere
            for caractere in texto_local
            if unicodedata.category(caractere) != "Mn"
        )

        texto_local = re.sub(
            r"\s+",
            " ",
            texto_local
        )

        return texto_local

    tema_normalizado = normalizar(tema)

    # --------------------------------------------------------
    # SEPARAR TEXTO DOS MARCADORES
    # --------------------------------------------------------

    linhas = texto.splitlines()

    paragrafos = []

    bloco_atual = None

    for linha in linhas:

        linha_limpa = linha.strip()

        if not linha_limpa:
            continue

        # ----------------------------------------------------
        # IGNORA MARCADORES DE BLOCO
        # ----------------------------------------------------

        if re.match(
            r"^\*\*BLOCO\s+\d+\*\*$",
            linha_limpa,
            re.IGNORECASE
        ):

            bloco_atual = linha_limpa

            continue

        # ----------------------------------------------------
        # IGNORA LINHAS DE SEGMENTOS
        # ----------------------------------------------------

        if linha_limpa.startswith("- "):
            continue

        # ----------------------------------------------------
        # IGNORA TÍTULOS / MARCADORES
        # ----------------------------------------------------

        if (
            linha_limpa.startswith("**")
            and linha_limpa.endswith("**")
        ):
            continue

        # ----------------------------------------------------
        # PARÁGRAFO
        # ----------------------------------------------------

        paragrafos.append({
            "texto": linha_limpa,
            "bloco": bloco_atual
        })

    print()
    print("PARÁGRAFOS ANALISADOS:", len(paragrafos))

    # --------------------------------------------------------
    # IDENTIFICAR ABERTURAS DIRETAS
    # --------------------------------------------------------

    ocorrencias = []

    for indice, item in enumerate(paragrafos):

        texto_paragrafo = item["texto"]

        inicio_normalizado = normalizar(
            texto_paragrafo
        )

        # Aceita:
        #
        # A bomba centrifuga...
        # A bomba centrífuga...
        #
        # Também permite artigo masculino/feminino
        # conforme o tema.

        if (
            inicio_normalizado.startswith(
                tema_normalizado
            )
            or inicio_normalizado.startswith(
                "a " + tema_normalizado
            )
            or inicio_normalizado.startswith(
                "o " + tema_normalizado
            )
        ):

            ocorrencias.append({
                "indice": indice,
                "bloco": item["bloco"],
                "texto": texto_paragrafo
            })

    # --------------------------------------------------------
    # PRINT DAS OCORRÊNCIAS
    # --------------------------------------------------------

    print()
    print("ABERTURAS DIRETAS ENCONTRADAS:", len(ocorrencias))

    if ocorrencias:

        for numero, ocorrencia in enumerate(
            ocorrencias,
            start=1
        ):

            print()
            print(
                f"ABERTURA {numero}"
            )

            print(
                "PARÁGRAFO:",
                ocorrencia["indice"] + 1
            )

            print(
                "BLOCO:",
                ocorrencia["bloco"]
            )

            print(
                "TEXTO:",
                ocorrencia["texto"][:180]
            )

    else:

        print(
            "NENHUMA ABERTURA DIRETA ENCONTRADA"
        )

    # --------------------------------------------------------
    # VERIFICAR LIMITE
    # --------------------------------------------------------

    if len(ocorrencias) <= limite:

        print()
        print("STATUS: DENTRO DO LIMITE")
        print(
            f"ABERTURAS: {len(ocorrencias)} / {limite}"
        )

        print(
            "NENHUMA CORREÇÃO NECESSÁRIA"
        )

        print("=" * 60)

        return texto

    # --------------------------------------------------------
    # EXISTE EXCESSO
    # --------------------------------------------------------

    print()
    print("STATUS: EXCESSO DE ABERTURAS")
    print(
        f"ABERTURAS ENCONTRADAS: {len(ocorrencias)}"
    )
    print(
        f"LIMITE PERMITIDO: {limite}"
    )

    # --------------------------------------------------------
    # SELECIONAR SOMENTE OS EXCEDENTES
    # --------------------------------------------------------

    excedentes = ocorrencias[limite:]

    print()
    print(
        "PARÁGRAFOS QUE SERÃO CORRIGIDOS:",
        len(excedentes)
    )

    for numero, ocorrencia in enumerate(
        excedentes,
        start=1
    ):

        print()
        print(
            f"EXCEDENTE {numero}"
        )

        print(
            "PARÁGRAFO:",
            ocorrencia["indice"] + 1
        )

        print(
            "BLOCO:",
            ocorrencia["bloco"]
        )

        print(
            "TEXTO:",
            ocorrencia["texto"][:200]
        )

    # --------------------------------------------------------
    # MONTAR PEDIDO DE CORREÇÃO
    # --------------------------------------------------------

    instrucoes = []

    instrucoes.append(
        "VARIAÇÃO EDITORIAL DE ABERTURAS"
    )

    instrucoes.append(
        f"Palavra-chave: {tema}"
    )

    instrucoes.append(
        "O texto já está pronto e não deve ser reescrito integralmente."
    )

    instrucoes.append(
        "Corrija SOMENTE os parágrafos fornecidos."
    )

    instrucoes.append(
        f"Não inicie o parágrafo com '{tema}'."
    )

    instrucoes.append(
        "A palavra-chave pode continuar aparecendo naturalmente dentro do parágrafo."
    )

    instrucoes.append(
        "Preserve integralmente o significado técnico."
    )

    instrucoes.append(
        "Não invente informações."
    )

    instrucoes.append(
        "Não transforme o texto em propaganda."
    )

    instrucoes.append(
        "Use uma abertura natural e diferente."
    )

    instrucoes.append(
        "Pode iniciar pelo cenário, operação, necessidade, aplicação, engenharia, processo, benefício, cliente ou outro elemento coerente."
    )

    instrucoes.append(
        "Não utilize outra fórmula repetitiva."
    )

    instrucoes.append(
        "Retorne somente os parágrafos corrigidos, na mesma ordem."
    )

    instrucoes.append(
        "Separe cada parágrafo corrigido por uma linha em branco."
    )

    instrucoes.append("")
    instrucoes.append("PARÁGRAFOS PARA CORRIGIR:")

    for numero, ocorrencia in enumerate(
        excedentes,
        start=1
    ):

        instrucoes.append(
            f"[PARÁGRAFO {numero}]"
        )

        instrucoes.append(
            ocorrencia["texto"]
        )

        instrucoes.append("")

    prompt_correcao = "\n".join(
        instrucoes
    )

    print()
    print("=" * 60)
    print("ENVIANDO CORREÇÃO DE ABERTURAS PARA OLLAMA")
    print("=" * 60)

    print(
        "MODELO: qwen2.5:3b"
    )

    print(
        "PROMPT:",
        len(prompt_correcao),
        "caracteres"
    )

    inicio_ollama = time.time()

    try:

        resposta = requests.post(

            "http://localhost:11434/api/generate",

            json={

                "model": "qwen2.5:3b",

                "prompt": prompt_correcao,

                "stream": False,

                "think": False,

                "options": {

                    "num_predict": 300,

                    "num_ctx": 4096,

                    "temperature": 0.0

                }

            },

            timeout=(30, 180)

        )

        tempo_ollama = (
            time.time() - inicio_ollama
        )

        print()
        print(
            "OLLAMA STATUS:",
            resposta.status_code
        )

        print(
            "TEMPO:",
            round(tempo_ollama, 1),
            "segundos"
        )

        if resposta.status_code != 200:

            print(
                "ERRO: OLLAMA NÃO RETORNOU 200"
            )

            print(
                "TEXTO ORIGINAL SERÁ PRESERVADO"
            )

            print("=" * 60)

            return texto

        dados = resposta.json()

        texto_corrigido = dados.get(
            "response",
            ""
        ).strip()

        print(
            "RESPOSTA OLLAMA:",
            len(texto_corrigido),
            "caracteres"
        )

        if not texto_corrigido:

            print(
                "RESPOSTA VAZIA"
            )

            print(
                "TEXTO ORIGINAL SERÁ PRESERVADO"
            )

            print("=" * 60)

            return texto

    except Exception as erro:

        print()
        print(
            "ERRO NA CORREÇÃO:",
            erro
        )

        print(
            "TEXTO ORIGINAL SERÁ PRESERVADO"
        )

        print("=" * 60)

        return texto

    # --------------------------------------------------------
    # EXTRAIR PARÁGRAFOS CORRIGIDOS
    # --------------------------------------------------------

    corrigidos = [

        p.strip()

        for p in re.split(
            r"\n\s*\n",
            texto_corrigido
        )

        if p.strip()

    ]

    print()
    print(
        "PARÁGRAFOS CORRIGIDOS RECEBIDOS:",
        len(corrigidos)
    )

    # --------------------------------------------------------
    # SEGURANÇA
    # --------------------------------------------------------

    if len(corrigidos) != len(excedentes):

        print()
        print(
            "ATENÇÃO: QUANTIDADE DE PARÁGRAFOS DIFERENTE"
        )

        print(
            "ESPERADO:",
            len(excedentes)
        )

        print(
            "RECEBIDO:",
            len(corrigidos)
        )

        print(
            "CORREÇÃO DESCARTADA"
        )

        print(
            "TEXTO ORIGINAL SERÁ PRESERVADO"
        )

        print("=" * 60)

        return texto

    # --------------------------------------------------------
    # SUBSTITUIR SOMENTE OS PARÁGRAFOS EXCEDENTES
    # --------------------------------------------------------

    mapa_correcoes = {}

    for ocorrencia, novo_texto in zip(
        excedentes,
        corrigidos
    ):

        mapa_correcoes[
            ocorrencia["indice"]
        ] = novo_texto

    novo_texto_final = texto

    # --------------------------------------------------------
    # RECONSTRUIR TEXTO PRESERVANDO BLOCOS
    # --------------------------------------------------------

    linhas_originais = texto.splitlines()

    resultado = []

    indice_paragrafo = 0

    for linha in linhas_originais:

        linha_limpa = linha.strip()

        # ----------------------------------------------------
        # LINHAS VAZIAS
        # ----------------------------------------------------

        if not linha_limpa:

            resultado.append("")

            continue

        # ----------------------------------------------------
        # MARCADORES / TÍTULOS / SEGMENTOS
        # ----------------------------------------------------

        if (
            re.match(
                r"^\*\*BLOCO\s+\d+\*\*$",
                linha_limpa,
                re.IGNORECASE
            )
            or linha_limpa.startswith("**")
            or linha_limpa.startswith("- ")
        ):

            resultado.append(linha)

            continue

        # ----------------------------------------------------
        # PARÁGRAFO
        # ----------------------------------------------------

        if indice_paragrafo in mapa_correcoes:

            resultado.append(
                mapa_correcoes[
                    indice_paragrafo
                ]
            )

        else:

            resultado.append(linha)

        indice_paragrafo += 1

    novo_texto_final = "\n".join(
        resultado
    )

    # --------------------------------------------------------
    # NOVA CONTAGEM
    # --------------------------------------------------------

    paragrafos_novos = []

    for linha in novo_texto_final.splitlines():

        linha_limpa = linha.strip()

        if not linha_limpa:
            continue

        if re.match(
            r"^\*\*BLOCO\s+\d+\*\*$",
            linha_limpa,
            re.IGNORECASE
        ):
            continue

        if linha_limpa.startswith("- "):
            continue

        if (
            linha_limpa.startswith("**")
            and linha_limpa.endswith("**")
        ):
            continue

        paragrafos_novos.append(
            linha_limpa
        )

    novas_ocorrencias = []

    for indice, paragrafo in enumerate(
        paragrafos_novos
    ):

        inicio_normalizado = normalizar(
            paragrafo
        )

        if (
            inicio_normalizado.startswith(
                tema_normalizado
            )
            or inicio_normalizado.startswith(
                "a " + tema_normalizado
            )
            or inicio_normalizado.startswith(
                "o " + tema_normalizado
            )
        ):

            novas_ocorrencias.append(
                indice
            )

    # --------------------------------------------------------
    # RESULTADO FINAL
    # --------------------------------------------------------

    print()
    print("=" * 60)
    print("RESULTADO DO CONTROLE")
    print("=" * 60)

    print(
        "ANTES:",
        len(ocorrencias),
        "aberturas diretas"
    )

    print(
        "DEPOIS:",
        len(novas_ocorrencias),
        "aberturas diretas"
    )

    print(
        "LIMITE:",
        limite
    )

    if len(novas_ocorrencias) <= limite:

        print(
            "STATUS FINAL: APROVADO"
        )

    else:

        print(
            "STATUS FINAL: AINDA ACIMA DO LIMITE"
        )

        print(
            "ATENÇÃO: NÃO SERÁ FEITA NOVA CHAMADA AO OLLAMA"
        )

    print("=" * 60)

    return novo_texto_final



# ============================================================
# AUDITAR E CORRIGIR ABERTURAS REPETIDAS
# ============================================================

def auditar_e_corrigir_aberturas(
    conteudo,
    tema
):

    print()
    print("==============================")
    print("AUDITORIA DE ABERTURAS")
    print("==============================")



    # ========================================================
    # 01. NORMALIZAÇÃO
    # ========================================================

    tema_limpo = (
        str(tema)
        .strip()
        .lower()
    )

    # ========================================================
    # 02. EXTRAIR PARÁGRAFOS
    # ========================================================

    paragrafos = [

        p.strip()

        for p in re.split(
            r"\n\s*\n",
            conteudo
        )

        if p.strip()

    ]

    # ========================================================
    # 03. CONTAR ABERTURAS DIRETAS
    # ========================================================

    aberturas_diretas = []

    for indice, paragrafo in enumerate(paragrafos):

        inicio = paragrafo.lower()

        # Remove pequenas marcações
        inicio = re.sub(
            r"^[\*\#\-\s]+",
            "",
            inicio
        ).strip()

        # ---------------------------------
        # Verificar início com o tema
        # ---------------------------------

        if inicio.startswith(
            tema_limpo
        ):

            aberturas_diretas.append(
                indice
            )

    # ========================================================
    # 04. RESULTADO
    # ========================================================

    quantidade_aberturas = len(
        aberturas_diretas
    )

    print(
        "PALAVRA-CHAVE:",
        tema
    )

    print(
        "ABERTURAS DIRETAS:",
        quantidade_aberturas
    )

    print(
        "LIMITE PERMITIDO:",
        3
    )

    print(
        "POSIÇÕES:",
        [
            x + 1
            for x in aberturas_diretas
        ]
    )

    # ========================================================
    # 05. DENTRO DO LIMITE
    # ========================================================

    if quantidade_aberturas <= 3:

        print()
        print(
            "ABERTURAS DENTRO DO LIMITE"
        )

        print(
            "CORREÇÃO NECESSÁRIA:",
            False
        )

        return conteudo

    # ========================================================
    # 06. LIMITE EXCEDIDO
    # ========================================================

    print()
    print("==============================")
    print("ABERTURAS EXCEDERAM O LIMITE")
    print("==============================")

    print(
        "ENCONTRADAS:",
        quantidade_aberturas
    )

    print(
        "PERMITIDAS:",
        3
    )

    print(
        "INICIANDO CORREÇÃO IA..."
    )

    # ========================================================
    # 07. PROMPT DE CORREÇÃO
    # ========================================================

    prompt_correcao = f"""
Você é um editor técnico especializado em naturalidade editorial.

Revise o conteúdo abaixo.

TEMA:
{tema}

==================================================
REGRA PRINCIPAL — ABERTURA DOS PARÁGRAFOS
==================================================

Existe uma regra editorial obrigatória:

No máximo 3 dos 15 parágrafos podem começar
diretamente com a palavra-chave:

"{tema}"

Os demais parágrafos devem começar naturalmente
por outros elementos da narrativa.

A palavra-chave NÃO está proibida no início dos
parágrafos.

Ela pode aparecer no início em até 3 parágrafos.

O objetivo é apenas impedir repetição excessiva.

Exemplos de entradas naturais:

o cenário;
a operação;
a necessidade;
o processo;
a aplicação;
um critério técnico;
uma característica;
uma consequência;
uma condição operacional;
a experiência;
o contexto industrial;
o sistema;
o ambiente;
a demanda;
o funcionamento;
a utilização.

==================================================
VARIAÇÃO NARRATIVA
==================================================

Evite também repetir excessivamente a mesma
estrutura de abertura.

Não substitua uma repetição por outra.

Evite sequências repetitivas como:

"Em ambientes..."
"Em ambientes..."
"Em ambientes..."

"Durante..."
"Durante..."
"Durante..."

"Para..."
"Para..."
"Para..."

"Quando..."
"Quando..."
"Quando..."

"A empresa..."
"A empresa..."
"A empresa..."

"Nossa empresa..."
"Nossa empresa..."
"Nossa empresa..."

A construção dos parágrafos deve parecer natural
e escrita por um especialista humano.

Varie:

- sujeito;
- perspectiva;
- posição da palavra-chave;
- estrutura sintática;
- ponto de entrada da informação;
- relação entre contexto e explicação técnica.

==================================================
PRESERVAÇÃO DO CONTEÚDO
==================================================

A revisão é exclusivamente editorial.

NÃO:

- altere o significado técnico;
- remova informações técnicas;
- invente informações;
- crie informações técnicas novas;
- altere especificações;
- altere números;
- altere características técnicas;
- altere aplicações técnicas;
- altere afirmações sustentadas pelas fontes;
- substitua o protagonista;
- transforme componentes em protagonistas;
- transforme manutenção em protagonista;
- transforme instalação em protagonista;
- transforme aplicações em protagonistas.

A palavra-chave continua sendo o protagonista
semântico da página.

==================================================
ESTRUTURA OBRIGATÓRIA
==================================================

Preserve exatamente:

5 blocos;

3 parágrafos em cada bloco;

15 parágrafos no total;

12 segmentos de aplicação.

NÃO:

- crie novos parágrafos;
- remova parágrafos;
- una parágrafos;
- divida parágrafos;
- altere a ordem dos blocos;
- altere os títulos dos blocos;
- altere a quantidade de segmentos;
- altere o conteúdo dos segmentos sem necessidade.

==================================================
EXTENSÃO
==================================================

Não reduza o conteúdo.

Não transforme parágrafos longos em
parágrafos curtos.

Não remova informações apenas para
corrigir uma abertura.

Quando uma abertura precisar ser alterada,
modifique somente a construção inicial
necessária para eliminar a repetição.

Preserve o restante do parágrafo.

==================================================
CRITÉRIO DE NATURALIDADE
==================================================

Antes de finalizar, verifique internamente:

1. Existem no máximo 3 parágrafos iniciados
   diretamente por "{tema}"?

2. As demais aberturas possuem variedade?

3. Existe alguma sequência excessiva de
   estruturas iniciadas por "Em", "Durante",
   "Para", "Quando", "A empresa" ou outra
   fórmula repetitiva?

4. A palavra-chave continua sendo o
   protagonista semântico?

5. O texto continua tecnicamente equivalente
   ao conteúdo original?

6. Foram preservados exatamente:
   - 5 blocos;
   - 15 parágrafos;
   - 12 segmentos?

Somente finalize quando todas essas condições
forem atendidas.

==================================================
SAÍDA
==================================================

Retorne somente o conteúdo corrigido.

Não explique as alterações.

Não apresente comentários editoriais.

Não mostre esta instrução.

Não mostre análise.

==================================================
CONTEÚDO
==================================================

{conteudo}
"""

    # ========================================================
    # 08. CHAMADA OLLAMA
    # ========================================================

    inicio_correcao = time.time()

    try:

        resposta = requests.post(

            "http://localhost:11434/api/generate",

            json={

                "model":
                    "qwen2.5:3b",

                "prompt":
                    prompt_correcao,

                "stream":
                    False,

                "think":
                    False,

                "options": {

                    "num_predict":
                        800,

                    "num_ctx":
                        8192,

                    "temperature":
                        0.0,

                    "top_p":
                        0.9,

                    "repeat_penalty":
                        1.08

                }

            },

            timeout=(
                30,
                900
            )

        )

    except requests.exceptions.Timeout:

        print()
        print(
            "TIMEOUT NA CORREÇÃO"
        )

        print(
            "MANTENDO CONTEÚDO ORIGINAL"
        )

        return conteudo

    except requests.exceptions.ConnectionError as e:

        print()
        print(
            "ERRO DE CONEXÃO NA CORREÇÃO"
        )

        print(
            repr(e)
        )

        print(
            "MANTENDO CONTEÚDO ORIGINAL"
        )

        return conteudo

    except Exception as e:

        print()
        print(
            "ERRO NA CORREÇÃO"
        )

        print(
            repr(e)
        )

        print(
            "MANTENDO CONTEÚDO ORIGINAL"
        )

        return conteudo

    # ========================================================
    # 09. RECEBER RESULTADO
    # ========================================================

    if resposta.status_code != 200:

        print()
        print(
            "ERRO HTTP NA CORREÇÃO:",
            resposta.status_code
        )

        print(
            "MANTENDO CONTEÚDO ORIGINAL"
        )

        return conteudo

    try:

        dados = resposta.json()

        novo_conteudo = dados.get(
            "response",
            ""
        ).strip()

    except Exception:

        novo_conteudo = ""

    # ========================================================
    # 10. VALIDAR RETORNO
    # ========================================================

    if not novo_conteudo:

        print()
        print(
            "CORREÇÃO IA RETORNOU VAZIO"
        )

        print(
            "MANTENDO CONTEÚDO ORIGINAL"
        )

        return conteudo

    tempo_correcao = (
        time.time()
        - inicio_correcao
    )

    print()
    print("==============================")
    print("CORREÇÃO IA FINALIZADA")
    print("==============================")

    print(
        "TEMPO:",
        round(
            tempo_correcao,
            1
        ),
        "segundos"
    )

    print(
        "CARACTERES ANTES:",
        len(conteudo)
    )

    print(
        "CARACTERES DEPOIS:",
        len(novo_conteudo)
    )

    # ========================================================
    # 11. AUDITORIA NOVAMENTE
    # ========================================================

    paragrafos_corrigidos = [

        p.strip()

        for p in re.split(
            r"\n\s*\n",
            novo_conteudo
        )

        if p.strip()

    ]

    aberturas_corrigidas = []

    for indice, paragrafo in enumerate(
        paragrafos_corrigidos
    ):

        inicio = paragrafo.lower()

        inicio = re.sub(
            r"^[\*\#\-\s]+",
            "",
            inicio
        ).strip()

        if inicio.startswith(
            tema_limpo
        ):

            aberturas_corrigidas.append(
                indice
            )

    quantidade_corrigida = len(
        aberturas_corrigidas
    )

    print()
    print("==============================")
    print("RESULTADO PÓS-CORREÇÃO")
    print("==============================")

    print(
        "ABERTURAS ANTES:",
        quantidade_aberturas
    )

    print(
        "ABERTURAS DEPOIS:",
        quantidade_corrigida
    )

    print(
        "LIMITE:",
        3
    )

    print(
        "CORREÇÃO FUNCIONOU:",
        quantidade_corrigida <= 3
    )

    print(
        "POSIÇÕES FINAIS:",
        [
            x + 1
            for x in aberturas_corrigidas
        ]
    )

    # ========================================================
    # 12. SEGURANÇA
    # ========================================================

    # Se a IA devolveu algo ainda pior,
    # não substituímos o conteúdo original.

    if quantidade_corrigida > 3:

        print()
        print(
            "ATENÇÃO: LIMITE AINDA EXCEDIDO"
        )

        print(
            "CONTEÚDO ORIGINAL SERÁ PRESERVADO"
        )

        return conteudo

    print()
    print(
        "CONTEÚDO CORRIGIDO ACEITO"
    )

    return novo_conteudo
    




# ============================================================
# REVISAR NATURALIDADE DO CONTEÚDO
# ============================================================

def revisar_naturalidade_conteudo(
    tema,
    conteudo
):

    print()
    print("==============================")
    print("INICIANDO REVISÃO DE NATURALIDADE")
    print("==============================")

    print(
        "TEMA:",
        tema
    )

    print(
        "CONTEÚDO ORIGINAL:",
        len(conteudo),
        "caracteres"
    )

    # ========================================================
    # 01. VALIDAR CONTEÚDO
    # ========================================================

    if not conteudo:

        print()
        print("==============================")
        print("CONTEÚDO VAZIO PARA REVISÃO")
        print("==============================")

        return ""

    # ========================================================
    # 02. PROMPT ENXUTO
    # ========================================================

    prompt = f"""
Você é um editor técnico experiente.

Faça uma revisão editorial do conteúdo abaixo.

OBJETIVO:
Melhorar a naturalidade da narrativa sem modificar
o conteúdo técnico.

TEMA PRINCIPAL:
{tema}

REGRAS:

1. "{tema}" continua sendo o protagonista absoluto
da narrativa.

2. Revise principalmente as aberturas dos parágrafos
e as transições entre ideias.

3. Evite que os parágrafos comecem repetidamente
com "{tema}" ou com a mesma estrutura sintática.

4. Varie naturalmente o ponto de entrada das frases:
contexto, situação, característica, consequência,
aplicação, necessidade, processo ou resultado.

5. Não force sinônimos para "{tema}" quando isso
prejudicar a precisão técnica.

6. Preserve integralmente os fatos e informações
técnicas existentes.

NÃO ALTERE:
- números;
- especificações;
- características;
- aplicações;
- processos;
- relações de causa e efeito;
- informações técnicas;
- sentido das afirmações.

7. Não introduza informações novas.

8. Não remova informações existentes.

9. Não resuma o conteúdo.

10. Preserve exatamente a estrutura existente:
- 5 blocos;
- 15 parágrafos;
- 12 segmentos de aplicação;
- títulos;
- ordem dos blocos.

11. Componentes, manutenção, instalação e aplicações
devem continuar como elementos de apoio.
O protagonista permanece sendo "{tema}".

12. Faça somente as alterações necessárias para que
o texto pareça escrito de forma natural por um
especialista humano.

IMPORTANTE:
Não tente reconstruir o texto.
Não altere sua estrutura.
Não transforme o conteúdo em outro texto.
Apenas refine as construções que apresentam
repetição ou artificialidade.

RETORNE SOMENTE O TEXTO REVISADO.

==================================================
CONTEÚDO
==================================================

{conteudo}
"""

    # ========================================================
    # 03. CONTROLE
    # ========================================================

    num_predict = 2600
    num_ctx = 12288

    print()
    print("==============================")
    print("CONTROLE REVISÃO")
    print("==============================")

    print(
        "PROMPT:",
        len(prompt),
        "caracteres"
    )

    print(
        "TOKENS:",
        num_predict
    )

    print(
        "CONTEXTO:",
        num_ctx
    )

    print(
        "MODELO:",
        "qwen2.5:3b"
    )

    # ========================================================
    # 04. ENVIAR PARA OLLAMA
    # ========================================================

    print()
    print("==============================")
    print("ENVIANDO REVISÃO PARA OLLAMA")
    print("==============================")

    inicio_ollama = time.time()

    try:

        resposta = requests.post(

            "http://localhost:11434/api/generate",

            json={

                "model":
                    "qwen2.5:3b",

                "prompt":
                    prompt,

                "stream":
                    True,

                "think":
                    False,

                "options": {

                    "num_predict":
                        num_predict,

                    "num_ctx":
                        num_ctx,

                    "temperature":
                        0.1,

                    "top_p":
                        0.9,

                    "repeat_penalty":
                        1.08

                }

            },

            stream=True,

            timeout=(
                30,
                900
            )

        )

    except requests.exceptions.Timeout:

        print()
        print("==============================")
        print("TIMEOUT OLLAMA — REVISÃO")
        print("==============================")

        return ""

    except requests.exceptions.ConnectionError as e:

        print()
        print("==============================")
        print("ERRO DE CONEXÃO OLLAMA — REVISÃO")
        print("==============================")

        print(
            repr(e)
        )

        return ""

    except Exception as e:

        print()
        print("==============================")
        print("ERRO OLLAMA — REVISÃO")
        print("==============================")

        print(
            repr(e)
        )

        return ""

    print()
    print("==============================")
    print("OLLAMA RESPONDENDO — REVISÃO")
    print("==============================")

    print(
        "STATUS:",
        resposta.status_code
    )

    if resposta.status_code != 200:

        try:

            print(
                resposta.text[:1000]
            )

        except Exception:

            pass

        return ""

    # ========================================================
    # 05. RECEBER STREAM
    # ========================================================

    conteudo_revisado = ""

    ultimo_print = time.time()

    try:

        for linha in resposta.iter_lines():

            if not linha:

                continue

            try:

                dados = json.loads(
                    linha.decode(
                        "utf-8"
                    )
                )

            except Exception:

                continue

            trecho = dados.get(
                "response",
                ""
            )

            if trecho:

                conteudo_revisado += trecho

                agora = time.time()

                if agora - ultimo_print >= 30:

                    print()
                    print("==============================")
                    print("STATUS REVISÃO OLLAMA")
                    print("==============================")

                    print(
                        "TEMPO DECORRIDO:",
                        round(
                            agora - inicio_ollama,
                            1
                        ),
                        "s"
                    )

                    print(
                        "CARACTERES:",
                        len(
                            conteudo_revisado
                        )
                    )

                    ultimo_print = agora

            if dados.get(
                "done",
                False
            ):

                break

    except Exception as e:

        print()
        print("==============================")
        print("ERRO DURANTE STREAM — REVISÃO")
        print("==============================")

        print(
            repr(e)
        )

        return ""

    fim_ollama = time.time()

    # ========================================================
    # 06. LIMPAR RESULTADO
    # ========================================================

    conteudo_revisado = (
        conteudo_revisado
        .strip()
    )

    print()
    print("==============================")
    print("REVISÃO DE NATURALIDADE FINALIZADA")
    print("==============================")

    print(
        "TEMPO TOTAL:",
        round(
            fim_ollama - inicio_ollama,
            1
        ),
        "segundos"
    )

    print(
        "CONTEÚDO ORIGINAL:",
        len(conteudo),
        "caracteres"
    )

    print(
        "CONTEÚDO REVISADO:",
        len(conteudo_revisado),
        "caracteres"
    )

    # ========================================================
    # 07. VALIDAR RETORNO
    # ========================================================

    if not conteudo_revisado:

        print()
        print("==============================")
        print("REVISÃO NÃO RETORNOU CONTEÚDO")
        print("==============================")

        print(
            "CONTEÚDO ORIGINAL SERÁ PRESERVADO"
        )

        return ""

    # ========================================================
    # 08. MOSTRAR RESULTADO
    # ========================================================

    print()
    print("==============================")
    print("CONTEÚDO REVISADO RECEBIDO")
    print("==============================")

    print(
        conteudo_revisado[:3000]
    )

    print()
    print("==============================")
    print("FIM REVISÃO DE NATURALIDADE")
    print("==============================")

    return conteudo_revisado



# ============================================================
# CORRIGIR CONTEÚDO COM QWEN3
# ============================================================

def corrigir_conteudo_com_qwen3(
    tema,
    conteudo,
    mapa_mead,
    paragrafos_por_bloco,
    segmentos_encontrados
):

    print()
    print("==============================")
    print("INICIANDO CORREÇÃO COM QWEN3")
    print("==============================")


    if not conteudo:

        print(
            "CONTEÚDO VAZIO"
        )

        return ""


    # ========================================================
    # 01. DIAGNÓSTICO
    # ========================================================

    problemas = []


    # ========================================================
    # 02. VERIFICAR BLOCOS
    # ========================================================

    blocos_encontrados = 0


    marcadores_blocos = [

        "**BLOCO 1**",
        "**BLOCO 2**",
        "**BLOCO 3**",
        "**BLOCO 4**",
        "**BLOCO 5**"

    ]


    for marcador in marcadores_blocos:

        if marcador in conteudo:

            blocos_encontrados += 1


    if blocos_encontrados != 5:

        problemas.append(
            "A estrutura deve possuir exatamente "
            "5 blocos."
        )


    # ========================================================
    # 03. VERIFICAR PARÁGRAFOS
    # ========================================================

    total_paragrafos = sum(
        paragrafos_por_bloco.values()
    )


    for i in range(
        1,
        6
    ):

        quantidade = paragrafos_por_bloco.get(
            i,
            0
        )


        if quantidade != 3:

            problemas.append(
                f"O BLOCO {i} possui "
                f"{quantidade} parágrafos. "
                f"Deve possuir exatamente 3."
            )


    if total_paragrafos != 15:

        problemas.append(
            "O conteúdo deve possuir "
            "exatamente 15 parágrafos."
        )


    # ========================================================
    # 04. VERIFICAR SEGMENTOS
    # ========================================================

    if segmentos_encontrados != 12:

        problemas.append(
            f"Foram encontrados "
            f"{segmentos_encontrados} segmentos. "
            f"Devem existir exatamente 12."
        )


    # ========================================================
    # 05. NENHUM PROBLEMA
    # ========================================================

    if not problemas:

        print()
        print("==============================")
        print("NENHUM PROBLEMA PARA QWEN3")
        print("==============================")

        return conteudo


    # ========================================================
    # 06. MOSTRAR PROBLEMAS
    # ========================================================

    print()
    print("==============================")
    print("PROBLEMAS ENCONTRADOS")
    print("==============================")


    for problema in problemas:

        print(
            "-",
            problema
        )


    # ========================================================
    # 07. CONTROLE DO MAPA
    # ========================================================

    mapa_resumido = str(
        mapa_mead
    )


    if len(mapa_resumido) > 5000:

        mapa_resumido = mapa_resumido[
            :5000
        ]


    # ========================================================
    # 08. PROMPT QWEN3
    # ========================================================

    prompt_correcao = f"""
Você é um editor técnico responsável por corrigir
somente problemas estruturais de uma página já escrita.

TEMA:
{tema}

O conteúdo abaixo foi produzido por outro modelo
e deve ser PRESERVADO sempre que estiver correto.

Não reescreva o conteúdo inteiro.

Não melhore estilo sem necessidade.

Não altere informações técnicas corretas.

Não mude o protagonista.

Não invente informações.

Não acrescente dados técnicos que não estejam
presentes no conteúdo ou no MAPA MEAD.

Sua função é corrigir SOMENTE os problemas
identificados.

==================================================
PROBLEMAS IDENTIFICADOS
==================================================

{chr(10).join("- " + p for p in problemas)}

==================================================
REGRAS OBRIGATÓRIAS
==================================================

A página final deve possuir:

5 blocos.

Cada bloco deve possuir exatamente
3 parágrafos.

Total de 15 parágrafos.

Depois dos blocos:

12 segmentos de aplicação.

Os segmentos devem aparecer somente na
seção:

**SEGMENTOS DE APLICAÇÃO**

==================================================
MAPA MEAD
==================================================

{mapa_resumido}

Use o MAPA somente como referência para
preservar o tema e o contexto.

==================================================
CONTEÚDO ORIGINAL
==================================================

{conteudo}

==================================================
REGRA PRINCIPAL
==================================================

Preserve o máximo possível do conteúdo original.

Corrija somente:

- blocos ausentes;
- quantidade incorreta de parágrafos;
- segmentos ausentes ou em quantidade incorreta;
- problemas necessários para recuperar a estrutura.

Se um bloco já possui 3 parágrafos corretos,
NÃO reescreva esse bloco.

Se um segmento já está correto,
NÃO substitua desnecessariamente.

Não altere o sentido técnico.

Não faça uma nova redação completa.

==================================================
FORMATO FINAL
==================================================

Entregue somente o conteúdo corrigido.

Use exatamente:

**BLOCO 1**

[3 parágrafos]

**BLOCO 2**

[3 parágrafos]

**BLOCO 3**

[3 parágrafos]

**BLOCO 4**

[3 parágrafos]

**BLOCO 5**

[3 parágrafos]

**SEGMENTOS DE APLICAÇÃO**

- segmento
- segmento
- segmento
- segmento
- segmento
- segmento
- segmento
- segmento
- segmento
- segmento
- segmento
- segmento

Não explique o que foi corrigido.
Não fale sobre inteligência artificial.
Não mostre o MAPA MEAD.
Entregue somente o conteúdo final.
"""


    # ========================================================
    # 09. DEBUG
    # ========================================================

    print()
    print("==============================")
    print("CONTROLE ENVIO QWEN3")
    print("==============================")


    print(
        "MODELO:",
        "qwen3:latest"
    )


    print(
        "CARACTERES CONTEÚDO:",
        len(conteudo)
    )


    print(
        "CARACTERES MAPA:",
        len(mapa_resumido)
    )


    print(
        "CARACTERES PROBLEMAS:",
        len(
            "\n".join(
                problemas
            )
        )
    )


    print(
        "PROMPT:",
        len(prompt_correcao)
    )


    # ========================================================
    # 10. CONFIGURAÇÃO QWEN3
    # ========================================================

    num_predict = 2600
    num_ctx = 12288


    inicio_qwen3 = time.time()


    # ========================================================
    # 11. OLLAMA
    # ========================================================

    try:

        resposta = requests.post(

            "http://localhost:11434/api/generate",

            json={

                "model":
                    "qwen3:latest",

                "prompt":
                    prompt_correcao,

                "stream":
                    True,

                "think":
                    False,

                "options": {

                    "num_predict":
                        num_predict,

                    "num_ctx":
                        num_ctx,

                    "temperature":
                        0.1,

                    "top_p":
                        0.9,

                    "repeat_penalty":
                        1.08

                }

            },

            stream=True,

            timeout=(
                30,
                900
            )

        )


    except requests.exceptions.Timeout:

        print()
        print("==============================")
        print("TIMEOUT QWEN3")
        print("==============================")

        return ""


    except requests.exceptions.ConnectionError as e:

        print()
        print("==============================")
        print("ERRO DE CONEXÃO QWEN3")
        print("==============================")

        print(
            repr(e)
        )

        return ""


    except Exception as e:

        print()
        print("==============================")
        print("ERRO QWEN3")
        print("==============================")

        print(
            repr(e)
        )

        return ""


    print()
    print("==============================")
    print("QWEN3 RESPONDENDO")
    print("==============================")


    print(
        "STATUS:",
        resposta.status_code
    )


    if resposta.status_code != 200:

        try:

            print(
                resposta.text[:1000]
            )

        except Exception:

            pass

        return ""


    # ========================================================
    # 12. RECEBER STREAM
    # ========================================================

    conteudo_corrigido = ""

    ultimo_print = time.time()


    try:

        for linha in resposta.iter_lines():

            if not linha:

                continue


            try:

                dados = json.loads(
                    linha.decode(
                        "utf-8"
                    )
                )

            except Exception:

                continue


            trecho = dados.get(
                "response",
                ""
            )


            if trecho:

                conteudo_corrigido += trecho


                agora = time.time()


                if agora - ultimo_print >= 30:

                    print()
                    print("==============================")
                    print("STATUS QWEN3")
                    print("==============================")


                    print(
                        "TEMPO DECORRIDO:",
                        round(
                            agora - inicio_qwen3,
                            1
                        ),
                        "s"
                    )


                    print(
                        "CARACTERES:",
                        len(
                            conteudo_corrigido
                        )
                    )


                    ultimo_print = agora


            if dados.get(
                "done",
                False
            ):

                break


    except Exception as e:

        print()
        print("==============================")
        print("ERRO DURANTE STREAM QWEN3")
        print("==============================")


        print(
            repr(e)
        )


        return ""


    fim_qwen3 = time.time()


    conteudo_corrigido = (
        conteudo_corrigido
        .strip()
    )


    # ========================================================
    # 13. RESULTADO
    # ========================================================

    print()
    print("==============================")
    print("QWEN3 FINALIZADO")
    print("==============================")


    print(
        "TEMPO TOTAL:",
        round(
            fim_qwen3 - inicio_qwen3,
            1
        ),
        "segundos"
    )


    print(
        "CARACTERES:",
        len(
            conteudo_corrigido
        )
    )


    # ========================================================
    # 14. VERIFICAR RESPOSTA
    # ========================================================

    if not conteudo_corrigido:

        print()
        print("==============================")
        print("QWEN3 RETORNOU VAZIO")
        print("==============================")

        return ""


    # ========================================================
    # 15. VALIDAÇÃO BÁSICA DA CORREÇÃO
    # ========================================================

    blocos_corrigidos = sum(

        1
        for marcador in marcadores_blocos

        if marcador in conteudo_corrigido

    )


    segmentos_corrigidos = 0


    if "**SEGMENTOS DE APLICAÇÃO**" in conteudo_corrigido:

        trecho_segmentos = (
            conteudo_corrigido.split(
                "**SEGMENTOS DE APLICAÇÃO**",
                1
            )[1]
        )


        for linha in trecho_segmentos.splitlines():

            linha = linha.strip()


            if (
                linha.startswith("- ")
                and len(linha) > 2
            ):

                segmentos_corrigidos += 1


    print()
    print("==============================")
    print("VALIDAÇÃO QWEN3")
    print("==============================")


    print(
        "BLOCOS:",
        blocos_corrigidos,
        "/ 5"
    )


    print(
        "SEGMENTOS:",
        segmentos_corrigidos,
        "/ 12"
    )


    if blocos_corrigidos != 5:

        print()
        print(
            "QWEN3 NÃO PRODUZIU ESTRUTURA VÁLIDA"
        )

        return ""


    if segmentos_corrigidos != 12:

        print()
        print(
            "QWEN3 NÃO PRODUZIU 12 SEGMENTOS"
        )

        return ""


    # ========================================================
    # 16. CORREÇÃO APROVADA
    # ========================================================

    print()
    print("==============================")
    print("CORREÇÃO QWEN3 APROVADA")
    print("==============================")


    return conteudo_corrigido
    

# ============================================================
# EXTRAIR TRECHOS RELEVANTES
# ============================================================

def extrair_trechos_relevantes(
    texto,
    tema,
    mapa_texto,
    limite=5000
):

    if not texto:

        return ""


    # ========================================================
    # 01. NORMALIZAR
    # ========================================================

    texto = str(
        texto
    ).strip()

    if not texto:

        return ""


    # ========================================================
    # 02. PALAVRAS IMPORTANTES
    # ========================================================

    base = (
        f"{tema} "
        f"{mapa_texto}"
    ).lower()


    palavras_mapa = set(

        palavra

        for palavra in re.findall(
            r"\b[a-záàâãéêíóôõúç0-9-]{4,}\b",
            base
        )

    )


    termos_tecnicos = {

        "funcionamento",
        "operação",
        "operacao",
        "pressão",
        "pressao",
        "vazão",
        "vazao",
        "temperatura",
        "eficiência",
        "eficiencia",
        "potência",
        "potencia",
        "rotação",
        "rotacao",
        "motor",
        "rotor",
        "carcaça",
        "carcaca",
        "selo",
        "vedação",
        "vedacao",
        "instalação",
        "instalacao",
        "manutenção",
        "manutencao",
        "segurança",
        "seguranca",
        "desempenho",
        "aplicação",
        "aplicacao",
        "componente",
        "componentes",
        "material",
        "modelo",
        "tipo",
        "capacidade",
        "temperatura",
        "pressão",
        "pressao",
        "altura",
        "fluido",
        "líquido",
        "liquido",
        "tubulação",
        "tubulacao",
        "energia",
        "desgaste",
        "corrosão",
        "corrosao",
        "vibração",
        "vibracao",
        "cavitação",
        "cavitacao"

    }


    termos_genericos = {

        "clique",
        "saiba",
        "contato",
        "comprar",
        "compre",
        "oferta",
        "promoção",
        "promocao",
        "preço",
        "preco",
        "consulte",
        "empresa líder",
        "empresa lider",
        "melhor preço",
        "melhor preco"

    }


    # ========================================================
    # 03. SEPARAR PARÁGRAFOS
    # ========================================================

    paragrafos = [

        p.strip()

        for p in re.split(
            r"\n\s*\n",
            texto
        )

        if p.strip()

    ]


    if not paragrafos:

        paragrafos = [

            texto

        ]


    candidatos = []


    # ========================================================
    # 04. ANALISAR CADA PARÁGRAFO
    # ========================================================

    for indice, paragrafo in enumerate(
        paragrafos
    ):

        texto_lower = paragrafo.lower()


        palavras = set(

            re.findall(
                r"\b[a-záàâãéêíóôõúç0-9-]{4,}\b",
                texto_lower
            )

        )


        pontuacao = 0


    # ========================================================
    # 05. RELAÇÃO COM O TEMA
    # ========================================================

        tema_lower = str(
            tema
        ).lower().strip()


        if tema_lower in texto_lower:

            pontuacao += 15


    # ========================================================
    # 06. TERMOS DO MAPA
    # ========================================================

        correspondencias_mapa = (
            palavras
            &
            palavras_mapa
        )


        pontuacao += min(
            len(correspondencias_mapa) * 2,
            20
        )


    # ========================================================
    # 07. INFORMAÇÃO TÉCNICA
    # ========================================================

        correspondencias_tecnicas = (
            palavras
            &
            termos_tecnicos
        )


        pontuacao += min(
            len(correspondencias_tecnicas) * 3,
            30
        )


    # ========================================================
    # 08. NÚMEROS / UNIDADES
    # ========================================================

        if re.search(
            r"\d+\s*(mm|cm|m|kg|g|bar|psi|°c|c|rpm|kw|cv|l/min|m³/h|hz|v|a)",
            texto_lower
        ):

            pontuacao += 15


    # ========================================================
    # 09. ESTRUTURA TÉCNICA
    # ========================================================

        if ":" in paragrafo:

            pontuacao += 2


        if ";" in paragrafo:

            pontuacao += 2


    # ========================================================
    # 10. PENALIZAR CONTEÚDO GENÉRICO
    # ========================================================

        for termo in termos_genericos:

            if termo in texto_lower:

                pontuacao -= 5


    # ========================================================
    # 11. TAMANHO
    # ========================================================

        tamanho = len(
            paragrafo
        )


        if tamanho < 100:

            pontuacao -= 3


        elif tamanho > 3000:

            pontuacao -= 2


        candidatos.append({

            "indice":
                indice,

            "texto":
                paragrafo,

            "pontuacao":
                pontuacao

        })


    # ========================================================
    # 12. ORDENAR
    # ========================================================

    candidatos.sort(

        key=lambda item:
            item["pontuacao"],

        reverse=True

    )


    # ========================================================
    # 13. MONTAR SELEÇÃO
    # ========================================================

    selecionados = []

    total = 0


    for candidato in candidatos:

        trecho = candidato[
            "texto"
        ]


        tamanho = len(
            trecho
        )


        if not tamanho:

            continue


    # ========================================================
    # 14. LIMITE ABSOLUTO
    # ========================================================

        if (
            total
            +
            tamanho
            >
            limite
        ):

            restante = (
                limite
                -
                total
            )


            if restante >= 250:

                trecho = trecho[
                    :restante
                ]


                # não cortar no meio da palavra

                ultimo_espaco = trecho.rfind(
                    " "
                )


                if ultimo_espaco > 0:

                    trecho = trecho[
                        :ultimo_espaco
                    ]


                selecionados.append(
                    trecho
                )


            break


        selecionados.append(
            trecho
        )


        total += tamanho


    # ========================================================
    # 15. PRESERVAR ORDEM ORIGINAL
    # ========================================================

    textos_selecionados = []


    indices_selecionados = set()


    for trecho in selecionados:

        for candidato in candidatos:

            if (
                candidato["texto"]
                ==
                trecho
            ):

                indices_selecionados.add(
                    candidato["indice"]
                )

                break


    for candidato in sorted(
        candidatos,
        key=lambda item:
            item["indice"]
    ):

        if candidato["indice"] in indices_selecionados:

            textos_selecionados.append(
                candidato["texto"]
            )


    resultado = "\n\n".join(
        textos_selecionados
    )


    return resultado[
        :limite
    ]    


# ============================================================
# SELECIONAR GRUPO
# ============================================================

def preparar_grupo_para_ia(
    grupo,
    tema,
    mapa_texto,
    limite=5000
):

    partes = []

    restante = limite


    for fonte in grupo:

        if restante <= 0:

            break


        trecho = extrair_trechos_relevantes(

            fonte["texto"],

            tema,

            mapa_texto,

            limite=restante

        )


        if not trecho:

            continue


        bloco = f"""
[FONTE {fonte["indice"]}]
TIPO: {fonte["tipo"]}
PDF TÉCNICO: {"SIM" if fonte["eh_pdf"] else "NÃO"}
URL: {fonte["url"]}

{trecho}
"""


        if len(bloco) > restante:

            bloco = bloco[
                :restante
            ]


        partes.append(
            bloco
        )


        restante -= len(
            bloco
        )


    return "\n".join(
        partes
    )[:limite]


# ============================================================
# CRIAR ESTRUTURA DO NOVO JSON DA PÁGINA
# ============================================================

def criar_estrutura_json_pagina(tema):

    tema = str(
        tema or ""
    ).strip()

    return {

        tema: {

            "tema":
                "",

            "nome_site":
                "",

            "grupo_principal_projeto":
                "",

            "segmentos_textuais":
                [],

            "fontes":
                [],

            "referencias":
                [],

            "trechos_utilizados":
                [],

            "grupo":
                "",

            "tipo":
                "",

            "tags":
                [],

            "controle_repeticoes": {

                "palavra_chave":
                    "",

                "meta_repeticoes":
                    60,

                "repeticoes_realizadas":
                    0,

                "repeticoes_faltantes":
                    60
            },

            "mapa_mead": {

                "status":
                    "",

                "texto":
                    ""
            },

            "pagina": {

                "tema":
                    "",

                "arquivo_origem":
                    "",

                "h1":
                    "",

                "titulo":
                    "",

                "subtitulo":
                    "",

                "descricao":
                    "",

                "bloco_1": {

                    "id":
                        "bloco_1",

                    "hash":
                        "",

                    "informacoes_relevantes":
                        "",

                    "titulo":
                        "",

                    "paragrafos":
                        [
                            "",
                            "",
                            ""
                        ]
                },

                "bloco_2": {

                    "id":
                        "bloco_2",

                    "hash":
                        "",

                    "informacoes_relevantes":
                        "",

                    "titulo":
                        "",

                    "paragrafos":
                        [
                            "",
                            "",
                            ""
                        ]
                },

                "bloco_3": {

                    "id":
                        "bloco_3",

                    "hash":
                        "",

                    "informacoes_relevantes":
                        "",

                    "titulo":
                        "",

                    "paragrafos":
                        [
                            "",
                            "",
                            ""
                        ]
                },

                "bloco_4": {

                    "id":
                        "bloco_4",

                    "hash":
                        "",

                    "informacoes_relevantes":
                        "",

                    "titulo":
                        "",

                    "paragrafos":
                        [
                            "",
                            "",
                            ""
                        ]
                },

                "bloco_5": {

                    "id":
                        "bloco_5",

                    "hash":
                        "",

                    "informacoes_relevantes":
                        "",

                    "titulo":
                        "",

                    "paragrafos":
                        [
                            "",
                            "",
                            ""
                        ]
                },

                "segmentos_listas": {

                    "segmento_1": [],
                    "segmento_2": [],
                    "segmento_3": [],
                    "segmento_4": [],
                    "segmento_5": [],
                    "segmento_6": [],
                    "segmento_7": [],
                    "segmento_8": [],
                    "segmento_9": [],
                    "segmento_10": [],
                    "segmento_11": [],
                    "segmento_12": []
                },

                "posicionamento_listas": {

                    "bloco":
                        None
                },

                "imagens": {

                    "imagem_1": {
                        "url": "",
                        "arquivo": "",
                        "alt": "",
                        "descricao": ""
                    },

                    "imagem_2": {
                        "url": "",
                        "arquivo": "",
                        "alt": "",
                        "descricao": ""
                    },

                    "imagem_3": {
                        "url": "",
                        "arquivo": "",
                        "alt": "",
                        "descricao": ""
                    },

                    "imagem_4": {
                        "url": "",
                        "arquivo": "",
                        "alt": "",
                        "descricao": ""
                    },

                    "imagem_5": {
                        "url": "",
                        "arquivo": "",
                        "alt": "",
                        "descricao": ""
                    },

                    "imagem_6": {
                        "url": "",
                        "arquivo": "",
                        "alt": "",
                        "descricao": ""
                    }
                },

                "caracteres":
                    0,

                "status":
                    "em_construcao"
            },

        }
    }


# ============================================================
# HASH DE TRECHO
# ============================================================

def gerar_hash_trecho(
    texto
):
    """
    Gera um identificador único para o trecho selecionado.

    O mesmo trecho, mesmo que tenha espaços diferentes,
    receberá o mesmo hash.
    """

    texto_normalizado = re.sub(
        r"\s+",
        " ",
        str(
            texto or ""
        ).strip().lower()
    )

    if not texto_normalizado:

        return ""

    return hashlib.sha256(
        texto_normalizado.encode(
            "utf-8"
        )
    ).hexdigest()


# ============================================================
# SELECIONAR INFORMAÇÕES RELEVANTES
# ============================================================

def selecionar_informacoes_relevantes(
    tema,
    textos,
    mapa_mead,
    estrutura_editorial
):



    id_execucao = f"{time.time():.6f}"

    print()
    print("##################################################")
    print("ENTRADA EM selecionar_informacoes_relevantes()")
    print("ID EXECUÇÃO:", id_execucao)
    print("TEMA:", tema)
    print("##################################################")

    print()
    print("==============================")
    print("SELECIONANDO INFORMAÇÕES RELEVANTES")
    print("==============================")

    print("TEMA:", tema)
    print("TEXTOS RECEBIDOS:", len(textos))

    # ========================================================
    # FUNÇÕES DE IDENTIFICAÇÃO
    # ========================================================

    def normalizar_texto_hash(texto):

        texto = str(
            texto or ""
        )

        texto = re.sub(
            r"\s+",
            " ",
            texto
        ).strip().lower()

        return texto

    def gerar_hash_trecho(texto):

        texto_normalizado = normalizar_texto_hash(
            texto
        )

        return hashlib.sha256(
            texto_normalizado.encode(
                "utf-8"
            )
        ).hexdigest()

    def gerar_id_trecho(hash_trecho):

        return (
            "IR_"
            + hash_trecho[:16]
        )

    # ========================================================
    # RETORNO VAZIO PADRONIZADO
    # ========================================================

    resultado_vazio = {

        "status":
            "sem_informacoes",

        "caracteres":
            0,

        "texto":
            "",

        "fontes":
            [],

        "fragmentos":
            [],

        "informacoes_relevantes":
            [],

        "blocos": {

            "bloco_1": [],
            "bloco_2": [],
            "bloco_3": [],
            "bloco_4": [],
            "bloco_5": []

        },

        "blocos_informacoes": {

            "bloco_1": {
                "informacoes_relevantes": []
            },

            "bloco_2": {
                "informacoes_relevantes": []
            },

            "bloco_3": {
                "informacoes_relevantes": []
            },

            "bloco_4": {
                "informacoes_relevantes": []
            },

            "bloco_5": {
                "informacoes_relevantes": []
            }

        }

    }

    if not textos:

        print("NENHUM TEXTO RECEBIDO")

        return resultado_vazio

    # ========================================================
    # 01. NORMALIZAR FONTES
    # ========================================================

    fontes = []

    for indice, item in enumerate(
        textos,
        start=1
    ):

        if isinstance(
            item,
            dict
        ):

            texto = (
                item.get("texto")
                or item.get("conteudo")
                or item.get("text")
                or ""
            )

            url = (
                item.get("url")
                or item.get("fonte")
                or ""
            )

            tipo = (
                item.get("tipo")
                or "texto"
            )

        else:

            texto = str(
                item
            )

            url = ""

            tipo = "texto"

        texto = str(
            texto
        ).strip()

        url = str(
            url
        ).strip()

        tipo = str(
            tipo
        ).strip()

        if not texto:
            continue

        eh_pdf = (
            ".pdf" in url.lower()
            or "pdf" in tipo.lower()
        )

        fontes.append({

            "indice":
                indice,

            "url":
                url,

            "tipo":
                tipo,

            "texto":
                texto,

            "eh_pdf":
                eh_pdf

        })

    if not fontes:

        print("NENHUMA FONTE VÁLIDA")

        return resultado_vazio

    print()
    print("==============================")
    print("FONTES NORMALIZADAS")
    print("==============================")

    print(
        "FONTES VÁLIDAS:",
        len(fontes)
    )

    fontes_pdf = [

        fonte
        for fonte in fontes
        if fonte["eh_pdf"]

    ]

    fontes_outros = [

        fonte
        for fonte in fontes
        if not fonte["eh_pdf"]

    ]

    print(
        "PDFS TÉCNICOS:",
        len(fontes_pdf)
    )

    print(
        "OUTRAS FONTES:",
        len(fontes_outros)
    )

    # ========================================================
    # 02. PREPARAR ESTRUTURA EDITORIAL
    # ========================================================

    if not isinstance(
        estrutura_editorial,
        dict
    ):

        estrutura_editorial = {}

    assuntos = [

        str(chave).replace(
            "_",
            " "
        )

        for chave, valor
        in estrutura_editorial.items()

        if valor

    ]

    print()
    print("==============================")
    print("CHECKBOXES RECEBIDOS")
    print("==============================")

    print(
        "ASSUNTOS:",
        assuntos
    )

    # ========================================================
    # 03. PREPARAR MAPA MEAD
    # ========================================================

    mapa_texto = str(
        mapa_mead or ""
    ).strip()

    if len(mapa_texto) > 5000:

        mapa_texto = mapa_texto[:5000]


    # ========================================================
    # 04. PREPARAR CANDIDATOS
    # ========================================================

    candidatos = []


    # --------------------------------------------------------
    #
    # Cada candidato terá entre 50 e 60 palavras.
    #
    # Frases pequenas serão acumuladas até formar um fragmento
    # de pelo menos 50 palavras, sem ultrapassar 60.
    #
    # Frases grandes serão divididas em partes de no máximo
    # 60 palavras.
    #
    # --------------------------------------------------------

    MIN_PALAVRAS_FRAGMENTO = 50
    MAX_PALAVRAS_FRAGMENTO = 60

    for fonte in fontes:

        texto = fonte["texto"]

        texto = re.sub(
            r"\s+",
            " ",
            texto
        ).strip()

        if not texto:
            continue

        # ----------------------------------------------------
        # SEPARAR O TEXTO EM FRASES
        # ----------------------------------------------------

        frases = re.split(
            r"(?<=[.!?])\s+",
            texto
        )

        frases = [

            frase.strip()

            for frase in frases

            if frase.strip()

        ]

        # ----------------------------------------------------
        # ACUMULADOR DO FRAGMENTO
        # ----------------------------------------------------

        acumulado = []

        palavras_acumuladas = 0

        # ----------------------------------------------------
        # PROCESSAR CADA FRASE
        # ----------------------------------------------------

        for frase in frases:

            palavras_frase = re.findall(
                r"\S+",
                frase
            )

            if not palavras_frase:
                continue

            # ------------------------------------------------
            # FRASE PEQUENA:
            # PODE SER ACUMULADA
            # ------------------------------------------------

            if len(palavras_frase) <= MAX_PALAVRAS_FRAGMENTO:

                # --------------------------------------------
                # SE A FRASE COUBER NO FRAGMENTO ATUAL
                # --------------------------------------------

                if (
                    palavras_acumuladas
                    + len(palavras_frase)
                    <= MAX_PALAVRAS_FRAGMENTO
                ):

                    acumulado.append(
                        frase
                    )

                    palavras_acumuladas += len(
                        palavras_frase
                    )

                    # ----------------------------------------
                    # SE JÁ ATINGIU O MÍNIMO,
                    # FECHAR O FRAGMENTO
                    # ----------------------------------------

                    if (
                        palavras_acumuladas
                        >= MIN_PALAVRAS_FRAGMENTO
                    ):

                        fragmento = " ".join(
                            acumulado
                        ).strip()

                        palavras_fragmento = len(
                            re.findall(
                                r"\S+",
                                fragmento
                            )
                        )

                        if (
                            MIN_PALAVRAS_FRAGMENTO
                            <= palavras_fragmento
                            <= MAX_PALAVRAS_FRAGMENTO
                        ):

                            candidatos.append({

                                "texto":
                                    fragmento,

                                "fonte":
                                    fonte["indice"],

                                "url":
                                    fonte["url"],

                                "tipo":
                                    fonte["tipo"],

                                "pdf":
                                    fonte["eh_pdf"],

                                "palavras":
                                    palavras_fragmento

                            })

                        acumulado = []

                        palavras_acumuladas = 0

                else:

                    # ----------------------------------------
                    # A NOVA FRASE NÃO CABE.
                    #
                    # SE O ACUMULADO JÁ TEM 50 OU MAIS,
                    # FECHAR O FRAGMENTO.
                    # ----------------------------------------

                    if (
                        acumulado
                        and palavras_acumuladas
                        >= MIN_PALAVRAS_FRAGMENTO
                    ):

                        fragmento = " ".join(
                            acumulado
                        ).strip()

                        palavras_fragmento = len(
                            re.findall(
                                r"\S+",
                                fragmento
                            )
                        )

                        if (
                            MIN_PALAVRAS_FRAGMENTO
                            <= palavras_fragmento
                            <= MAX_PALAVRAS_FRAGMENTO
                        ):

                            candidatos.append({

                                "texto":
                                    fragmento,

                                "fonte":
                                    fonte["indice"],

                                "url":
                                    fonte["url"],

                                "tipo":
                                    fonte["tipo"],

                                "pdf":
                                    fonte["eh_pdf"],

                                "palavras":
                                    palavras_fragmento

                            })

                    # ----------------------------------------
                    # COMEÇAR NOVO FRAGMENTO
                    # ----------------------------------------

                    acumulado = [
                        frase
                    ]

                    palavras_acumuladas = len(
                        palavras_frase
                    )

            # ------------------------------------------------
            # FRASE GRANDE:
            # DIVIDIR EM PARTES DE 60 PALAVRAS
            # ------------------------------------------------

            else:

                # --------------------------------------------
                # PRIMEIRO FECHAR O ACUMULADO EXISTENTE
                # --------------------------------------------

                if (
                    acumulado
                    and palavras_acumuladas
                    >= MIN_PALAVRAS_FRAGMENTO
                ):

                    fragmento = " ".join(
                        acumulado
                    ).strip()

                    palavras_fragmento = len(
                        re.findall(
                            r"\S+",
                            fragmento
                        )
                    )

                    if (
                        MIN_PALAVRAS_FRAGMENTO
                        <= palavras_fragmento
                        <= MAX_PALAVRAS_FRAGMENTO
                    ):

                        candidatos.append({

                            "texto":
                                fragmento,

                            "fonte":
                                fonte["indice"],

                            "url":
                                fonte["url"],

                            "tipo":
                                fonte["tipo"],

                            "pdf":
                                fonte["eh_pdf"],

                            "palavras":
                                palavras_fragmento

                        })

                    acumulado = []

                    palavras_acumuladas = 0

                # --------------------------------------------
                # QUEBRAR A FRASE GRANDE
                # EM PARTES DE ATÉ 60 PALAVRAS
                # --------------------------------------------

                for inicio in range(
                    0,
                    len(palavras_frase),
                    MAX_PALAVRAS_FRAGMENTO
                ):

                    parte = palavras_frase[
                        inicio:
                        inicio
                        + MAX_PALAVRAS_FRAGMENTO
                    ]

                    if not parte:
                        continue

                    fragmento = " ".join(
                        parte
                    ).strip()

                    palavras_fragmento = len(
                        parte
                    )

                    # ----------------------------------------
                    # PARTES COM 50 A 60 PALAVRAS
                    # ----------------------------------------

                    if (
                        MIN_PALAVRAS_FRAGMENTO
                        <= palavras_fragmento
                        <= MAX_PALAVRAS_FRAGMENTO
                    ):

                        candidatos.append({

                            "texto":
                                fragmento,

                            "fonte":
                                fonte["indice"],

                            "url":
                                fonte["url"],

                            "tipo":
                                fonte["tipo"],

                            "pdf":
                                fonte["eh_pdf"],

                            "palavras":
                                palavras_fragmento

                        })

                    # ----------------------------------------
                    # PARTE FINAL COM MENOS DE 50 PALAVRAS:
                    # GUARDAR PARA TENTAR COMPLETAR COM
                    # A PRÓXIMA FRASE
                    # ----------------------------------------

                    elif (
                        palavras_fragmento
                        < MIN_PALAVRAS_FRAGMENTO
                    ):

                        acumulado = [
                            fragmento
                        ]

                        palavras_acumuladas = (
                            palavras_fragmento
                        )

        # ----------------------------------------------------
        # ÚLTIMO FRAGMENTO DA FONTE
        # ----------------------------------------------------

        if (
            acumulado
            and MIN_PALAVRAS_FRAGMENTO
            <= palavras_acumuladas
            <= MAX_PALAVRAS_FRAGMENTO
        ):

            fragmento = " ".join(
                acumulado
            ).strip()

            palavras_fragmento = len(
                re.findall(
                    r"\S+",
                    fragmento
                )
            )

            if (
                MIN_PALAVRAS_FRAGMENTO
                <= palavras_fragmento
                <= MAX_PALAVRAS_FRAGMENTO
            ):

                candidatos.append({

                    "texto":
                        fragmento,

                    "fonte":
                        fonte["indice"],

                    "url":
                        fonte["url"],

                    "tipo":
                        fonte["tipo"],

                    "pdf":
                        fonte["eh_pdf"],

                    "palavras":
                        palavras_fragmento

                })

    print()
    print("==============================")
    print("CANDIDATOS DE FRAGMENTOS")
    print("==============================")

    print(
        "TOTAL DE CANDIDATOS:",
        len(candidatos)
    )

    # --------------------------------------------------------
    # VERIFICAÇÃO DE SEGURANÇA
    # --------------------------------------------------------
    #
    # Esta verificação é feita antes da remoção de duplicados.
    #
    # Nenhum candidato fora da faixa de 50 a 60 palavras
    # poderá seguir para a etapa seguinte.
    #
    # --------------------------------------------------------

    candidatos_controlados = []

    for candidato in candidatos:

        quantidade_palavras = len(
            re.findall(
                r"\S+",
                candidato.get(
                    "texto",
                    ""
                )
            )
        )

        if (
            quantidade_palavras
            < MIN_PALAVRAS_FRAGMENTO
            or quantidade_palavras
            > MAX_PALAVRAS_FRAGMENTO
        ):

            print(
                "FRAGMENTO DESCARTADO POR TAMANHO:",
                quantidade_palavras,
                "palavras | FONTE:",
                candidato.get(
                    "fonte",
                    ""
                )
            )

            continue

        candidato["palavras"] = (
            quantidade_palavras
        )

        candidatos_controlados.append(
            candidato
        )

    candidatos = candidatos_controlados

    print(
        "CANDIDATOS APÓS CONTROLE:",
        len(candidatos)
    )



    # ========================================================
    # 05. REMOVER DUPLICADOS
    # ========================================================

    candidatos_unicos = []

    fragmentos_vistos = set()

    for candidato in candidatos:

        texto = candidato[
            "texto"
        ].strip()

        chave = normalizar_texto_hash(
            texto
        )

        if not chave:
            continue

        if chave in fragmentos_vistos:
            continue

        fragmentos_vistos.add(
            chave
        )

        # ----------------------------------------------------
        # ID + HASH DO TRECHO
        # ----------------------------------------------------

        hash_trecho = gerar_hash_trecho(
            texto
        )

        id_trecho = gerar_id_trecho(
            hash_trecho
        )

        candidato["id"] = id_trecho

        candidato["hash"] = hash_trecho

        candidatos_unicos.append(
            candidato
        )

    candidatos = candidatos_unicos

    print(
        "CANDIDATOS ÚNICOS:",
        len(candidatos)
    )


    # ========================================================
    # 06. PRIORIZAR PELOS ASSUNTOS DOS CHECKBOXES
    #     + CONTEXTO EDITORIAL DOS 5 BLOCOS DO MEAD
    # ========================================================

    if not isinstance(
        estrutura_editorial,
        dict
    ):

        estrutura_editorial = {}

    # --------------------------------------------------------
    # CHECKBOXES ATIVOS
    # --------------------------------------------------------

    assuntos = [

        str(chave).replace(
            "_",
            " "
        )

        for chave, valor
        in estrutura_editorial.items()

        if valor

    ]

    # --------------------------------------------------------
    # NORMALIZAÇÃO
    # --------------------------------------------------------

    def normalizar_assunto_texto(
        valor
    ):

        valor = str(
            valor or ""
        ).lower()

        substituicoes = str.maketrans(

            "áàãâäéèêëíìîïóòõôöúùûüç",

            "aaaaaeeeeiiiiooooouuuuc"

        )

        return valor.translate(
            substituicoes
        )

    assuntos_normalizados = []

    for assunto in assuntos:

        assunto_normalizado = (
            normalizar_assunto_texto(
                assunto
            ).strip()
        )

        if assunto_normalizado:

            assuntos_normalizados.append(
                assunto_normalizado
            )

    # ========================================================
    # TERMOS DOS CHECKBOXES
    # ========================================================

    termos_assuntos = {

        "apresentacao": [

            "contexto",
            "finalidade",
            "importância",
            "importancia",
            "necessidade",
            "conceito",
            "definição",
            "definicao",
            "introdução",
            "introducao",
            "característica",
            "caracteristicas"

        ],

        "funcionamento": [

            "funcionamento",
            "funciona",
            "operação",
            "operacao",
            "processo",
            "mecanismo",
            "acionamento",
            "desempenho",
            "movimento",
            "pressão",
            "pressao",
            "vazão",
            "vazao"

        ],

        "aplicacoes": [

            "aplicação",
            "aplicacao",
            "aplicações",
            "aplicacoes",
            "utilização",
            "utilizacao",
            "uso",
            "empregado",
            "empregada",
            "atende",
            "atendimento",
            "sistema",
            "processo"

        ],

        "criterios": [

            "critério",
            "criterio",
            "critérios",
            "criterios",
            "seleção",
            "selecao",
            "dimensionamento",
            "especificação",
            "especificacao",
            "escolha",
            "avaliação",
            "avaliacao"

        ],

        "comercial": [

            "empresa",
            "fabricante",
            "fornecedor",
            "produto",
            "solução",
            "solucao",
            "serviço",
            "servico",
            "atendimento",
            "suporte",
            "equipe",
            "experiência",
            "experiencia"

        ],

        "informacao_tecnica": [

            "informação técnica",
            "informacao tecnica",
            "especificação",
            "especificacao",
            "característica técnica",
            "caracteristica tecnica",
            "material",
            "dimensão",
            "dimensao",
            "capacidade",
            "potência",
            "potencia",
            "pressão",
            "pressao",
            "vazão",
            "vazao",
            "temperatura",
            "rendimento",
            "eficiência",
            "eficiencia"

        ],

        "instalacao_execucao": [

            "instalação",
            "instalacao",
            "montagem",
            "execução",
            "execucao",
            "implantação",
            "implantacao",
            "operação",
            "operacao",
            "manutenção",
            "manutencao",
            "segurança",
            "seguranca",
            "ajuste",
            "inspeção",
            "inspecao"

        ],

        "beneficios": [

            "benefício",
            "beneficio",
            "benefícios",
            "beneficios",
            "vantagem",
            "vantagens",
            "redução",
            "reducao",
            "economia",
            "desempenho",
            "confiabilidade",
            "durabilidade",
            "eficiência",
            "eficiencia",
            "produtividade"

        ]

    }

    # ========================================================
    # 06.1 CONTEXTO EDITORIAL DO MEAD
    # ========================================================

    contexto_blocos_mead = {}

    try:

        if isinstance(
            MEAD,
            dict
        ):

            contexto_blocos_mead = (
                MEAD.get(
                    "contexto_blocos",
                    {}
                )
            )

    except Exception:

        contexto_blocos_mead = {}

    if not isinstance(
        contexto_blocos_mead,
        dict
    ):

        contexto_blocos_mead = {}

    # --------------------------------------------------------
    # OBJETIVOS DOS CINCO BLOCOS
    #
    # Estes termos traduzem os objetivos editoriais do MEAD
    # em critérios de seleção.
    #
    # O texto original do MEAD continua sendo a autoridade.
    # Estes termos apenas permitem ao Python localizar
    # fragmentos compatíveis com aquela finalidade.
    # --------------------------------------------------------

    termos_blocos = {

        "bloco_1": [

            "contextualização",
            "contextualizacao",
            "contexto",
            "importância",
            "importancia",
            "cenário",
            "cenario",
            "necessidade",
            "protagonista",
            "finalidade",
            "conceito",
            "definição",
            "definicao"

        ],

        "bloco_2": [

            "funcionamento",
            "funciona",
            "aplicação",
            "aplicacao",
            "aplicações",
            "aplicacoes",
            "característica",
            "caracteristicas",
            "utilização",
            "utilizacao",
            "uso",
            "operação",
            "operacao",
            "processo",
            "mecanismo"

        ],

        "bloco_3": [

            "critério",
            "criterio",
            "critérios",
            "criterios",
            "cuidado",
            "cuidados",
            "instalação",
            "instalacao",
            "operação",
            "operacao",
            "manutenção",
            "manutencao",
            "segurança",
            "seguranca",
            "aspecto técnico",
            "aspecto tecnico",
            "especificação",
            "especificacao",
            "dimensionamento",
            "inspeção",
            "inspecao",
            "ajuste"

        ],

        "bloco_4": [

            "empresa",
            "conhecimento",
            "conhecimento técnico",
            "conhecimento tecnico",
            "atendimento",
            "suporte",
            "produto",
            "produtos",
            "solução",
            "solucao",
            "experiência",
            "experiencia",
            "serviço",
            "servico",
            "fornecedor",
            "fabricante"

        ],

        "bloco_5": [

            "necessidade",
            "necessidades",
            "aplicação",
            "aplicacao",
            "aplicações",
            "aplicacoes",
            "conhecimento técnico",
            "conhecimento tecnico",
            "solução",
            "solucao",
            "problema",
            "problemas",
            "atendimento",
            "desempenho",
            "resultado",
            "eficiência",
            "eficiencia"

        ]

    }

    # --------------------------------------------------------
    # TAMBÉM LÊ O TEXTO REAL DO MEAD
    #
    # Isso mantém a seleção vinculada ao MEAD carregado,
    # sem enviar o MEAD inteiro para o Ollama.
    # --------------------------------------------------------

    for numero_bloco in range(
        1,
        6
    ):

        chave_bloco = (
            f"bloco_{numero_bloco}"
        )

        dados_bloco_mead = (
            contexto_blocos_mead.get(
                chave_bloco,
                {}
            )
        )

        if not isinstance(
            dados_bloco_mead,
            dict
        ):

            continue

        objetivo_mead = normalizar_assunto_texto(
            dados_bloco_mead.get(
                "objetivo",
                ""
            )
        )

        funcao_mead = normalizar_assunto_texto(
            dados_bloco_mead.get(
                "funcao",
                ""
            )
        )

        texto_mead = (
            f"{objetivo_mead} "
            f"{funcao_mead}"
        ).strip()

        if texto_mead:

            termos_existentes = (
                termos_blocos.get(
                    chave_bloco,
                    []
                )
            )

            termos_blocos[
                chave_bloco
            ] = (

                termos_existentes
                +
                [
                    termo
                    for termo in texto_mead.split()
                    if len(termo) > 4
                ]

            )

    # --------------------------------------------------------
    # NORMALIZAR TERMOS DOS BLOCOS
    # --------------------------------------------------------

    termos_blocos_normalizados = {}

    for chave_bloco, termos in termos_blocos.items():

        termos_normalizados = []

        for termo in termos:

            termo_normalizado = (
                normalizar_assunto_texto(
                    termo
                ).strip()
            )

            if termo_normalizado:

                termos_normalizados.append(
                    termo_normalizado
                )

        termos_blocos_normalizados[
            chave_bloco
        ] = termos_normalizados

    # ========================================================
    # 06.2 PONTUAÇÃO DOS CHECKBOXES
    # ========================================================

    def calcular_pontuacao_checkbox(
        texto_normalizado
    ):

        pontuacao = 0

        for assunto in assuntos_normalizados:

            termos = termos_assuntos.get(
                assunto,
                []
            )

            for termo in termos:

                termo_normalizado = (
                    normalizar_assunto_texto(
                        termo
                    )
                )

                if (
                    termo_normalizado
                    and
                    termo_normalizado
                    in texto_normalizado
                ):

                    pontuacao += 1

        return pontuacao

    # ========================================================
    # 06.3 PONTUAÇÃO ESPECÍFICA DE CADA BLOCO
    # ========================================================

    def calcular_pontuacao_bloco(
        candidato,
        chave_bloco
    ):

        texto_normalizado = (
            normalizar_assunto_texto(
                candidato.get(
                    "texto",
                    ""
                )
            )
        )

        if not texto_normalizado:

            return 0

        pontuacao = 0

        # ----------------------------------------------------
        # A. COMPATIBILIDADE COM O OBJETIVO DO BLOCO
        # ----------------------------------------------------

        termos_bloco = termos_blocos_normalizados.get(
            chave_bloco,
            []
        )

        termos_vistos = set()

        for termo in termos_bloco:

            if termo in termos_vistos:

                continue

            termos_vistos.add(
                termo
            )

            if (
                termo
                and
                termo
                in texto_normalizado
            ):

                pontuacao += 2

        # ----------------------------------------------------
        # B. COMPATIBILIDADE COM OS CHECKBOXES
        # ----------------------------------------------------

        pontuacao_checkbox = (
            calcular_pontuacao_checkbox(
                texto_normalizado
            )
        )

        pontuacao += (
            pontuacao_checkbox
        )

        # ----------------------------------------------------
        # C. PRIORIDADE TÉCNICA
        #
        # PDFs continuam tendo vantagem, mas somente como
        # critério complementar.
        # ----------------------------------------------------

        if candidato.get(
            "pdf"
        ):

            pontuacao += 1

        # ----------------------------------------------------
        # D. RELEVÂNCIA DIRETA PARA O TEMA
        # ----------------------------------------------------

        tema_normalizado = (
            normalizar_assunto_texto(
                tema
            )
        )

        palavras_tema = [
            palavra
            for palavra
            in tema_normalizado.split()
            if len(palavra) > 2
        ]

        for palavra in palavras_tema:

            if palavra in texto_normalizado:

                pontuacao += 2

        return pontuacao

    # ========================================================
    # 06.4 ORGANIZAR CANDIDATOS POR BLOCO
    # ========================================================

    candidatos_por_bloco = {

        "bloco_1": [],
        "bloco_2": [],
        "bloco_3": [],
        "bloco_4": [],
        "bloco_5": []

    }

    for candidato in candidatos:

        for numero_bloco in range(
            1,
            6
        ):

            chave_bloco = (
                f"bloco_{numero_bloco}"
            )

            pontuacao = (
                calcular_pontuacao_bloco(
                    candidato,
                    chave_bloco
                )
            )

            candidatos_por_bloco[
                chave_bloco
            ].append({

                "candidato":
                    candidato,

                "pontuacao":
                    pontuacao

            })

    # ========================================================
    # 06.5 ORDENAR CADA BLOCO
    # ========================================================

    for chave_bloco in candidatos_por_bloco:

        candidatos_por_bloco[
            chave_bloco
        ] = sorted(

            candidatos_por_bloco[
                chave_bloco
            ],

            key=lambda item: (
                item["pontuacao"],
                1
                if item["candidato"].get("pdf")
                else 0
            ),

            reverse=True

        )

    print()
    print(
        "=============================="
    )
    print(
        "PRIORIZAÇÃO EDITORIAL POR BLOCO"
    )
    print(
        "=============================="
    )

    print(
        "ASSUNTOS ATIVOS:",
        assuntos
    )

    for numero_bloco in range(
        1,
        6
    ):

        chave_bloco = (
            f"bloco_{numero_bloco}"
        )

        print()
        print(
            chave_bloco.upper()
        )

        dados_mead = (
            contexto_blocos_mead.get(
                chave_bloco,
                {}
            )
        )

        if isinstance(
            dados_mead,
            dict
        ):

            print(
                "OBJETIVO MEAD:",
                dados_mead.get(
                    "objetivo",
                    ""
                )
            )

            print(
                "FUNÇÃO MEAD:",
                dados_mead.get(
                    "funcao",
                    ""
                )
            )

        print(
            "CANDIDATOS:",
            len(
                candidatos_por_bloco[
                    chave_bloco
                ]
            )
        )

    # ========================================================
    # 07. SELECIONAR 15 FRAGMENTOS
    #     3 PARA CADA BLOCO
    # ========================================================

    fragmentos_selecionados = []

    hashes_selecionados = set()

    # --------------------------------------------------------
    # PRIMEIRA PASSAGEM:
    # GARANTIR 3 FRAGMENTOS POR BLOCO
    # --------------------------------------------------------

    for numero_bloco in range(
        1,
        6
    ):

        chave_bloco = (
            f"bloco_{numero_bloco}"
        )

        quantidade_bloco = 0

        for item in candidatos_por_bloco[
            chave_bloco
        ]:

            candidato = item[
                "candidato"
            ]

            hash_trecho = candidato.get(
                "hash",
                ""
            )

            if not hash_trecho:

                hash_trecho = gerar_hash_trecho(
                    candidato["texto"]
                )

                candidato["hash"] = (
                    hash_trecho
                )

            if hash_trecho in hashes_selecionados:

                continue

            hashes_selecionados.add(
                hash_trecho
            )

            # ------------------------------------------------
            # REGISTRAR O BLOCO DE ORIGEM
            # ------------------------------------------------

            candidato[
                "bloco_mead"
            ] = chave_bloco

            fragmentos_selecionados.append(
                candidato
            )

            quantidade_bloco += 1

            if quantidade_bloco >= 3:

                break

    # --------------------------------------------------------
    # SEGUNDA PASSAGEM:
    # SE ALGUM BLOCO NÃO CONSEGUIU 3,
    # COMPLETAR COM OS MELHORES CANDIDATOS RESTANTES.
    #
    # Isso evita perder fragmentos quando o patrimônio não
    # contém material suficiente para determinada finalidade.
    # --------------------------------------------------------

    if len(
        fragmentos_selecionados
    ) < 15:

        candidatos_complementares = []

        for candidato in candidatos:

            hash_trecho = candidato.get(
                "hash",
                ""
            )

            if not hash_trecho:

                hash_trecho = gerar_hash_trecho(
                    candidato["texto"]
                )

                candidato["hash"] = (
                    hash_trecho
                )

            if hash_trecho in hashes_selecionados:

                continue

            melhor_pontuacao = 0
            melhor_bloco = ""

            for numero_bloco in range(
                1,
                6
            ):

                chave_bloco = (
                    f"bloco_{numero_bloco}"
                )

                pontuacao = (
                    calcular_pontuacao_bloco(
                        candidato,
                        chave_bloco
                    )
                )

                if pontuacao > melhor_pontuacao:

                    melhor_pontuacao = (
                        pontuacao
                    )

                    melhor_bloco = (
                        chave_bloco
                    )

            candidatos_complementares.append({

                "candidato":
                    candidato,

                "pontuacao":
                    melhor_pontuacao,

                "bloco":
                    melhor_bloco

            })

        candidatos_complementares = sorted(

            candidatos_complementares,

            key=lambda item: (
                item["pontuacao"],
                1
                if item["candidato"].get("pdf")
                else 0
            ),

            reverse=True

        )

        for item in candidatos_complementares:

            if len(
                fragmentos_selecionados
            ) >= 15:

                break

            candidato = item[
                "candidato"
            ]

            hash_trecho = candidato.get(
                "hash",
                ""
            )

            if (
                not hash_trecho
                or
                hash_trecho
                in hashes_selecionados
            ):

                continue

            hashes_selecionados.add(
                hash_trecho
            )

            candidato[
                "bloco_mead"
            ] = (
                item.get(
                    "bloco",
                    ""
                )
            )

            fragmentos_selecionados.append(
                candidato
            )

    # --------------------------------------------------------
    # CONTROLE FINAL
    # --------------------------------------------------------

    if len(
        fragmentos_selecionados
    ) > 15:

        fragmentos_selecionados = (
            fragmentos_selecionados[:15]
        )

    print()
    print(
        "=============================="
    )
    print(
        "SELEÇÃO MEAD CONCLUÍDA"
    )
    print(
        "=============================="
    )

    print(
        "TOTAL:",
        len(
            fragmentos_selecionados
        ),
        "/ 15"
    )

    for numero_bloco in range(
        1,
        6
    ):

        chave_bloco = (
            f"bloco_{numero_bloco}"
        )

        quantidade = sum(

            1

            for fragmento
            in fragmentos_selecionados

            if fragmento.get(
                "bloco_mead"
            )
            ==
            chave_bloco

        )

        print(
            chave_bloco,
            ":",
            quantidade,
            "fragmentos"
        )


    # ========================================================
    # 08. VERIFICAR QUANTIDADE
    # ========================================================

    print()
    print("==============================")
    print("FRAGMENTOS SELECIONADOS")
    print("==============================")

    print(
        "TOTAL:",
        len(fragmentos_selecionados),
        "/ 15"
    )

    if len(
        fragmentos_selecionados
    ) < 15:

        print(
            "AVISO: NÃO EXISTEM 15 FRAGMENTOS "
            "BRUTOS DISPONÍVEIS."
        )

        print(
            "DISPONÍVEIS:",
            len(
                fragmentos_selecionados
            )
        )


    # ========================================================
    # 09. DISTRIBUIR EM 5 BLOCOS
    # ========================================================
    #
    # Cada bloco passa a ser um objeto completo.
    #
    # Estrutura:
    #
    # bloco_1
    #   ├── id
    #   ├── hash
    #   ├── informacoes_relevantes
    #   ├── titulo
    #   └── paragrafos
    #
    # São 3 fragmentos por bloco.
    #
    # ========================================================

    blocos = {

        "bloco_1": {
            "id": "bloco_1",
            "hash": "",
            "informacoes_relevantes": [],
            "titulo": "",
            "paragrafos": ["", "", ""]
        },

        "bloco_2": {
            "id": "bloco_2",
            "hash": "",
            "informacoes_relevantes": [],
            "titulo": "",
            "paragrafos": ["", "", ""]
        },

        "bloco_3": {
            "id": "bloco_3",
            "hash": "",
            "informacoes_relevantes": [],
            "titulo": "",
            "paragrafos": ["", "", ""]
        },

        "bloco_4": {
            "id": "bloco_4",
            "hash": "",
            "informacoes_relevantes": [],
            "titulo": "",
            "paragrafos": ["", "", ""]
        },

        "bloco_5": {
            "id": "bloco_5",
            "hash": "",
            "informacoes_relevantes": [],
            "titulo": "",
            "paragrafos": ["", "", ""]
        }

    }

    for indice, fragmento in enumerate(
        fragmentos_selecionados
    ):

        numero_bloco = (
            indice // 3
        ) + 1

        chave_bloco = (
            f"bloco_{numero_bloco}"
        )

        # ----------------------------------------------------
        # PRESERVAR O FRAGMENTO COMPLETO
        # ----------------------------------------------------

        dados_fragmento = {

            "id":
                fragmento["id"],

            "hash":
                fragmento["hash"],

            "texto":
                fragmento["texto"],

            "fonte":
                fragmento["fonte"],

            "url":
                fragmento["url"],

            "tipo":
                fragmento["tipo"],

            "pdf":
                fragmento["pdf"],

            "palavras":
                fragmento["palavras"]

        }

        # ----------------------------------------------------
        # COLOCAR O FRAGMENTO NO BLOCO CORRESPONDENTE
        # ----------------------------------------------------

        blocos[
            chave_bloco
        ][
            "informacoes_relevantes"
        ].append(
            dados_fragmento
        )

    # ========================================================
    # GERAR HASH DE CADA BLOCO
    # ========================================================

    for numero_bloco in range(
        1,
        6
    ):

        chave_bloco = (
            f"bloco_{numero_bloco}"
        )

        fragmentos_bloco = (
            blocos[
                chave_bloco
            ][
                "informacoes_relevantes"
            ]
        )

        texto_bloco = " ".join(

            str(
                fragmento.get(
                    "texto",
                    ""
                )
            )

            for fragmento
            in fragmentos_bloco

        ).strip()

        if texto_bloco:

            blocos[
                chave_bloco
            ][
                "hash"
            ] = gerar_hash_trecho(
                texto_bloco
            )

    # ========================================================
    # 09.1 COMPATIBILIDADE COM BLOCOS_INFORMACOES
    # ========================================================

    blocos_informacoes = {}

    for numero_bloco in range(
        1,
        6
    ):

        chave_bloco = (
            f"bloco_{numero_bloco}"
        )

        blocos_informacoes[
            chave_bloco
        ] = {

            "informacoes_relevantes":
                blocos[
                    chave_bloco
                ][
                    "informacoes_relevantes"
                ]

        }



    # ========================================================
    # 10. MOSTRAR DISTRIBUIÇÃO
    # ========================================================

    print()
    print("==============================")
    print("DISTRIBUIÇÃO DOS FRAGMENTOS")
    print("==============================")

    for numero in range(
        1,
        6
    ):

        chave = f"bloco_{numero}"

        quantidade = len(
            blocos[chave]
        )

        print(
            chave,
            ":",
            quantidade,
            "fragmentos"
        )

    # ========================================================
    # 11. MONTAR TEXTO DE COMPATIBILIDADE
    # ========================================================

    partes_texto = []

    for indice, fragmento in enumerate(
        fragmentos_selecionados,
        start=1
    ):

        partes_texto.append(

            f"[FRAGMENTO {indice}]\n"
            f"[ID {fragmento['id']}]\n"
            f"[HASH {fragmento['hash']}]\n"
            f"[FONTE {fragmento['fonte']}]\n"
            f"{fragmento['texto']}"

        )

    material_selecionado = (
        "\n\n".join(
            partes_texto
        ).strip()
    )

    # ========================================================
    # 12. FONTES UTILIZADAS
    # ========================================================

    fontes_resultado = []

    fontes_vistas = set()

    for fragmento in fragmentos_selecionados:

        indice = fragmento[
            "fonte"
        ]

        if indice in fontes_vistas:

            continue

        fontes_vistas.add(
            indice
        )

        fontes_resultado.append({

            "indice":
                fragmento["fonte"],

            "url":
                fragmento["url"],

            "tipo":
                fragmento["tipo"],

            "pdf":
                fragmento["pdf"]

        })

    # ========================================================
    # 13. INFORMAÇÕES_RELEVANTES PRINCIPAIS
    # ========================================================
    #
    # ESTA É A ESTRUTURA QUE SERÁ GRAVADA NO JSON.
    #
    # Cada trecho fica identificado por:
    #
    #   id
    #   hash
    #   texto
    #   fonte
    #   url
    #   tipo
    #   pdf
    #   palavras
    #
    # O HASH É O IDENTIFICADOR PERMANENTE DO CONTEÚDO.
    #
    # ========================================================

    informacoes_relevantes = []

    for fragmento in fragmentos_selecionados:

        informacoes_relevantes.append({

            "id":
                fragmento["id"],

            "hash":
                fragmento["hash"],

            "texto":
                fragmento["texto"],

            "fonte":
                fragmento["fonte"],

            "url":
                fragmento["url"],

            "tipo":
                fragmento["tipo"],

            "pdf":
                fragmento["pdf"],

            "palavras":
                fragmento["palavras"]

        })

    # ========================================================
    # 14. OBJETO FINAL
    # ========================================================

    resultado = {

        "status":
            (
                "selecionado"
                if len(
                    fragmentos_selecionados
                ) == 15
                else "selecionado_parcial"
            ),

        "caracteres":
            len(
                material_selecionado
            ),

        "texto":
            material_selecionado,

        "fontes":
            fontes_resultado,

        "fragmentos":
            fragmentos_selecionados,

        # ----------------------------------------------------
        # NOVO:
        # INFORMAÇÕES RELEVANTES COM ID + HASH
        # ----------------------------------------------------

        "informacoes_relevantes":
            informacoes_relevantes,

        "blocos":
            blocos,

        "blocos_informacoes":
            blocos_informacoes

    }

    # ========================================================
    # 15. GRAVAR DIRETAMENTE NO JSON
    # ========================================================
    
    print()
    print("==============================")
    print("GRAVANDO INFORMAÇÕES RELEVANTES NO JSON")
    print("==============================")
    
    try:
    
        salvar_banco(
            tema,
            "informacoes_relevantes",
            informacoes_relevantes,
            blocos=blocos,
            grupo_principal_projeto=grupo_principal_projeto
        )
    
        print(
            "INFORMAÇÕES_RELEVANTES SALVAS:",
            len(
                informacoes_relevantes
            )
        )
    
    except Exception as erro:
    
        print()
        print("==============================")
        print("ERRO AO SALVAR INFORMAÇÕES RELEVANTES")
        print("==============================")
    
        print(
            erro
        )
        
        
    # ========================================================
    # 16. DEBUG DOS IDS E HASHES
    # ========================================================

    print()
    print("==============================")
    print("ID + HASH DOS TRECHOS SELECIONADOS")
    print("==============================")

    for indice, fragmento in enumerate(
        fragmentos_selecionados,
        start=1
    ):

        print()
        print(
            f"FRAGMENTO {indice}"
        )

        print(
            "ID:",
            fragmento["id"]
        )

        print(
            "HASH:",
            fragmento["hash"]
        )

        print(
            "PALAVRAS:",
            fragmento["palavras"]
        )

        print(
            "FONTE:",
            fragmento["fonte"]
        )

    # ========================================================
    # 17. DEBUG FINAL
    # ========================================================

    print()
    print("==============================")
    print("RESULTADO DA SELEÇÃO")
    print("==============================")

    print(
        "STATUS:",
        resultado["status"]
    )

    print(
        "FRAGMENTOS:",
        len(
            resultado["fragmentos"]
        )
    )

    print(
        "INFORMAÇÕES_RELEVANTES:",
        len(
            resultado[
                "informacoes_relevantes"
            ]
        )
    )

    print(
        "CARACTERES:",
        resultado["caracteres"]
    )

    print(
        "FONTES UTILIZADAS:",
        len(
            resultado["fontes"]
        )
    )

    print()

    for indice, fragmento in enumerate(
        fragmentos_selecionados,
        start=1
    ):

        print(

            f"FRAGMENTO {indice}: "
            f"{fragmento['palavras']} palavras | "
            f"FONTE {fragmento['fonte']} | "
            f"ID {fragmento['id']} | "
            f"HASH {fragmento['hash'][:16]}..."

        )

    print()
    print("==============================")
    print("SELEÇÃO CONCLUÍDA")
    print("==============================")

    return resultado



# ============================================================
# GERAR SEGMENTOS DA PÁGINA
# ============================================================
#
# REGRA MEAD:
#
# - Existe um banco fixo com mais de 20 segmentos.
# - O tema/produto/serviço aparece obrigatoriamente
#   em cada segmento.
# - O Python sorteia exatamente 12 segmentos.
# - Não pode haver repetição dentro da mesma página.
# - Os checkboxes NÃO participam da criação dos segmentos.
# - Segmento genérico sem referência ao tema é proibido.
#
# ============================================================

def gerar_segmentos_pagina(tema):

    import random

    tema = str(
        tema or ""
    ).strip()

    if not tema:
        return []

    # --------------------------------------------------------
    # BANCO FIXO DE SEGMENTOS
    # --------------------------------------------------------
    #
    # O tema é inserido diretamente em todos os segmentos.
    # Portanto nenhum segmento pode ficar sem referência
    # ao produto ou serviço.
    #
    # Mais de 20 opções para permitir variação entre páginas.
    # --------------------------------------------------------

    banco_segmentos = [

        f"Aplicações industriais de {tema}",

        f"Funcionamento e características de {tema}",

        f"Critérios técnicos para seleção de {tema}",

        f"Manutenção preventiva de {tema}",

        f"Instalação adequada de {tema}",

        f"Benefícios operacionais de {tema}",

        f"Dimensionamento de {tema} para diferentes sistemas",

        f"Desempenho técnico de {tema}",

        f"Eficiência operacional de {tema}",

        f"Características construtivas de {tema}",

        f"Cuidados na operação de {tema}",

        f"Soluções industriais com {tema}",

        f"Aplicações de {tema} em processos industriais",

        f"Segurança na operação de {tema}",

        f"Diagnóstico de problemas em {tema}",

        f"Inspeção e conservação de {tema}",

        f"Escolha de {tema} conforme a aplicação",

        f"Condições de operação de {tema}",

        f"Vantagens técnicas de {tema}",

        f"Integração de {tema} em sistemas industriais",

        f"Cuidados durante a instalação de {tema}",

        f"Procedimentos de manutenção de {tema}",

        f"Especificação técnica de {tema}",

        f"Confiabilidade operacional de {tema}",

        f"Desempenho de {tema} em diferentes aplicações",

        f"Aplicação de {tema} em sistemas industriais",

        f"Operação adequada de {tema}",

        f"Seleção de {tema} para processos industriais",

        f"Manutenção e conservação de {tema}",

        f"Aspectos técnicos de {tema}"

    ]

    # --------------------------------------------------------
    # LIMPAR DUPLICIDADES
    # --------------------------------------------------------

    segmentos_validos = []

    vistos = set()

    for segmento in banco_segmentos:

        segmento = str(
            segmento or ""
        ).strip()

        if not segmento:
            continue

        chave = segmento.casefold()

        if chave in vistos:
            continue

        vistos.add(chave)

        segmentos_validos.append(
            segmento
        )

    # --------------------------------------------------------
    # GARANTIR QUANTIDADE MÍNIMA
    # --------------------------------------------------------

    if len(segmentos_validos) < 20:

        print()
        print(
            "❌ ERRO: banco de segmentos possui menos "
            "de 20 opções."
        )

        return []

    # --------------------------------------------------------
    # SORTEAR EXATAMENTE 12
    # --------------------------------------------------------
    #
    # random.sample() garante que não haverá repetição.
    # --------------------------------------------------------

    segmentos_escolhidos = random.sample(
        segmentos_validos,
        12
    )

    # --------------------------------------------------------
    # VALIDAÇÃO FINAL
    # --------------------------------------------------------
    #
    # Todo segmento precisa conter referência ao tema.
    # --------------------------------------------------------

    tema_normalizado = tema.casefold()

    segmentos_finais = []

    for segmento in segmentos_escolhidos:

        if tema_normalizado not in segmento.casefold():

            print()
            print(
                "⚠️ SEGMENTO REJEITADO POR NÃO "
                "REFERENCIAR O TEMA:"
            )
            print(segmento)

            continue

        segmentos_finais.append(
            segmento
        )

    # --------------------------------------------------------
    # SEGURANÇA
    # --------------------------------------------------------

    if len(segmentos_finais) != 12:

        print()
        print(
            "❌ ERRO: não foi possível montar "
            "12 segmentos válidos."
        )

        return []

    # --------------------------------------------------------
    # LOG
    # --------------------------------------------------------

    print()
    print(
        "=========================================="
    )
    print(
        "SEGMENTOS DA PÁGINA"
    )
    print(
        "=========================================="
    )

    print(
        f"TEMA: {tema}"
    )

    print(
        f"BANCO DISPONÍVEL: "
        f"{len(segmentos_validos)}"
    )

    print(
        "SEGMENTOS SORTEADOS: 12"
    )

    for numero, segmento in enumerate(
        segmentos_finais,
        start=1
    ):

        print(
            f"SEGMENTO_{numero}: "
            f"{segmento}"
        )

    print(
        "=========================================="
    )

    return segmentos_finais


# ============================================================
# FORMAS GRAMATICAIS DO TEMA
# ============================================================
#
# Função independente.
#
# NÃO fica dentro de gerar_conteudo_completo().
#
# Objetivo:
# - identificar masculino/feminino do tema;
# - fornecer as formas necessárias para segmentos e tags;
# - evitar o uso do tema sozinho;
# - manter concordância dos adjetivos.
#
# ============================================================

def obter_formas_gramaticais_tema(tema):

    import re

    tema = re.sub(
        r"\s+",
        " ",
        str(tema or "").strip()
    )

    if not tema:
        return {
            "tema": "",
            "genero": "masculino",
            "artigo": "o",
            "artigo_com_tema": "o",
            "de": "do",
            "de_com_tema": "do",
            "em": "no",
            "em_com_tema": "no",
            "especializado": "especializado",
            "tecnico": "técnico",
            "preventivo": "preventivo",
            "corretivo": "corretivo",
            "adequado": "adequado"
        }

    palavras = tema.lower().split()

    primeira_palavra = palavras[0]

    # --------------------------------------------------------
    # PALAVRAS FEMININAS CONHECIDAS
    # --------------------------------------------------------

    palavras_femininas = {
        "bomba",
        "válvula",
        "valvula",
        "máquina",
        "maquina",
        "empresa",
        "instalação",
        "instalacao",
        "manutenção",
        "manutencao",
        "assistência",
        "assistencia",
        "consultoria",
        "inspeção",
        "inspecao",
        "calibração",
        "calibracao",
        "engenharia",
        "solução",
        "solucao",
        "recuperação",
        "recuperacao",
        "adequação",
        "adequacao",
        "aplicação",
        "aplicacao",
        "operação",
        "operacao",
        "seleção",
        "selecao",
        "configuração",
        "configuracao",
        "especificação",
        "especificacao",
        "integração",
        "integracao",
        "proteção",
        "protecao",
        "produção",
        "producao",
        "transmissão",
        "transmissao",
        "gestão",
        "gestao",
        "análise",
        "analise"
    }

    # --------------------------------------------------------
    # PALAVRAS MASCULINAS CONHECIDAS
    # --------------------------------------------------------

    palavras_masculinas = {
        "equipamento",
        "sistema",
        "produto",
        "serviço",
        "servico",
        "processo",
        "motor",
        "compressor",
        "gerador",
        "painel",
        "sensor",
        "controlador",
        "acionamento",
        "dimensionamento",
        "fornecimento",
        "reparo",
        "diagnóstico",
        "diagnostico",
        "projeto",
        "orçamento",
        "orcamento",
        "suporte",
        "atendimento",
        "funcionamento",
        "desempenho",
        "benefício",
        "beneficio",
        "critério",
        "criterio",
        "modelo",
        "fabricante"
    }

    # --------------------------------------------------------
    # IDENTIFICAR GÊNERO
    # --------------------------------------------------------

    if primeira_palavra in palavras_femininas:

        genero = "feminino"

    elif primeira_palavra in palavras_masculinas:

        genero = "masculino"

    elif (
        primeira_palavra.endswith("a")
        or primeira_palavra.endswith("ção")
        or primeira_palavra.endswith("são")
        or primeira_palavra.endswith("ssão")
        or primeira_palavra.endswith("dade")
        or primeira_palavra.endswith("agem")
    ):

        genero = "feminino"

    else:

        genero = "masculino"

    # --------------------------------------------------------
    # FORMAS FEMININAS
    # --------------------------------------------------------

    if genero == "feminino":

        return {
            "tema": tema,
            "genero": "feminino",

            "artigo": "a",
            "artigo_com_tema": f"a {tema}",

            "de": "da",
            "de_com_tema": f"da {tema}",

            "em": "na",
            "em_com_tema": f"na {tema}",

            "especializado": "especializada",
            "tecnico": "técnica",
            "preventivo": "preventiva",
            "corretivo": "corretiva",
            "adequado": "adequada"
        }

    # --------------------------------------------------------
    # FORMAS MASCULINAS
    # --------------------------------------------------------

    return {
        "tema": tema,
        "genero": "masculino",

        "artigo": "o",
        "artigo_com_tema": f"o {tema}",

        "de": "do",
        "de_com_tema": f"do {tema}",

        "em": "no",
        "em_com_tema": f"no {tema}",

        "especializado": "especializado",
        "tecnico": "técnico",
        "preventivo": "preventivo",
        "corretivo": "corretivo",
        "adequado": "adequado"
    }

# ============================================================
# GERAR CONTEÚDO COMPLETO
# ============================================================

def gerar_conteudo_completo(
    tema,
    textos,
    mapa_mead,
    estrutura_editorial,
    arquivo_origem=None,
    dados_coleta=None
):


    
    print()
    print("======================================")
    print("INICIANDO IA - CONTEÚDO COMPLETO")
    print("======================================")
    
    print(
        "TEMA:",
        tema
    )
    
    print(
        "ARQUIVO:",
        arquivo_origem
    )
    
    print(
        "TEXTOS RECEBIDOS:",
        len(textos)
    )
    
    print(
        "MAPA:",
        len(str(mapa_mead))
    )
    
    # ========================================================
    # 01. IDENTIDADE DA PÁGINA
    # ========================================================
    
    if arquivo_origem:
    
        nome_arquivo = os.path.basename(
            arquivo_origem
        )
    
        nome_sem_extensao = os.path.splitext(
            nome_arquivo
        )[0]
    
    else:
    
        nome_arquivo = ""
        nome_sem_extensao = tema
    
    print()
    print("======================================")
    print("IDENTIDADE DA PÁGINA")
    print("======================================")
    
    print(
        "ARQUIVO:",
        nome_arquivo
    )
    
    print(
        "IDENTIDADE:",
        normalizar_tema_chave(
            tema
        )
    )
    
    # ========================================================
    # 02. CHECKBOXES EDITORIAIS
    # ========================================================
    
    print()
    print("======================================")
    print("CHECKBOXES EDITORIAIS")
    print("======================================")
    
    if not isinstance(
        estrutura_editorial,
        dict
    ):
    
        estrutura_editorial = {}
    
    for chave, valor in estrutura_editorial.items():
    
        print(
            f"{chave}: {valor}"
        )
    
    assuntos = [
        str(chave).replace("_", " ")
        for chave, valor
        in estrutura_editorial.items()
        if valor
    ]
    
    print()
    print(
        "ASSUNTOS SELECIONADOS:",
        assuntos
    )
    
    print(
        "TOTAL DE ASSUNTOS:",
        len(assuntos)
    )
    
    # ========================================================
    # 03. ESTRUTURA OBRIGATÓRIA
    # ========================================================
    
    total_blocos = 5
    paragrafos_por_bloco = 3
    total_paragrafos = 15
    total_segmentos = 12
    total_tags = 30
    
    print()
    print("======================================")
    print("ESTRUTURA OBRIGATÓRIA")
    print("======================================")
    
    print(
        "BLOCOS:",
        total_blocos
    )
    
    print(
        "PARÁGRAFOS POR BLOCO:",
        paragrafos_por_bloco
    )
    
    print(
        "TOTAL DE PARÁGRAFOS:",
        total_paragrafos
    )
    
    print(
        "SEGMENTOS:",
        total_segmentos
    )
    
    print(
        "TAGS:",
        total_tags
    )
    
    # ========================================================
    # 04. SELECIONAR INFORMAÇÕES RELEVANTES
    # ========================================================
    
    print()
    print("======================================")
    print("SELECIONANDO INFORMAÇÕES RELEVANTES")
    print("======================================")
    
    textos_relevantes = selecionar_informacoes_relevantes(
        tema,
        textos,
        mapa_mead,
        estrutura_editorial
    )
    
    # ========================================================
    # CORREÇÃO:
    # A função selecionar_informacoes_relevantes()
    # retorna um DICIONÁRIO.
    #
    # Portanto, a validação precisa verificar o conteúdo
    # efetivamente selecionado, e não apenas se o dict existe.
    # ========================================================
    
    if not isinstance(
        textos_relevantes,
        dict
    ):
    
        print()
        print("======================================")
        print("NENHUMA INFORMAÇÃO RELEVANTE")
        print("======================================")
    
        return None
    
    contexto_inicial = str(
        textos_relevantes.get(
            "texto",
            ""
        ) or ""
    ).strip()
    
    fragmentos_iniciais = textos_relevantes.get(
        "fragmentos",
        []
    )
    
    if not contexto_inicial and not fragmentos_iniciais:
    
        print()
        print("======================================")
        print("NENHUMA INFORMAÇÃO RELEVANTE")
        print("======================================")
    
        return None
    
    # ========================================================
    # 05. PREPARAR INFORMAÇÕES
    # ========================================================
    
    print()
    print("======================================")
    print("PREPARANDO INFORMAÇÕES SELECIONADAS")
    print("======================================")
    
    if not isinstance(
        textos_relevantes,
        dict
    ):
    
        print()
        print("======================================")
        print("NENHUMA INFORMAÇÃO RELEVANTE")
        print("======================================")
    
        return None
    
    contexto = str(
        textos_relevantes.get(
            "texto",
            ""
        ) or ""
    ).strip()
    
    blocos_informacoes = (
        textos_relevantes.get(
            "blocos_informacoes",
            {}
        )
    )
    
    if not isinstance(
        blocos_informacoes,
        dict
    ):
    
        blocos_informacoes = {}
    
    informacoes_blocos = {}
    
    for numero_bloco in range(
        1,
        6
    ):
    
        chave_bloco = (
            f"bloco_{numero_bloco}"
        )
    
        informacoes = (
            blocos_informacoes.get(
                chave_bloco,
                ""
            )
        )
    
        if isinstance(
            informacoes,
            list
        ):
    
            partes = []
    
            for item in informacoes:
    
                if isinstance(
                    item,
                    dict
                ):
    
                    texto_item = str(
                        item.get(
                            "texto",
                            ""
                        ) or ""
                    ).strip()
    
                else:
    
                    texto_item = str(
                        item or ""
                    ).strip()
    
                if texto_item:
    
                    partes.append(
                        texto_item
                    )
    
            informacoes = "\n\n".join(
                partes
            )
    
        elif isinstance(
            informacoes,
            dict
        ):
    
            lista_informacoes = (
                informacoes.get(
                    "informacoes_relevantes",
                    []
                )
            )
    
            if isinstance(
                lista_informacoes,
                list
            ):
    
                partes = []
    
                for item in lista_informacoes:
    
                    if isinstance(
                        item,
                        dict
                    ):
    
                        texto_item = str(
                            item.get(
                                "texto",
                                ""
                            ) or ""
                        ).strip()
    
                    else:
    
                        texto_item = str(
                            item or ""
                        ).strip()
    
                    if texto_item:
    
                        partes.append(
                            texto_item
                        )
    
                informacoes = "\n\n".join(
                    partes
                )
    
            else:
    
                informacoes = str(
                    lista_informacoes or ""
                ).strip()
    
        else:
    
            informacoes = str(
                informacoes or ""
            ).strip()
    
        informacoes_blocos[
            chave_bloco
        ] = informacoes
    
    caracteres_selecionados = len(
        contexto
    )
    
    caracteres_blocos = sum(
        len(
            informacoes_blocos.get(
                f"bloco_{numero}",
                ""
            )
        )
        for numero in range(
            1,
            6
        )
    )
    
    print()
    print("======================================")
    print("DADOS APÓS SELEÇÃO")
    print("======================================")
    
    print(
        "SELEÇÃO RECEBIDA:",
        "SIM" if contexto or fragmentos_iniciais else "NÃO"
    )
    
    print(
        "TIPO RECEBIDO:",
        type(
            textos_relevantes
        )
    )
    
    print(
        "CARACTERES TEXTO:",
        caracteres_selecionados
    )
    
    print(
        "CARACTERES BLOCOS:",
        caracteres_blocos
    )
    
    for numero_bloco in range(
        1,
        6
    ):
    
        texto_bloco = informacoes_blocos.get(
            f"bloco_{numero_bloco}",
            ""
        )
    
        print(
            f"BLOCO {numero_bloco}:",
            len(
                texto_bloco
            ),
            "caracteres"
        )
    
    print()
    print("======================================")
    print("INFORMAÇÕES POR BLOCO PREPARADAS")
    print("======================================")
    
    for numero_bloco in range(
        1,
        6
    ):
    
        texto_bloco = informacoes_blocos.get(
            f"bloco_{numero_bloco}",
            ""
        )
    
        print()
        print(
            f"--- BLOCO {numero_bloco} ---"
        )
    
        print(
            texto_bloco[:500]
        )
    
    print()
    print("======================================")
    print("DADOS APÓS SELEÇÃO")
    print("======================================")
    
    print(
        "SELEÇÃO RECEBIDA:",
        "SIM" if contexto or fragmentos_iniciais else "NÃO"
    )
    
    print(
        "CARACTERES SELECIONADOS:",
        caracteres_selecionados
    )
    
    # ========================================================
    # 06. SALVAR INFORMAÇÕES DA PÁGINA
    # ========================================================
    
    print()
    print("======================================")
    print("SALVANDO INFORMAÇÕES DA PÁGINA")
    print("======================================")
    
    salvar_banco(
        tema,
        "informacoes_relevantes",
        textos_relevantes
    )
    
    # ========================================================
    # 07. CONTEXTO PARA IA
    # ========================================================
    
    print()
    print("======================================")
    print("CONTEXTO PARA IA")
    print("======================================")
    
    contexto_geracao = {}
    
    for numero_bloco in range(
        1,
        6
    ):
    
        chave_bloco = (
            f"bloco_{numero_bloco}"
        )
    
        contexto_geracao[
            chave_bloco
        ] = str(
            informacoes_blocos.get(
                chave_bloco,
                ""
            ) or ""
        ).strip()
    
    # ========================================================
    # CORREÇÃO:
    # Guardamos explicitamente o total de caracteres.
    # len(contexto_geracao) retornaria apenas 5, pois
    # contexto_geracao é um dicionário com 5 chaves.
    # ========================================================
    
    caracteres_contexto_geracao = sum(
        len(
            texto_bloco
        )
        for texto_bloco
        in contexto_geracao.values()
    )
    
    print()
    print("======================================")
    print("CONTEXTO SELECIONADO POR BLOCO")
    print("======================================")
    
    for numero_bloco in range(
        1,
        6
    ):
    
        chave_bloco = (
            f"bloco_{numero_bloco}"
        )
    
        texto_bloco = contexto_geracao.get(
            chave_bloco,
            ""
        )
    
        print(
            f"{chave_bloco}:",
            len(texto_bloco),
            "caracteres"
        )
    
    print()
    print(
        "TOTAL CARACTERES CONTEXTO:",
        caracteres_contexto_geracao
    )
    
    # ========================================================
    # 08. NORMALIZAR MAPA
    # ========================================================
    
    if isinstance(
        mapa_mead,
        dict
    ):
    
        mapa_texto = str(
            mapa_mead
        )
    
    else:
    
        mapa_texto = str(
            mapa_mead or ""
        )
    
    mapa_texto = mapa_texto[:3000]
    
    # ========================================================
    # 09. PREPARAR MEAD
    # ========================================================
    
    contexto_mead = ""
    
    try:
    
        contexto_mead = preparar_mead(MEAD)
    
    except Exception as erro:
    
        print()
        print("======================================")
        print("ERRO AO PREPARAR MEAD PARA IA")
        print("======================================")
    
        print(
            erro
        )
    
        contexto_mead = ""
    
    contexto_mead = str(
        contexto_mead or ""
    )
    
    # ========================================================
    # 10. REGRA DE INDEPENDÊNCIA
    # ========================================================
    
    regra_independencia = f"""
    
    REGRA FUNDAMENTAL DESTA PÁGINA:
    
    Esta é uma página independente.
    
    TEMA:
    {tema}
    
    ARQUIVO:
    {nome_arquivo}
    
    Nunca copie ou adapte conteúdo de outra página.
    
    Use somente:
    
    o tema atual;
    as informações técnicas selecionadas;
    o mapa estratégico atual;
    as regras editoriais.
    
    O conhecimento técnico pode ser semelhante entre páginas.
    
    A REDAÇÃO NÃO PODE SER.
    
    Crie uma narrativa nova, própria e independente.
    """
    

    # ============================================================
    # 11. PROMPT FINAL — OLLAMA
    #
    # RESPONSABILIDADE:
    # O Python já selecionou os 15 fragmentos.
    #
    # O Ollama recebe SOMENTE:
    # - tema
    # - instruções editoriais
    # - 15 fragmentos separados pelo Python
    #
    # NÃO enviar:
    # - textos brutos
    # - 35 fontes
    # - mapa_mead
    # - MEAD
    # - contexto adicional
    # ============================================================
    
    print()
    print("=" * 60)
    print("PREPARANDO CONTEXTO FINAL PARA O OLLAMA")
    print("=" * 60)
    
    
    # ------------------------------------------------------------
    # MONTAR SOMENTE OS 5 BLOCOS JÁ PREPARADOS PELO PYTHON
    # ------------------------------------------------------------
    
    fragmentos_para_ia = []
    
    for numero_bloco in range(1, 6):
    
        chave_bloco = (
            f"bloco_{numero_bloco}"
        )
    
        contexto_bloco = str(
            contexto_geracao.get(
                chave_bloco,
                ""
            )
            or ""
        ).strip()
    
        fragmentos_para_ia.append(
            f"BLOCO {numero_bloco}\n"
            f"{contexto_bloco}"
        )
    
    
    contexto_fragmentos = "\n\n".join(
        fragmentos_para_ia
    ).strip()
    
    
    # ------------------------------------------------------------
    # CONTROLE
    # ------------------------------------------------------------
    
    print()
    print("=" * 60)
    print("CONTEXTO QUE SERÁ ENVIADO AO OLLAMA")
    print("=" * 60)
    
    print(
        "FRAGMENTOS ENVIADOS:",
        15
    )
    
    print(
        "BLOCOS ENVIADOS:",
        5
    )
    
    print(
        "CARACTERES DOS FRAGMENTOS:",
        len(contexto_fragmentos)
    )
    
    for numero_bloco in range(1, 6):
    
        chave_bloco = (
            f"bloco_{numero_bloco}"
        )
    
        tamanho = len(
            str(
                contexto_geracao.get(
                    chave_bloco,
                    ""
                )
                or ""
            )
        )
    
        print(
            f"{chave_bloco}:",
            tamanho,
            "caracteres"
        )
    
    
    # ============================================================
    # PROMPT
    # ============================================================
    
    prompt = f"""
Você é um redator técnico especializado.

TEMA:
{tema}

O Python já fez toda a seleção das informações.

Os textos abaixo são os ÚNICOS dados técnicos que você
deve utilizar para escrever a página.

NÃO pesquise.
NÃO use conhecimento externo.
NÃO invente informações.
NÃO acrescente especificações.
NÃO acrescente números.
NÃO acrescente aplicações que não estejam nos fragmentos.

Sua função é transformar os fragmentos selecionados em
uma redação técnica, natural e original.

==================================================
ESTRUTURA OBRIGATÓRIA
==================================================

Crie:

1 título principal;
1 subtítulo;

5 blocos;

cada bloco deve possuir:
1 título;
3 parágrafos;

total:
15 parágrafos.

Cada parágrafo deve possuir aproximadamente
60 a 70 palavras.

Os 3 fragmentos de cada bloco pertencem somente
àquele bloco.

Não misture os fragmentos entre os blocos.

==================================================
PROTAGONISTA
==================================================

O protagonista absoluto é:

{tema}

Os demais elementos são apenas apoio à explicação.

==================================================
NARRATIVA
==================================================

Escreva de forma técnica, natural e humana.

Varie as aberturas dos parágrafos.

Não comece todos os parágrafos com a palavra-chave.

Evite repetição de estrutura.

Não faça propaganda exagerada.

Não invente experiência, clientes, certificações,
números, desempenho ou características técnicas.

==================================================
MARCADORES — USE EXATAMENTE ASSIM
==================================================

IMPORTANTE:

Os marcadores abaixo devem ser copiados
EXATAMENTE como estão escritos.

NÃO coloque acentos nos marcadores.

NÃO altere letras.

NÃO altere maiúsculas ou minúsculas.

NÃO troque os nomes.

Use exatamente:

[TITULO_PRINCIPAL]
[/TITULO_PRINCIPAL]

[SUBTITULO]
[/SUBTITULO]

[BLOCO_1]
[/BLOCO_1]

[BLOCO_2]
[/BLOCO_2]

[BLOCO_3]
[/BLOCO_3]

[BLOCO_4]
[/BLOCO_4]

[BLOCO_5]
[/BLOCO_5]

[TITULO_BLOCO]
[/TITULO_BLOCO]

[PARAGRAFO_1]
[/PARAGRAFO_1]

[PARAGRAFO_2]
[/PARAGRAFO_2]

[PARAGRAFO_3]
[/PARAGRAFO_3]

==================================================
ORDEM OBRIGATÓRIA
==================================================

[TITULO_PRINCIPAL]
Título
[/TITULO_PRINCIPAL]

[SUBTITULO]
Subtítulo
[/SUBTITULO]

[BLOCO_1]

[TITULO_BLOCO]
Título do bloco
[/TITULO_BLOCO]

[PARAGRAFO_1]
Texto
[/PARAGRAFO_1]

[PARAGRAFO_2]
Texto
[/PARAGRAFO_2]

[PARAGRAFO_3]
Texto
[/PARAGRAFO_3]

[/BLOCO_1]

Repita exatamente a mesma estrutura
para BLOCO_2, BLOCO_3, BLOCO_4 e BLOCO_5.

NÃO escreva nada antes dos marcadores.

NÃO escreva nada depois dos marcadores.

NÃO repita o título principal no final.

==================================================
FRAGMENTOS SELECIONADOS PELO PYTHON
==================================================

{contexto_fragmentos}

==================================================
INÍCIO DA GERAÇÃO
==================================================

Retorne somente o conteúdo estruturado pelos
marcadores obrigatórios.
"""


    # ============================================================
    # CONTROLE FINAL DO PROMPT
    # ============================================================
    
    print()
    print("=" * 60)
    print("PROMPT FINAL PARA OLLAMA")
    print("=" * 60)
    
    print(
        "TEMA:",
        tema
    )
    
    print(
        "CARACTERES DOS FRAGMENTOS:",
        len(contexto_fragmentos)
    )
    
    print(
        "CARACTERES DO PROMPT:",
        len(prompt)
    )
    
    print(
        "FONTES BRUTAS NÃO ENVIADAS:",
        len(textos) if isinstance(textos, list) else 0
    )
    
    print(
        "MAPA MEAD NÃO ENVIADO: SIM"
    )
    
    print(
        "MEAD NÃO ENVIADO: SIM"
    )
    
    print("=" * 60)
    
    
    # ============================================================
    # 12. CHAMADA OLLAMA
    # ============================================================
    
    print()
    print("=" * 60)
    print("ENVIANDO SOMENTE OS 15 FRAGMENTOS PARA O OLLAMA")
    print("=" * 60)
    
    inicio_ollama = time.time()
    
    try:
    
        resposta = requests.post(
    
            "http://localhost:11434/api/generate",
    
            json={
    
                "model":
                    "qwen2.5:3b",
    
                "prompt":
                    prompt,
    
                "stream":
                    False,
    
                "think":
                    False,
    
                "options": {
    
                    "num_predict":
                        5000,
    
                    "num_ctx":
                        8192,
    
                    "temperature":
                        0.2,
    
                    "top_p":
                        0.9,
    
                    "repeat_penalty":
                        1.05
    
                }
    
            },
    
            timeout=(
                30,
                900
            )
    
        )
    
    except requests.exceptions.Timeout:
    
        print()
        print("=" * 60)
        print("TIMEOUT OLLAMA")
        print("=" * 60)
    
        print(
            "A IA demorou mais de 900 segundos."
        )
    
        return None
    
    except requests.exceptions.ConnectionError as erro:
    
        print()
        print("=" * 60)
        print("ERRO DE CONEXÃO COM OLLAMA")
        print("=" * 60)
    
        print(
            repr(erro)
        )
    
        return None
    
    except Exception as erro:
    
        print()
        print("=" * 60)
        print("ERRO NA CHAMADA OLLAMA")
        print("=" * 60)
    
        print(
            repr(erro)
        )
    
        return None
    
    
    fim_ollama = time.time()
    
    
    # ============================================================
    # RESPOSTA
    # ============================================================
    
    print()
    print("=" * 60)
    print("OLLAMA RESPONDEU")
    print("=" * 60)
    
    print(
        "STATUS:",
        resposta.status_code
    )
    
    print(
        "TEMPO:",
        round(
            fim_ollama - inicio_ollama,
            2
        ),
        "segundos"
    )
    
    
    if resposta.status_code != 200:
    
        print(
            "ERRO HTTP:",
            resposta.text[:1000]
        )
    
        return None
    
    
    try:
    
        dados_ollama = resposta.json()
    
    except Exception as erro:
    
        print(
            "ERRO AO LER RESPOSTA JSON:",
            repr(erro)
        )
    
        print(
            resposta.text[:1000]
        )
    
        return None
    
    
    conteudo_bruto = str(
        dados_ollama.get(
            "response",
            ""
        )
        or ""
    ).strip()
    
    
    print(
        "CARACTERES RETORNADOS:",
        len(conteudo_bruto)
    )
    
    
    if not conteudo_bruto:
    
        print(
            "OLLAMA RETORNOU VAZIO"
        )
    
        return None
    
    
    print()
    print("=" * 60)
    print("RESPOSTA BRUTA DO OLLAMA")
    print("=" * 60)
    
    print(
        conteudo_bruto
    )
    
    print("=" * 60)

    # ========================================================
    # 14. FUNÇÃO AUXILIAR PARA EXTRAIR MARCADORES
    # ========================================================
    
    def extrair_marcador(
        texto,
        marcador
    ):
    
        padrao = (
            r"\["
            + re.escape(marcador)
            + r"\]"
            r"(.*?)"
            r"\[/"
            + re.escape(marcador)
            + r"\]"
        )
    
        resultado = re.search(
            padrao,
            texto,
            re.IGNORECASE |
            re.DOTALL
        )
    
        if not resultado:
    
            return ""
    
        return resultado.group(
            1
        ).strip()
    
    # ========================================================
    # 15 / 16 
    # ========================================================
    
    # ========================================================
    # EXTRAIR TÍTULO, SUBTÍTULO, BLOCOS E PARÁGRAFOS
    # ========================================================
    
    titulo = extrair_marcador(
        conteudo_bruto,
        "TITULO_PRINCIPAL"
    )
    
    subtitulo = extrair_marcador(
        conteudo_bruto,
        "SUBTITULO"
    )
    
    blocos = []
    
    for numero_bloco in range(1, total_blocos + 1):
    
        marcador_inicio = f"[BLOCO_{numero_bloco}]"
    
        if numero_bloco < total_blocos:
            marcador_proximo = f"[BLOCO_{numero_bloco + 1}]"
    
            padrao_bloco = (
                re.escape(marcador_inicio)
                + r"(.*?)"
                + re.escape(marcador_proximo)
            )
    
        else:
            padrao_bloco = (
                re.escape(marcador_inicio)
                + r"(.*)$"
            )
    
        resultado_bloco = re.search(
            padrao_bloco,
            str(conteudo_bruto),
            re.IGNORECASE | re.DOTALL
        )
    
        if not resultado_bloco:
            continue
    
        texto_bloco = resultado_bloco.group(1).strip()
    
        titulo_bloco = extrair_marcador(
            texto_bloco,
            "TITULO_BLOCO"
        )
    
        # ----------------------------------------------------
        # REMOVER O TÍTULO DO BLOCO
        # ----------------------------------------------------
    
        texto_paragrafos = re.sub(
            r"\[TITULO_BLOCO\].*?\[/TITULO_BLOCO\]",
            "",
            texto_bloco,
            flags=re.IGNORECASE | re.DOTALL
        ).strip()
    
        paragrafos = []
    
        # ----------------------------------------------------
        # EXTRAIR OS 3 PARÁGRAFOS
        #
        # OLLAMA ATUAL:
        #
        # texto do parágrafo
        # [/PARAGRAFO_1]
        #
        # texto do parágrafo
        # [/PARAGRAFO_2]
        #
        # texto do parágrafo
        # [/PARAGRAFO_3]
        #
        # Portanto não exigimos marcador de abertura.
        # ----------------------------------------------------
    
        for numero_paragrafo in range(
            1,
            paragrafos_por_bloco + 1
        ):
    
            marcador_fechamento = (
                f"[/PARAGRAFO_{numero_paragrafo}]"
            )
    
            pos_fechamento = (
                texto_paragrafos.lower().find(
                    marcador_fechamento.lower()
                )
            )
    
            if pos_fechamento == -1:
                continue
    
            if numero_paragrafo == 1:
    
                inicio = 0
    
            else:
    
                marcador_anterior = (
                    f"[/PARAGRAFO_{numero_paragrafo - 1}]"
                )
    
                pos_anterior = (
                    texto_paragrafos.lower().find(
                        marcador_anterior.lower()
                    )
                )
    
                if pos_anterior == -1:
                    continue
    
                inicio = (
                    pos_anterior
                    + len(marcador_anterior)
                )
    
            paragrafo = texto_paragrafos[
                inicio:pos_fechamento
            ].strip()
    
            if paragrafo:
                paragrafos.append(
                    paragrafo
                )
    
        conteudo_bloco = "\n\n".join(
            paragrafos
        )
    
        blocos.append({
            "numero": numero_bloco,
            "titulo": titulo_bloco,
            "paragrafos": paragrafos,
            "conteudo": conteudo_bloco
        })
    


    # ========================================================
    # 17. CRIAR SEGMENTOS COM PYTHON
    # ========================================================
    #
    # Os segmentos NÃO são gerados pelos checkboxes.
    #
    # O Python possui um banco fixo com mais de 20
    # possibilidades e sorteia exatamente 12 para a página.
    #
    # Cada segmento obrigatoriamente referencia o tema,
    # que representa o produto ou serviço.
    #
    # ========================================================
    
    lista_segmentos = gerar_segmentos_pagina(
        tema
    )
    
    if len(lista_segmentos) != 12:
    
        print()
        print(
            "❌ FALHA NA GERAÇÃO DOS SEGMENTOS."
        )
    
        print(
            "A página não será considerada válida "
            "sem exatamente 12 segmentos."
        )
    
        return None



    # --------------------------------------------------------
    # IDENTIFICAR SE O TEMA REPRESENTA UM SERVIÇO
    # --------------------------------------------------------

    termos_servico = (
        "manutenção",
        "manutencao",
        "reparo",
        "assistência",
        "assistencia",
        "instalação",
        "instalacao",
        "montagem",
        "consultoria",
        "inspeção",
        "inspecao",
        "calibração",
        "calibracao",
        "recuperação",
        "recuperacao",
        "revisão",
        "revisao",
        "diagnóstico",
        "diagnostico",
        "serviço",
        "servico",
        "engenharia",
        "projeto",
        "recondicionamento",
        "adequação",
        "adequacao"
    )

    tema_eh_servico = any(
        termo in tema_lower
        for termo in termos_servico
    )

    # --------------------------------------------------------
    # MONTAR OS 12 SEGMENTOS
    # --------------------------------------------------------

    if tema_eh_servico:

        segmentos_base = [
            tema_base,
            f"{tema_base} industrial",
            f"{tema_base} preventiva",
            f"{tema_base} corretiva",
            f"{tema_base} técnica",
            f"{tema_base} especializada",
            f"{tema_base} para equipamentos industriais",
            f"{tema_base} para sistemas industriais",
            f"{tema_base} para processos industriais",
            f"assistência técnica em {tema_base}",
            f"diagnóstico e {tema_base}",
            f"orçamento de {tema_base}"
        ]

    else:

        segmentos_base = [
            tema_base,
            f"{tema_base} industrial",
            f"{tema_base} para sistemas industriais",
            f"{tema_base} para processos industriais",
            f"{tema_base} para aplicações industriais",
            f"{tema_base} para transferência de fluidos",
            f"{tema_base} para sistemas hidráulicos",
            f"{tema_base} para diferentes aplicações",
            f"aplicações de {tema_base}",
            f"manutenção de {tema_base}",
            f"assistência técnica para {tema_base}",
            f"orçamento de {tema_base}"
        ]

    lista_segmentos = []

    for segmento in segmentos_base:

        segmento = re.sub(
            r"\s+",
            " ",
            str(segmento).strip()
        )

        if not segmento:
            continue

        lista_segmentos.append(
            segmento
        )

    # Garantia estrutural: exatamente 12 segmentos
    lista_segmentos = remover_duplicados(
        lista_segmentos
    ) if "remover_duplicados" in locals() else lista_segmentos

    lista_segmentos = lista_segmentos[:12]


    # ========================================================
    # 18. CRIAR TAGS COM PYTHON
    # ========================================================
    #
    # As tags NÃO são mais geradas pelo Ollama.
    #
    # O Python monta exatamente 30 tags a partir do tema,
    # mantendo o padrão PRODUTO/SERVIÇO.
    #
    # REGRA OBRIGATÓRIA:
    # A palavra-chave NUNCA pode aparecer sozinha como tag.
    #
    # Exemplo:
    #
    # ❌ "bomba centrifuga"
    #
    # ✅ "bomba centrifuga industrial"
    # ✅ "manutenção de bomba centrifuga"
    # ✅ "aplicações de bomba centrifuga"
    #
    # ========================================================

    if tema_eh_servico:

        tags_base = [
            f"{tema_base} industrial",
            f"{tema_base} industrial especializada"
            f"{tema_base} técnica",
            f"{tema_base} especializada",
            f"{tema_base} preventiva",
            f"{tema_base} corretiva",
            f"{tema_base} profissional",
            f"{tema_base} em equipamentos",
            f"{tema_base} em sistemas",
            f"{tema_base} em processos",
            f"serviço de {tema_base}",
            f"serviços de {tema_base}",
            f"empresa de {tema_base}",
            f"especialista em {tema_base}",
            f"assistência técnica {tema_base}",
            f"assistência em {tema_base}",
            f"manutenção relacionada a {tema_base}",
            f"reparo relacionado a {tema_base}",
            f"diagnóstico de {tema_base}",
            f"inspeção de {tema_base}",
            f"solução em {tema_base}",
            f"atendimento de {tema_base}",
            f"suporte técnico {tema_base}",
            f"consultoria em {tema_base}",
            f"orçamento de {tema_base}",
            f"cotação de {tema_base}",
            f"preço de {tema_base}",
            f"contratação de {tema_base}",
            f"empresa especializada em {tema_base}",
            f"serviço especializado {tema_base}"
        ]

    else:

        tags_base = [
            f"{tema_base} industrial",
            f"{tema_base} industrial especializada",
            f"{tema_base} profissional",
            f"{tema_base} técnica",
            f"{tema_base} especializada",
            f"{tema_base} para indústria",
            f"{tema_base} industrial aplicação",
            f"{tema_base} para sistemas",
            f"{tema_base} para processos",
            f"{tema_base} hidráulica",
            f"{tema_base} industrial hidráulica",
            f"aplicações de {tema_base}",
            f"uso de {tema_base}",
            f"soluções com {tema_base}",
            f"equipamento {tema_base}",
            f"sistema com {tema_base}",
            f"fornecedor de {tema_base}",
            f"fabricante de {tema_base}",
            f"empresa de {tema_base}",
            f"especialista em {tema_base}",
            f"assistência técnica {tema_base}",
            f"manutenção de {tema_base}",
            f"reparo de {tema_base}",
            f"diagnóstico de {tema_base}",
            f"orçamento de {tema_base}",
            f"cotação de {tema_base}",
            f"preço de {tema_base}",
            f"compra de {tema_base}",
            f"fornecimento de {tema_base}",
            f"solução industrial {tema_base}"
        ]

    # ========================================================
    # LIMPAR E VALIDAR TAGS
    # ========================================================

    lista_tags = []

    tema_normalizado = re.sub(
        r"\s+",
        " ",
        str(tema_base).strip()
    ).casefold()

    for tag in tags_base:

        tag = re.sub(
            r"\s+",
            " ",
            str(tag).strip()
        )

        if not tag:
            continue

        # ----------------------------------------------------
        # REGRA FUNDAMENTAL:
        # A tag não pode ser exatamente a palavra-chave.
        # ----------------------------------------------------

        if tag.casefold() == tema_normalizado:

            print(
                "⚠️ TAG REJEITADA: "
                "palavra-chave isolada ->",
                tag
            )

            continue

        # ----------------------------------------------------
        # EVITAR DUPLICIDADES
        # ----------------------------------------------------

        if tag.casefold() in [
            item.casefold()
            for item in lista_tags
        ]:

            continue

        lista_tags.append(
            tag
        )

    # ========================================================
    # GARANTIA ESTRUTURAL
    # ========================================================

    lista_tags = remover_duplicados(
        lista_tags
    ) if "remover_duplicados" in locals() else lista_tags

    # ========================================================
    # GARANTIA FINAL:
    # NENHUMA TAG PODE SER O TEMA PURO
    # ========================================================

    lista_tags = [
        tag
        for tag in lista_tags
        if str(tag).strip().casefold()
        != tema_normalizado
    ]

    # --------------------------------------------------------
    # O banco acima possui 29 opções válidas.
    #
    # Como a palavra-chave isolada foi retirada, não devemos
    # simplesmente fazer [:30] e aceitar quantidade menor.
    #
    # Se houver menos de 30 tags válidas, interrompemos a
    # geração para evitar um JSON estruturalmente incorreto.
    # --------------------------------------------------------

    if len(lista_tags) < 30:

        print()
        print(
            "❌ ERRO: quantidade insuficiente de "
            "tags válidas."
        )

        print(
            "TAGS VÁLIDAS:",
            len(lista_tags)
        )

        print(
            "MÍNIMO NECESSÁRIO:",
            30
        )

        return None

    lista_tags = lista_tags[:30]

    # ========================================================
    # VALIDAÇÃO ABSOLUTA
    # ========================================================

    for tag in lista_tags:

        if str(tag).strip().casefold() == tema_normalizado:

            print()
            print(
                "❌ ERRO GRAVE: palavra-chave isolada "
                "detectada nas tags:"
            )

            print(tag)

            return None

    # ========================================================
    # LOG
    # ========================================================

    print()
    print("======================================")
    print("SEGMENTOS E TAGS CRIADOS PELO PYTHON")
    print("======================================")

    print(
        "TIPO:",
        "SERVIÇO" if tema_eh_servico else "PRODUTO"
    )

    print(
        "SEGMENTOS:",
        len(lista_segmentos),
        "/",
        total_segmentos
    )

    print(
        "TAGS:",
        len(lista_tags),
        "/",
        total_tags
    )

    print()
    print("TAGS FINAIS:")

    for numero, tag in enumerate(
        lista_tags,
        start=1
    ):

        print(
            f"TAG_{numero}:",
            tag
        )



    # ========================================================
    # SALVAR INFORMAÇÕES RELEVANTES + SEGMENTOS + TAGS
    # ========================================================
    #
    # Neste ponto o Python já criou:
    # - as informações relevantes
    # - exatamente 12 segmentos
    # - exatamente 30 tags
    #
    # Tudo é enviado junto ao salvar_banco() para que seja
    # gravado na estrutura correta do JSON.
    #
    # ========================================================

    salvar_banco(
        tema,
        "informacoes_relevantes",
        contexto,
        blocos=blocos_informacoes,
        tags=lista_tags,
        segmentos=lista_segmentos
    )

    print()
    print("======================================")
    print("DADOS SALVOS NO BANCO")
    print("======================================")

    print(
        "DADOS ENVIADOS AO salvar_banco():",
        "segmentos =", len(lista_segmentos),
        "tags =", len(lista_tags)
    )

    print(
        "SEGMENTOS SALVOS:",
        len(lista_segmentos)
    )

    print(
        "TAGS SALVAS:",
        len(lista_tags)
    )

    print(
        "TAGS SALVAS:",
        len(lista_tags)
    )



    # ========================================================
    # 19. LIMPAR DUPLICIDADES
    # ========================================================
    
    def remover_duplicados(
        lista
    ):
    
        resultado = []
        vistos = set()
    
        for item in lista:
    
            item_limpo = str(
                item
            ).strip()
    
            chave = item_limpo.lower()
    
            if not item_limpo:
                continue
    
            if chave in vistos:
                continue
    
            vistos.add(
                chave
            )
    
            resultado.append(
                item_limpo
            )
    
        return resultado
    
    lista_segmentos = remover_duplicados(
        lista_segmentos
    )
    
    lista_tags = remover_duplicados(
        lista_tags
    )
    
    # ========================================================
    # 20. VALIDAR ESTRUTURA REAL
    # ========================================================
    
    total_blocos_real = len(
        blocos
    )
    
    total_paragrafos_real = sum(
        len(
            bloco.get(
                "paragrafos",
                []
            )
        )
        for bloco in blocos
    )
    
    total_segmentos_real = len(
        lista_segmentos
    )
    
    total_tags_real = len(
        lista_tags
    )
    
    print()
    print("======================================")
    print("VALIDAÇÃO REAL DA ESTRUTURA")
    print("======================================")
    
    print(
        "TÍTULO:",
        "OK" if titulo else "FALTANDO"
    )
    
    print(
        "SUBTÍTULO:",
        "OK" if subtitulo else "FALTANDO"
    )
    
    print(
        "BLOCOS:",
        total_blocos_real,
        "/",
        total_blocos
    )
    
    print(
        "PARÁGRAFOS:",
        total_paragrafos_real,
        "/",
        total_paragrafos
    )
    
    print(
        "SEGMENTOS:",
        total_segmentos_real,
        "/",
        total_segmentos
    )
    
    print(
        "TAGS:",
        total_tags_real,
        "/",
        total_tags
    )
    
    estrutura_valida = True
    
    if not titulo:
        estrutura_valida = False
    
    if not subtitulo:
        estrutura_valida = False
    
    if total_blocos_real != total_blocos:
        estrutura_valida = False
    
    for bloco in blocos:
    
        if len(
            bloco.get(
                "paragrafos",
                []
            )
        ) != paragrafos_por_bloco:
    
            estrutura_valida = False
    
    if total_paragrafos_real != total_paragrafos:
        estrutura_valida = False
    
    if total_segmentos_real != total_segmentos:
        estrutura_valida = False
    
    if total_tags_real != total_tags:
        estrutura_valida = False
    
    if not estrutura_valida:
    
        print()
        print("======================================")
        print("ESTRUTURA INVÁLIDA")
        print("======================================")
    
        print(
            "A página NÃO será salva como estrutura final."
        )
    
        print(
            "O conteúdo bruto será descartado para evitar"
        )
    
        print(
            "gravar uma estrutura falsa no banco."
        )
    
        return None
    
    # ========================================================
    # 21. MONTAR CONTEÚDO FINAL
    # ========================================================

    partes_conteudo = []

    partes_conteudo.append(
        "TÍTULO PRINCIPAL: "
        + titulo
    )

    partes_conteudo.append(
        "SUBTÍTULO: "
        + subtitulo
    )

    partes_conteudo.append(
        "### Conteúdo"
    )

    for bloco in blocos:

        partes_conteudo.append(
            "#### Bloco "
            + str(
                bloco["numero"]
            )
            + ": "
            + bloco["titulo"]
        )

        for paragrafo in bloco[
            "paragrafos"
        ]:

            partes_conteudo.append(
                paragrafo
            )

    partes_conteudo.append(
        "### Segmentos"
    )

    for indice, segmento in enumerate(
        lista_segmentos,
        start=1
    ):

        partes_conteudo.append(
            f"{indice}. {segmento}"
        )

    partes_conteudo.append(
        "### Tags"
    )

    for indice, tag in enumerate(
        lista_tags,
        start=1
    ):

        partes_conteudo.append(
            f"{indice}. {tag}"
        )

    conteudo = "\n\n".join(
        partes_conteudo
    ).strip()

    # ========================================================
    # 22. ESTRUTURA REAL
    # ========================================================

    estrutura_real = {

        "blocos":
            total_blocos_real,

        "paragrafos_por_bloco":
            paragrafos_por_bloco,

        "total_paragrafos":
            total_paragrafos_real,

        "segmentos":
            total_segmentos_real,

        "tags":
            total_tags_real,

        "independente":
            True
    }



    # ========================================================
    # 23. ENVIAR ESTRUTURA OFICIAL PARA SALVAR_BANCO
    # ========================================================
    
    print()
    print("======================================")
    print("SALVANDO ESTRUTURA OFICIAL NO BANCO")
    print("======================================")
    
    salvar_banco(
        tema,
        "pagina",
        conteudo,
    
        informacoes_adicionais={
    
            # NOME DO SITE — INTERFACE
            "nome_site":
                entrada_site.get().strip()
                if "entrada_site" in globals()
                else "",
    
            # GRUPO PRINCIPAL DO PROJETO — INTERFACE
            "grupo_principal_projeto":
                entrada_grupo.get().strip()
                if "entrada_grupo" in globals()
                else "",
    
            # CHECKBOXES EDITORIAIS MARCADOS
            "segmentos_textuais": [
                NOMES_BLOCOS_EDITORIAIS[chave]
                for chave, marcado
                in obter_estrutura_editorial().items()
                if marcado
                and chave in NOMES_BLOCOS_EDITORIAIS
            ],
    
            # REFERÊNCIAS DAS FONTES COLETADAS
            "referencias": list(
                dict.fromkeys(
                    [
                        str(item.get("url", "")).strip()
                        for item in dados_coleta
                        if isinstance(item, dict)
                        and str(item.get("url", "")).strip()
                    ]
                )
            ),
    
            # ARQUIVO DE ORIGEM
            "arquivo_origem":
                nome_arquivo
        },
    
        blocos=blocos,
        segmentos=lista_segmentos,
        tags=lista_tags,
    
        grupo_principal_projeto=
            entrada_grupo.get().strip()
            if "entrada_grupo" in globals()
            else ""
    )


    # ========================================================
    # 24. SALVAR MAPA MEAD
    # ========================================================

    salvar_banco(
        tema,
        "mapa_mead",
        mapa_texto
    )

    # ========================================================
    # 25. RESULTADO
    # ========================================================

    print()
    print("======================================")
    print("PÁGINA GERADA E SALVA")
    print("======================================")

    print(
        "TEMA:",
        tema
    )

    print(
        "ARQUIVO:",
        nome_arquivo
    )

    print(
        "CARACTERES:",
        len(conteudo)
    )

    print(
        "BLOCOS:",
        total_blocos_real
    )

    print(
        "PARÁGRAFOS:",
        total_paragrafos_real
    )

    print(
        "SEGMENTOS:",
        total_segmentos_real
    )

    print(
        "TAGS:",
        total_tags_real
    )

    print(
        "TEMPO:",
        round(
            tempo_total,
            1
        ),
        "segundos"
    )

    return conteudo



# ============================================================
# LIMPAR REFERÊNCIAS COMERCIAIS
# ============================================================

def limpar_referencias_comerciais(texto, tema):

    if not texto:
        return texto


    tema_lower = tema.lower()


    # =====================================
    # REMOVER MARCAS SOMENTE PARA SELAGEM
    # CORTA FOGO
    # =====================================

    if any(
        palavra in tema_lower
        for palavra in [

            "firestop",
            "selagem",
            "corta fogo",
            "passagem corta fogo"

        ]
    ):

        remover = [

            "CP 636",
            "CKC",
            "Firestop",
            "Hilti",
            "3M",
            "Promat",
            "Noneifire",
            "Promaseal",
            "Tecbor",
            "FFC",
            "FCR"

        ]


        for item in remover:

            texto = re.sub(
                item,
                "",
                texto,
                flags=re.IGNORECASE
            )



    # ========================================================
    # 01. NORMALIZA TEXTO
    # ========================================================

    texto = texto.replace(
        "corta-fogo",
        "corta fogo"
    )


    texto = texto.replace(
        "Corta-Fogo",
        "Corta Fogo"
    )


    texto = re.sub(
        r"\s+",
        " ",
        texto
    )


    return texto.strip()
    
    # ========================================================
    # 02. AUDITORIA DE NATURALIDADE
    # ========================================================

    print()
    print("==============================")
    print("INICIANDO AUDITORIA DE NATURALIDADE")
    print("==============================")

    conteudo = auditar_e_corrigir_aberturas(
        conteudo,
        tema
    )

    print()
    print("==============================")
    print("AUDITORIA DE NATURALIDADE FINALIZADA")
    print("==============================")    


    

# ============================================================
# LIMPAR LISTA DE REFERÊNCIAS
# ============================================================

def limpar_lista_referencias(
    textos,
    tema
):

    textos_limpos = []


    for item in textos:


        if not item:
            continue



    # ========================================================
    # 01. NORMALIZAR FORMATO
    # ========================================================

        if isinstance(item, dict):

            texto = item.get(
                "texto",
                ""
            )

            url = item.get(
                "url",
                ""
            )

            tipo = item.get(
                "tipo",
                "texto"
            )


        else:

            texto = str(item)

            url = ""

            tipo = "texto"



        if not texto:

            continue



        texto_upper = texto.upper()



    # ========================================================
    # 02. REMOVER MAPAS MEAD ANTIGOS
    # ========================================================

        if "### MAPA_MEAD" in texto_upper:

            print(
                "MAPA MEAD ANTIGO REMOVIDO"
            )

            continue



        if "PROTAGONISTA:" in texto_upper:

            print(
                "REFERÊNCIA COM PROTAGONISTA REMOVIDA"
            )

            continue



    # ========================================================
    # 03. LIMPEZA COMERCIAL
    # ========================================================

        texto = limpar_referencias_comerciais(
            texto,
            tema
        )
        
        
    # ========================================================
    # 04. LIMPEZA BÁSICA PARA IA
    # ========================================================
        
        texto = texto.strip()
        
        texto = " ".join(
            texto.split()
        )



        if len(texto) < 300:

            print(
                "TEXTO DESCARTADO - PEQUENO:",
                len(texto)
            )

            continue



    # ========================================================
    # 05. SALVAR TEXTO LIMPO
    # ========================================================

        textos_limpos.append(
            {
                "url": url,
                "tipo": tipo,
                "texto": texto
            }
        )



    print()
    print("==============================")
    print("TEXTOS DISPONÍVEIS")
    print("==============================")
    print(len(textos))



    print()
    print("==============================")
    print("TEXTOS APÓS LIMPEZA")
    print("==============================")
    print(len(textos_limpos))



    for i, item in enumerate(textos_limpos):

        print(
            f"{i+1}: {len(item.get('texto',''))} caracteres | "
            f"{item.get('url','')[:80]}"
        )



    return textos_limpos
    

# ============================================================
# EXTRAIR MAPA MEAD DA RESPOSTA
# ============================================================

def extrair_mapa_mead(conteudo):

    if not conteudo:
        return ""

    inicio = conteudo.find("### MAPA_MEAD")

    if inicio == -1:
        return ""

    secoes = [
        "### DEFINICAO",
        "### BENEFICIOS",
        "### VANTAGENS",
        "### MATERIA_PRIMA",
        "### APLICACOES",
        "### FABRICACAO",
        "### MANUTENCAO",
        "### ATIVOS_NARRATIVOS",
        "### DUVIDAS_FREQUENTES"
    ]

    fim = len(conteudo)

    for secao in secoes:

        pos = conteudo.find(secao)

        if pos > inicio:
            fim = pos
            break

    return conteudo[inicio:fim].strip()
    
    
    

# ============================================================
# VALIDAR MAPA MEAD
# ============================================================

def validar_mapa_mead(texto):

    if not texto:
        return False


    texto_upper = texto.upper()


    # remover acentos
    texto_upper = (
        texto_upper
        .replace("Á","A")
        .replace("À","A")
        .replace("Ã","A")
        .replace("Â","A")
        .replace("É","E")
        .replace("Ê","E")
        .replace("Í","I")
        .replace("Ó","O")
        .replace("Ô","O")
        .replace("Õ","O")
        .replace("Ú","U")
        .replace("Ç","C")
    )


    obrigatorios = [

        "PROTAGONISTA",
        "CENARIO",
        "PROBLEMA",
        "SOLUCAO"

    ]


    for item in obrigatorios:

        if item not in texto_upper:

            print()
            print("==============================")
            print("MAPA MEAD FALTANDO:")
            print(item)
            print("==============================")

            return False


    return True
    
    

# ============================================================
# LER BANCO DE DADOS
# ============================================================

def carregar_banco():

    try:

        if not os.path.exists(
            ARQUIVO_BANCO
        ):

            return {}

        with open(
            ARQUIVO_BANCO,
            "r",
            encoding="utf-8"
        ) as f:

            return json.load(f)

    except Exception as e:

        print()
        print("ERRO AO LER BANCO:")
        print(e)

        return {}




# ============================================================
# BUSCAR MAPA MEAD DO PRÓPRIO TEMA
# ============================================================

def obter_mapa_mead_tema(tema):


    banco = carregar_banco()


    tema_busca = remover_acentos(
        tema.lower()
    )

    tema_busca = " ".join(
        tema_busca.split()
    )


    for chave, dados in banco.items():


        chave_normalizada = remover_acentos(
            chave.lower()
        )

        chave_normalizada = " ".join(
            chave_normalizada.split()
        )


        if chave_normalizada == tema_busca:


            print()
            print("==============================")
            print("MAPA ENCONTRADO NO BANCO")
            print("==============================")
            print(chave)


            if not isinstance(
                dados,
                dict
            ):

                return ""


            mapa = dados.get(
                "mapa_mead",
                {}
            )


            if not isinstance(
                mapa,
                dict
            ):

                return ""


            texto = mapa.get(
                "texto",
                ""
            )


            if not texto:

                return ""


            if not validar_protagonista_mead(
                texto,
                tema
            ):

                print()
                print("==============================")
                print("MAPA MEAD INVÁLIDO")
                print("==============================")

                return ""


            return texto



    print()
    print("==============================")
    print("MAPA NÃO ENCONTRADO")
    print("==============================")

    return ""
    

# ============================================================
# VERIFICAR TEMA NO BANCO
# ============================================================

def verificar_tema_banco(tema):

    banco = carregar_banco()


    if tema in banco:

        dados = banco[tema]


        resultado = {

            "existe": True,

            "mapa_mead": "",

            "conteudo": ""

        }


        mapa = dados.get(
            "mapa_mead",
            {}
        )


        if isinstance(
            mapa,
            dict
        ):

            resultado["mapa_mead"] = mapa.get(
                "texto",
                ""
            )


        categorias = dados.get(
            "categorias",
            {}
        )


        if isinstance(
            categorias,
            dict
        ):

            resultado["conteudo"] = categorias.get(
                "conteudo_completo",
                ""
            )


        return resultado



    return {

        "existe": False,

        "mapa_mead": "",

        "conteudo": ""

    }    


# ============================================================
# ALIMENTAR A IA COM O MAPA MEAD
# ============================================================

def obter_textos_banco():

    banco = carregar_banco()

    textos = []


    for tema, dados in banco.items():


        if not isinstance(
            dados,
            dict
        ):
            continue



    # ========================================================
    # 01. CARREGAR MAPA MEAD
    # ========================================================

        mapa = dados.get(
            "mapa_mead",
            {}
        )


        if isinstance(
            mapa,
            dict
        ):

            texto_mapa = mapa.get(
                "texto",
                ""
            )


            if texto_mapa:

                textos.append(
                    texto_mapa
                )



    # ========================================================
    # 02. CARREGAR CATEGORIAS
    # ========================================================

        categorias = dados.get(
            "categorias",
            {}
        )


        if isinstance(
            categorias,
            dict
        ):


            for categoria, conteudo in categorias.items():


                if isinstance(
                    conteudo,
                    str
                ):

                    textos.append(
                        conteudo
                    )



    return textos


def obter_contexto_banco(
    limite=5000
):

    textos = obter_textos_banco()

    contexto = "\n\n".join(
        item.get("texto", "")
        for item in textos
        if isinstance(item, dict)
    )
    
    print()
    print("==============================")
    print("DEBUG CONTEXTO MAPA")
    print("==============================")
    
    print(
        "ITENS RECEBIDOS:",
        len(textos)
    )
    
    print(
        "CARACTERES GERADOS:",
        len(contexto)
    )
    
    print(
        contexto[:500]
    )

    return contexto[:limite]



# ============================================================
# NORMALIZAR IDENTIDADE DO TEMA
# ============================================================

def normalizar_tema_chave(tema):

    if not tema:

        return ""

    texto = str(tema).strip()

    texto = unicodedata.normalize(
        "NFD",
        texto
    )

    texto = "".join(

        caractere

        for caractere in texto

        if unicodedata.category(caractere) != "Mn"

    )

    texto = texto.lower()

    texto = " ".join(
        texto.split()
    )

    return texto



# ============================================================
# VALIDAR ESTRUTURA DA PÁGINA
# ============================================================

def validar_estrutura_pagina(
    conteudo
):

    print()
    print("==============================")
    print("VALIDANDO ESTRUTURA DA PÁGINA")
    print("==============================")

    texto = str(
        conteudo or ""
    ).strip()

    if not texto:

        print(
            "CONTEÚDO VAZIO"
        )

        return False

    # ========================================================
    # 01. CONTAR PARÁGRAFOS
    # ========================================================

    paragrafos = []

    for bloco in texto.split("\n\n"):

        bloco = bloco.strip()

        if not bloco:
            continue

        # Ignorar títulos de bloco
        if bloco.startswith("#"):
            continue

        if bloco.startswith("BLOCO"):
            continue

        if bloco.startswith("Segmentos"):
            continue

        paragrafos.append(
            bloco
        )

    # ========================================================
    # 02. CONTAR SEGMENTOS
    # ========================================================

    segmentos = []

    encontrou_segmentos = False

    for linha in texto.splitlines():

        linha_limpa = linha.strip()

        if (
            "segmentos atendidos" in
            linha_limpa.lower()
            or
            linha_limpa.lower() == "segmentos"
        ):

            encontrou_segmentos = True
            continue

        if encontrou_segmentos:

            if not linha_limpa:
                continue

            if (
                linha_limpa.startswith("-")
                or
                linha_limpa.startswith("•")
                or
                linha_limpa[:2].isdigit()
            ):

                segmentos.append(
                    linha_limpa
                )

    # ========================================================
    # 03. RESULTADO
    # ========================================================

    total_paragrafos = len(
        paragrafos
    )

    total_segmentos = len(
        segmentos
    )

    print(
        "PARÁGRAFOS ENCONTRADOS:",
        total_paragrafos
    )

    print(
        "SEGMENTOS ENCONTRADOS:",
        total_segmentos
    )

    # ========================================================
    # 04. REGRA OBRIGATÓRIA
    # ========================================================

    if total_paragrafos < 15:

        print()
        print(
            "❌ PÁGINA REPROVADA"
        )

        print(
            "Motivo: menos de 15 parágrafos."
        )

        return False

    if total_segmentos < 12:

        print()
        print(
            "❌ PÁGINA REPROVADA"
        )

        print(
            "Motivo: menos de 12 segmentos."
        )

        return False

    print()
    print(
        "✅ ESTRUTURA APROVADA"
    )

    print(
        "15 PARÁGRAFOS + 12 SEGMENTOS"
    )

    return True



# ============================================================
# GERAR TÍTULOS
# ============================================================

def gerar_titulos(
    tema,
    conteudo,
    mapa_mead
):

    print()
    print("==============================")
    print("GERANDO TÍTULOS")
    print("==============================")

    tema = str(
        tema or ""
    ).strip()

    # ========================================================
    # 01. TÍTULO PRINCIPAL
    # ========================================================

    h1 = tema

    # ========================================================
    # 02. TÍTULO SEO
    # ========================================================

    title = (
        tema.capitalize()
        + " | Soluções e Informações Técnicas"
    )

    # ========================================================
    # 03. RETORNO
    # ========================================================

    resultado = {

        "h1":
        h1,

        "title":
        title

    }

    print()
    print("==============================")
    print("TÍTULOS GERADOS")
    print("==============================")

    print(
        "H1:",
        h1
    )

    print(
        "TITLE:",
        title
    )

    return resultado
    


# ============================================================
# SALVAR BANCO
# ============================================================
    
def salvar_banco(
    tema,
    categoria,
    texto,
    informacoes_adicionais=None,
    categorias=None,
    blocos=None,
    segmentos=None,
    tags=None,
    trechos_utilizados=None,
    grupo_principal_projeto=None
):
    """
    Persiste a página no formato oficial do conteudo-site.json.
    
    ```
    REGRAS:
    
    - Uma única estrutura oficial por tema.
    - Uma nova execução do mesmo tema substitui a versão anterior.
    - Os cinco blocos recebidos da seleção são gravados diretamente:
        pagina.bloco_1
        pagina.bloco_2
        pagina.bloco_3
        pagina.bloco_4
        pagina.bloco_5
    
    - Depois de gravados, os blocos NÃO são apagados
    por chamadas posteriores que não enviem blocos.
    
    - NÃO cria informacoes_relevantes no nível global.
    - Cada bloco possui:
        id
        hash
        informacoes_relevantes
        titulo
        paragrafos[3]
    
    - Mantém:
        5 blocos
        15 parágrafos
        12 segmentos
        6 imagens
        até 30 tags
    """
    
    global PAGINAS_EM_PROCESSAMENTO
    
    # ========================================================
    # PRESERVAR A DIFERENÇA ENTRE:
    #
    # None  = dado não enviado nesta chamada
    #
    # valor = dado efetivamente enviado
    #
    # Isso é fundamental para impedir que chamadas posteriores
    # de salvar_banco() apaguem dados já gravados.
    # ========================================================
    
    recebeu_informacoes_adicionais = (
        informacoes_adicionais is not None
    )
    
    recebeu_categorias = (
        categorias is not None
    )
    
    recebeu_blocos = (
        blocos is not None
    )
    
    recebeu_segmentos = (
        segmentos is not None
    )
    
    recebeu_tags = (
        tags is not None
    )
    
    recebeu_trechos_utilizados = (
        trechos_utilizados is not None
    )
    
    # ========================================================
    # NORMALIZAÇÃO
    # ========================================================
    
    if not isinstance(
        informacoes_adicionais,
        dict
    ):
    
        informacoes_adicionais = {}
    
    if not isinstance(
        categorias,
        dict
    ):
    
        categorias = {}
    
    # ========================================================
    # NORMALIZAR BLOCOS
    #
    # A geração do conteúdo envia os 5 blocos como LISTA:
    #
    # [
    #     {"numero": 1, ...},
    #     {"numero": 2, ...},
    #     ...
    # ]
    #
    # O salvamento trabalha internamente com DICT:
    #
    # {
    #     "bloco_1": {...},
    #     "bloco_2": {...},
    #     ...
    # }
    #
    # Portanto convertemos a lista para o formato oficial
    # antes de qualquer processamento.
    # ========================================================

    if isinstance(
        blocos,
        list
    ):

        blocos_normalizados = {}

        for bloco in blocos:

            if not isinstance(
                bloco,
                dict
            ):
                continue

            try:

                numero_bloco = int(
                    bloco.get(
                        "numero",
                        0
                    )
                )

            except Exception:

                numero_bloco = 0

            if numero_bloco < 1 or numero_bloco > 5:
                continue

            blocos_normalizados[
                f"bloco_{numero_bloco}"
            ] = bloco

        blocos = blocos_normalizados

    elif not isinstance(
        blocos,
        dict
    ):

        blocos = {}
    
    if not isinstance(
        segmentos,
        list
    ):
    
        segmentos = []
    
    if not isinstance(
        tags,
        list
    ):
    
        tags = []
    
    if not isinstance(
        trechos_utilizados,
        list
    ):
    
        trechos_utilizados = []
    
    # ========================================================
    # VALIDAR TEMA
    # ========================================================
    
    if not tema:
    
        print(
            "ERRO: tema vazio ao salvar banco."
        )
    
        return False
    
    tema_original = str(
        tema
    ).strip()
    
    if not tema_original:
    
        print(
            "ERRO: tema vazio ao salvar banco."
        )
    
        return False
    
    try:
    
        tema_normalizado = normalizar_tema_chave(
            tema_original
        )
    
    except Exception:
    
        tema_normalizado = (
            tema_original
            .strip()
            .lower()
        )
    
    # ========================================================
    # GARANTIR DIRETÓRIO
    # ========================================================
    
    try:
    
        diretorio = os.path.dirname(
            ARQUIVO_BANCO
        )
    
        if diretorio:
    
            os.makedirs(
                diretorio,
                exist_ok=True
            )
    
    except Exception as erro:
    
        print(
            "AVISO: não foi possível preparar diretório:",
            repr(erro)
        )
    
    # ========================================================
    # CARREGAR BANCO
    # ========================================================
    
    banco = carregar_banco()
    
    if not isinstance(
        banco,
        dict
    ):
    
        banco = {}
    
    # ========================================================
    # LOCALIZAR TEMA EXISTENTE
    # ========================================================
    
    chave_existente = None
    
    for chave in list(
        banco.keys()
    ):
    
        if not isinstance(
            chave,
            str
        ):
    
            continue
    
        try:
    
            chave_normalizada = normalizar_tema_chave(
                chave
            )
    
        except Exception:
    
            chave_normalizada = (
                chave
                .strip()
                .lower()
            )
    
        if chave_normalizada == tema_normalizado:
    
            chave_existente = chave
    
            break
    
    # ========================================================
    # PRIMEIRA GRAVAÇÃO DA EXECUÇÃO
    # ========================================================
    
    primeira_gravacao = (
        tema_normalizado
        not in PAGINAS_EM_PROCESSAMENTO
    )
    
    if primeira_gravacao:
    
        # ----------------------------------------------------
        # Remover versão anterior do mesmo tema.
        # ----------------------------------------------------
    
        for chave in list(
            banco.keys()
        ):
    
            if not isinstance(
                chave,
                str
            ):
    
                continue
    
            try:
    
                chave_normalizada = normalizar_tema_chave(
                    chave
                )
    
            except Exception:
    
                chave_normalizada = (
                    chave
                    .strip()
                    .lower()
                )
    
            if chave_normalizada == tema_normalizado:
    
                del banco[
                    chave
                ]
    
        # ----------------------------------------------------
        # A chave externa será o tema original.
        # ----------------------------------------------------
    
        chave_existente = tema_original
    
        # ----------------------------------------------------
        # Criar estrutura oficial inicial.
        # ----------------------------------------------------
    
        estrutura_inicial = criar_estrutura_json_pagina(
            tema_original
        )
    
        dados_iniciais = estrutura_inicial.get(
            tema_original,
            {}
        )
    
        if not isinstance(
            dados_iniciais,
            dict
        ):
    
            dados_iniciais = {}
    
        banco[
            chave_existente
        ] = dados_iniciais
    
        PAGINAS_EM_PROCESSAMENTO.add(
            tema_normalizado
        )
    
    elif chave_existente is None:
    
        chave_existente = tema_original
    
        estrutura_inicial = criar_estrutura_json_pagina(
            tema_original
        )
    
        dados_iniciais = estrutura_inicial.get(
            tema_original,
            {}
        )
    
        if not isinstance(
            dados_iniciais,
            dict
        ):
    
            dados_iniciais = {}
    
        banco[
            chave_existente
        ] = dados_iniciais
    
    # ========================================================
    # RECUPERAR DADOS DO TEMA
    # ========================================================
    
    dados_tema = banco.get(
        chave_existente
    )
    
    if not isinstance(
        dados_tema,
        dict
    ):
    
        estrutura_inicial = criar_estrutura_json_pagina(
            tema_original
        )
    
        dados_tema = estrutura_inicial.get(
            tema_original,
            {}
        )
    
        if not isinstance(
            dados_tema,
            dict
        ):
    
            dados_tema = {}
    
        banco[
            chave_existente
        ] = dados_tema
    
    # ========================================================
    # RECUPERAR PÁGINA
    # ========================================================
    
    pagina = dados_tema.get(
        "pagina"
    )
    
    if not isinstance(
        pagina,
        dict
    ):
    
        pagina = {}
    
    # ========================================================
    # FUNÇÕES AUXILIARES LOCAIS
    # ========================================================
    
    def normalizar_lista_local(valor):
    
        if isinstance(
            valor,
            list
        ):
    
            resultado = []
    
            for item in valor:
    
                item_limpo = str(
                    item or ""
                ).strip()
    
                if item_limpo:
    
                    resultado.append(
                        item_limpo
                    )
    
            return resultado
    
        if valor is None:
    
            return []
    
        resultado = []
    
        for item in re.split(
            r"[,;\n]+",
            str(valor)
        ):
    
            item_limpo = item.strip()
    
            if item_limpo:
    
                resultado.append(
                    item_limpo
                )
    
        return resultado
    
    def extrair_paragrafos_bloco(bloco):
    
        if not isinstance(
            bloco,
            dict
        ):
    
            return []
    
        paragrafos = bloco.get(
            "paragrafos"
        )
    
        if isinstance(
            paragrafos,
            list
        ):
    
            resultado = []
    
            for item in paragrafos:
    
                item_limpo = str(
                    item or ""
                ).strip()
    
                if item_limpo:
    
                    resultado.append(
                        item_limpo
                    )
    
            return resultado[:3]
    
        conteudo_bloco = str(
            bloco.get(
                "conteudo",
                ""
            )
            or ""
        ).strip()
    
        if not conteudo_bloco:
    
            return []
    
        partes = re.split(
            r"\n\s*\n+",
            conteudo_bloco
        )
    
        resultado = []
    
        for parte in partes:
    
            parte_limpa = parte.strip()
    
            if parte_limpa:
    
                resultado.append(
                    parte_limpa
                )
    
        return resultado[:3]
    
    def criar_bloco_vazio(numero):
    
        return {
            "id": f"bloco_{numero}",
            "hash": "",
            "informacoes_relevantes": "",
            "titulo": "",
            "paragrafos": [
                "",
                "",
                ""
            ]
        }
    
    def atualizar_bloco(numero, dados_bloco):
    
        if numero < 1 or numero > 5:
    
            return
    
        chave_bloco = (
            f"bloco_{numero}"
        )
    
        alvo = pagina.get(
            chave_bloco
        )
    
        if not isinstance(
            alvo,
            dict
        ):
    
            alvo = criar_bloco_vazio(
                numero
            )
    
        # ----------------------------------------------------
        # ID
        # ----------------------------------------------------
    
        id_bloco = str(
            alvo.get(
                "id",
                ""
            )
            or ""
        ).strip()
    
        if not id_bloco:
    
            id_bloco = (
                f"bloco_{numero}"
            )
    
        alvo[
            "id"
        ] = id_bloco
    
        # ----------------------------------------------------
        # HASH
        # ----------------------------------------------------
    
        alvo[
            "hash"
        ] = str(
            alvo.get(
                "hash",
                ""
            )
            or ""
        ).strip()
    
        # ----------------------------------------------------
        # INFORMAÇÕES RELEVANTES
        # ----------------------------------------------------
    
        alvo[
            "informacoes_relevantes"
        ] = str(
            alvo.get(
                "informacoes_relevantes",
                ""
            )
            or ""
        ).strip()
    
        # ----------------------------------------------------
        # TÍTULO
        # ----------------------------------------------------
    
        alvo[
            "titulo"
        ] = str(
            alvo.get(
                "titulo",
                ""
            )
            or ""
        ).strip()
    
        # ----------------------------------------------------
        # PARÁGRAFOS
        # ----------------------------------------------------
    
        paragrafos_existentes = (
            alvo.get(
                "paragrafos",
                []
            )
        )
    
        if not isinstance(
            paragrafos_existentes,
            list
        ):
    
            paragrafos_existentes = []
    
        paragrafos_normalizados = []
    
        for item in paragrafos_existentes[:3]:
    
            paragrafos_normalizados.append(
                str(
                    item or ""
                ).strip()
            )
    
        alvo[
            "paragrafos"
        ] = (
            paragrafos_normalizados
            + [
                "",
                "",
                ""
            ]
        )[:3]
    
        # ====================================================
        # APLICAR DADOS RECEBIDOS
        # ====================================================
    
        if isinstance(
            dados_bloco,
            dict
        ):
    
            # ------------------------------------------------
            # TÍTULO
            # ------------------------------------------------
    
            titulo_recebido = str(
                dados_bloco.get(
                    "titulo",
                    ""
                )
                or ""
            ).strip()
    
            if titulo_recebido:
    
                alvo[
                    "titulo"
                ] = titulo_recebido
    
            # ------------------------------------------------
            # INFORMAÇÕES RELEVANTES
            # ------------------------------------------------
    
            info_recebida = dados_bloco.get(
                "informacoes_relevantes",
                ""
            )
    
            if isinstance(
                info_recebida,
                list
            ):
    
                info_recebida = "\n\n".join(
                    str(
                        item or ""
                    ).strip()
                    for item in info_recebida
                    if str(
                        item or ""
                    ).strip()
                )
    
            info_recebida = str(
                info_recebida or ""
            ).strip()
    
            # ------------------------------------------------
            # IMPORTANTE:
            #
            # Só substitui se realmente recebeu conteúdo.
            #
            # Uma chamada posterior sem conteúdo não apaga
            # a seleção anterior.
            # ------------------------------------------------
    
            if info_recebida:
    
                alvo[
                    "informacoes_relevantes"
                ] = info_recebida
    
            # ------------------------------------------------
            # PARÁGRAFOS
            # ------------------------------------------------
    
            paragrafos_recebidos = (
                extrair_paragrafos_bloco(
                    dados_bloco
                )
            )
    
            if paragrafos_recebidos:
    
                alvo[
                    "paragrafos"
                ] = (
                    paragrafos_recebidos
                    + [
                        "",
                        "",
                        ""
                    ]
                )[:3]
    
        # ====================================================
        # HASH FINAL
        # ====================================================
    
        conteudo_hash = (
            str(
                alvo.get(
                    "informacoes_relevantes",
                    ""
                )
                or ""
            ).strip()
            + "|"
            + str(
                alvo.get(
                    "titulo",
                    ""
                )
                or ""
            ).strip()
            + "|"
            + "|".join(
                str(
                    paragrafo or ""
                ).strip()
                for paragrafo in alvo.get(
                    "paragrafos",
                    []
                )[:3]
            )
        )
    
        alvo[
            "hash"
        ] = hashlib.sha256(
            conteudo_hash.encode(
                "utf-8"
            )
        ).hexdigest()
    
        pagina[
            chave_bloco
        ] = alvo
    
    # ========================================================
    # METADADOS PRINCIPAIS
    # ========================================================
    
    dados_tema[
        "tema"
    ] = tema_original
    
    try:
    
        grupo_identificado = identificar_grupo_tema(
            tema_original
        )
    
    except Exception:
    
        grupo_identificado = ""
    
    dados_tema[
        "grupo"
    ] = (
        grupo_identificado
        or dados_tema.get(
            "grupo",
            ""
        )
    )
    
    try:
    
        tipo_identificado = identificar_tipo_tema(
            tema_original
        )
    
    except Exception:
    
        tipo_identificado = ""
    
    dados_tema[
        "tipo"
    ] = (
        tipo_identificado
        or dados_tema.get(
            "tipo",
            ""
        )
    )
    
    # ========================================================
    # TAGS
    #
    # CORREÇÃO:
    #
    # Só atualiza tags quando tags foram realmente enviadas.
    # Uma chamada posterior sem tags preserva as existentes.
    # ========================================================
    
    if recebeu_tags:
    
        tags_normalizadas = []
    
        vistos_tags = set()
    
        for tag in normalizar_lista_local(
            tags
        ):
    
            try:
    
                chave_tag = normalizar_tema_chave(
                    tag
                )
    
            except Exception:
    
                chave_tag = tag.lower()
    
            if chave_tag in vistos_tags:
    
                continue
    
            vistos_tags.add(
                chave_tag
            )
    
            tags_normalizadas.append(
                tag
            )
    
        dados_tema[
            "tags"
        ] = tags_normalizadas[:30]
    
    else:
    
        if not isinstance(
            dados_tema.get(
                "tags"
            ),
            list
        ):
    
            dados_tema[
                "tags"
            ] = []
    
    # ========================================================
    # SEGMENTOS TEXTUAIS
    #
    # CORREÇÃO:
    #
    # Só atualiza quando segmentos foram enviados.
    # ========================================================
    
    if recebeu_informacoes_adicionais:
    
        segmentos_textuais = (
            informacoes_adicionais.get(
                "segmentos_textuais",
                []
            )
        )
    
        segmentos_textuais = (
            normalizar_lista_local(
                segmentos_textuais
            )
        )
    
        segmentos_textuais_unicos = []
    
        vistos_segmentos = set()
    
        for segmento in segmentos_textuais:
    
            try:
    
                chave_segmento = normalizar_tema_chave(
                    segmento
                )
    
            except Exception:
    
                chave_segmento = segmento.lower()
    
            if chave_segmento in vistos_segmentos:
    
                continue
    
            vistos_segmentos.add(
                chave_segmento
            )
    
            segmentos_textuais_unicos.append(
                segmento
            )
    
        segmentos_textuais = (
            segmentos_textuais_unicos[:12]
        )
    
    elif recebeu_segmentos:
    
        segmentos_textuais = (
            normalizar_lista_local(
                segmentos
            )[:12]
        )
    
    else:
    
        segmentos_textuais = (
            dados_tema
            .get(
                "informacoes_adicionais",
                {}
            )
            .get(
                "segmentos_textuais",
                []
            )
            if isinstance(
                dados_tema.get(
                    "informacoes_adicionais",
                    {}
                ),
                dict
            )
            else []
        )
    
        if isinstance(
            segmentos_textuais,
            dict
        ):
    
            segmentos_textuais = [
                item
                for lista in segmentos_textuais.values()
                if isinstance(lista, list)
                for item in lista
                if item
            ]
    
        segmentos_textuais = (
            normalizar_lista_local(
                segmentos_textuais
            )[:12]
        )
    
    # ========================================================
    # FONTES
    # ========================================================
    
    if recebeu_informacoes_adicionais:
    
        fontes = informacoes_adicionais.get(
            "fontes",
            []
        )
    
        if not isinstance(
            fontes,
            list
        ):
    
            fontes = []
    
        fontes = [
            fonte
            for fonte in fontes
            if fonte
        ]
    
    else:
    
        informacoes_existentes = (
            dados_tema.get(
                "informacoes_adicionais",
                {}
            )
        )
    
        if isinstance(
            informacoes_existentes,
            dict
        ):
    
            fontes = informacoes_existentes.get(
                "fontes",
                []
            )
    
        else:
    
            fontes = []
    
    # ========================================================
    # REFERÊNCIAS
    # ========================================================
    
    if recebeu_informacoes_adicionais:
    
        referencias = informacoes_adicionais.get(
            "referencias",
            []
        )
    
        if not isinstance(
            referencias,
            list
        ):
    
            referencias = []
    
        referencias = [
            referencia
            for referencia in referencias
            if referencia
        ]
    
    else:
    
        informacoes_existentes = (
            dados_tema.get(
                "informacoes_adicionais",
                {}
            )
        )
    
        if isinstance(
            informacoes_existentes,
            dict
        ):
    
            referencias = informacoes_existentes.get(
                "referencias",
                []
            )
    
        else:
    
            referencias = []
    
    # ========================================================
    # NOME DO SITE
    # ========================================================
    
    if recebeu_informacoes_adicionais:
    
        nome_site = str(
            informacoes_adicionais.get(
                "nome_site",
                ""
            )
            or ""
        ).strip()
    
    else:
    
        nome_site = str(
            dados_tema.get(
                "informacoes_adicionais",
                {}
            ).get(
                "nome_site",
                ""
            )
            if isinstance(
                dados_tema.get(
                    "informacoes_adicionais",
                    {}
                ),
                dict
            )
            else ""
        ).strip()
    
    # ========================================================
    # GRUPO PRINCIPAL DO PROJETO
    #
    # O valor vem diretamente do fluxo principal/GUI.
    #
    # NÃO buscar em informacoes_adicionais.
    # NÃO criar informacoes_adicionais.
    # ========================================================

    grupo_principal_projeto = str(
        grupo_principal_projeto or ""
    ).strip()
    
    
    # ========================================================
    # TRECHOS UTILIZADOS
    # ========================================================
    
    dados_tema_existente = dados_tema.get(
        "informacoes_adicionais",
        {}
    )
    
    if not isinstance(
        dados_tema_existente,
        dict
    ):
    
        dados_tema_existente = {}
    
    trechos_existentes = (
        dados_tema_existente.get(
            "trechos_utilizados",
            []
        )
    )
    
    if not isinstance(
        trechos_existentes,
        list
    ):
    
        trechos_existentes = []
    
    hashes_existentes = set()
    
    for existente in trechos_existentes:
    
        if not isinstance(
            existente,
            dict
        ):
    
            continue
    
        hash_existente = str(
            existente.get(
                "hash",
                ""
            )
            or ""
        ).strip()
    
        if hash_existente:
    
            hashes_existentes.add(
                hash_existente
            )
    
    if recebeu_trechos_utilizados:
    
        for item in trechos_utilizados:
    
            if isinstance(
                item,
                dict
            ):
    
                trecho = str(
                    item.get(
                        "trecho",
                        ""
                    )
                    or ""
                ).strip()
    
                hash_trecho = str(
                    item.get(
                        "hash",
                        ""
                    )
                    or ""
                ).strip()
    
                if not hash_trecho and trecho:
    
                    try:
    
                        hash_trecho = gerar_hash_trecho(
                            trecho
                        )
    
                    except Exception:
    
                        hash_trecho = ""
    
                if not hash_trecho:
    
                    continue
    
                if hash_trecho in hashes_existentes:
    
                    continue
    
                trechos_existentes.append({
                    "id": str(
                        item.get(
                            "id",
                            hash_trecho
                        )
                        or hash_trecho
                    ),
                    "hash": hash_trecho,
                    "tema_origem": str(
                        item.get(
                            "tema_origem",
                            tema_original
                        )
                        or tema_original
                    ).strip(),
                    "fonte": str(
                        item.get(
                            "fonte",
                            ""
                        )
                        or ""
                    ).strip()
                })
    
                hashes_existentes.add(
                    hash_trecho
                )
    
            else:
    
                trecho = str(
                    item or ""
                ).strip()
    
                if not trecho:
    
                    continue
    
                try:
    
                    hash_trecho = gerar_hash_trecho(
                        trecho
                    )
    
                except Exception:
    
                    hash_trecho = ""
    
                if not hash_trecho:
    
                    continue
    
                if hash_trecho in hashes_existentes:
    
                    continue
    
                trechos_existentes.append({
                    "id": hash_trecho,
                    "hash": hash_trecho,
                    "tema_origem": tema_original,
                    "fonte": ""
                })
    
                hashes_existentes.add(
                    hash_trecho
                )
    
    
    # ========================================================
    # GARANTIR BLOCOS RECEBIDOS
    #
    # CORREÇÃO CRÍTICA:
    #
    # SOMENTE atualizar blocos se blocos foram enviados.
    #
    # Uma chamada posterior:
    #
    # salvar_banco(tema, "mapa_mead", mapa)
    #
    # não toca nos cinco blocos.
    # ========================================================
    
    if recebeu_blocos:
    
        for numero in range(
            1,
            6
        ):
    
            chave_bloco = (
                f"bloco_{numero}"
            )
    
            if chave_bloco not in blocos:
    
                continue
    
            dados_bloco = blocos.get(
                chave_bloco
            )
    
            if isinstance(
                dados_bloco,
                dict
            ):
    
                dados_bloco_final = dict(
                    dados_bloco
                )
    
            else:
    
                dados_bloco_final = {
                    "informacoes_relevantes":
                        dados_bloco
                }
    
            atualizar_bloco(
                numero,
                dados_bloco_final
            )
    
    # ========================================================
    # CATEGORIA: PÁGINA
    # ========================================================
    
    if categoria == "pagina":
    
        if isinstance(
            texto,
            dict
        ):
    
            pagina_recebida = texto
    
            pagina[
                "tema"
            ] = str(
                pagina_recebida.get(
                    "tema",
                    tema_original
                )
                or tema_original
            ).strip()
    
            pagina[
                "arquivo_origem"
            ] = str(
                pagina_recebida.get(
                    "arquivo_origem",
                    pagina.get(
                        "arquivo_origem",
                        ""
                    )
                )
                or ""
            ).strip()
    
            pagina[
                "h1"
            ] = tema_original
    
            pagina[
                "titulo"
            ] = tema_original
    
            pagina[
                "subtitulo"
            ] = str(
                pagina_recebida.get(
                    "subtitulo",
                    pagina.get(
                        "subtitulo",
                        ""
                    )
                )
                or ""
            ).strip()
    
            pagina[
                "descricao"
            ] = str(
                pagina_recebida.get(
                    "descricao",
                    pagina.get(
                        "descricao",
                        ""
                    )
                )
                or ""
            ).strip()
    
            blocos_recebidos = pagina_recebida.get(
                "blocos",
                []
            )
    
            if isinstance(
                blocos_recebidos,
                list
            ):
    
                for bloco in blocos_recebidos:
    
                    if not isinstance(
                        bloco,
                        dict
                    ):
    
                        continue
    
                    try:
    
                        numero = int(
                            bloco.get(
                                "numero",
                                0
                            )
                        )
    
                    except Exception:
    
                        numero = 0
    
                    atualizar_bloco(
                        numero,
                        bloco
                    )
    
            segmentos_recebidos = (
                pagina_recebida.get(
                    "segmentos",
                    []
                )
            )
    
            if isinstance(
                segmentos_recebidos,
                list
            ) and segmentos_recebidos:
    
                segmentos = segmentos_recebidos
    
                recebeu_segmentos = True
    
    # ========================================================
    # CATEGORIA: CONTEÚDO COMPLETO
    # ========================================================
    
    elif categoria == "conteudo_completo":
    
        conteudo = str(
            texto or ""
        ).strip()
    
        # ----------------------------------------------------
        # TÍTULO
        # ----------------------------------------------------
    
        padroes_titulo = [
    
            r"(?im)^\s*(?:#\s*)?T[ÍI]TULO\s*:\s*(.+)$",
    
            r"(?im)^\s*##\s*T[ÍI]TULO\s*:?\s*(.+)$",
    
            r"(?im)^\s*TITLE\s*:\s*(.+)$"
        ]
    
        for padrao in padroes_titulo:
    
            resultado = re.search(
                padrao,
                conteudo
            )
    
            if resultado:
    
                pagina[
                    "titulo"
                ] = resultado.group(
                    1
                ).strip()
    
                break
    
        # ----------------------------------------------------
        # SUBTÍTULO
        # ----------------------------------------------------
    
        padroes_subtitulo = [
    
            r"(?im)^\s*(?:#\s*)?SUBT[ÍI]TULO\s*:\s*(.+)$",
    
            r"(?im)^\s*##\s*SUBT[ÍI]TULO\s*:?\s*(.+)$",
    
            r"(?im)^\s*SUBTITLE\s*:\s*(.+)$"
        ]
    
        for padrao in padroes_subtitulo:
    
            resultado = re.search(
                padrao,
                conteudo
            )
    
            if resultado:
    
                pagina[
                    "subtitulo"
                ] = resultado.group(
                    1
                ).strip()
    
                break
    
        # ----------------------------------------------------
        # BLOCOS 1 A 5
        # ----------------------------------------------------
    
        padrao_bloco = re.compile(
            r"(?im)"
            r"^\s*"
            r"(?:#+\s*)?"
            r"BLOCO\s*"
            r"([1-5])"
            r"\s*:?"
            r"(.*)$"
        )
    
        ocorrencias = list(
            padrao_bloco.finditer(
                conteudo
            )
        )
    
        for idx, match in enumerate(
            ocorrencias
        ):
    
            inicio = match.end()
    
            if idx + 1 < len(
                ocorrencias
            ):
    
                fim = ocorrencias[
                    idx + 1
                ].start()
    
            else:
    
                fim = len(
                    conteudo
                )
    
            conteudo_bloco = conteudo[
                inicio:fim
            ].strip()
    
            titulo_bloco = (
                match.group(
                    2
                ).strip()
            )
    
            paragrafos = []
    
            partes = re.split(
                r"\n\s*\n+",
                conteudo_bloco
            )
    
            for parte in partes:
    
                parte_limpa = parte.strip()
    
                if not parte_limpa:
    
                    continue
    
                if re.match(
                    r"(?i)^\s*(?:TITULO|SUBTITULO|TAGS?|SEGMENTOS?)\s*:",
                    parte_limpa
                ):
    
                    continue
    
                paragrafos.append(
                    parte_limpa
                )
    
            atualizar_bloco(
                int(
                    match.group(
                        1
                    )
                ),
                {
                    "titulo":
                        titulo_bloco,
    
                    "paragrafos":
                        paragrafos[:3]
                }
            )
    
        # ----------------------------------------------------
        # SEGMENTOS
        # ----------------------------------------------------
    
        padroes_segmentos = [
    
            r"(?is)"
            r"(?:^|\n)"
            r"\s*(?:#+\s*)?"
            r"SEGMENTOS?\s*:?\s*"
            r"\n?"
            r"(.*?)(?="
            r"\n\s*(?:#+\s*)?"
            r"(?:TAGS?|BLOCO|FIM|$)"
            r")",
    
            r"(?is)"
            r"\[\s*SEGMENTOS?\s*\]"
            r"\s*(.*?)(?="
            r"\[\s*TAGS?\s*\]|$)"
        ]
    
        bloco_segmentos = None
    
        for padrao in padroes_segmentos:
    
            resultado = re.search(
                padrao,
                conteudo
            )
    
            if resultado:
    
                bloco_segmentos = resultado.group(
                    1
                ).strip()
    
                break
    
        if bloco_segmentos:
    
            linhas_segmentos = (
                bloco_segmentos.splitlines()
            )
    
            segmentos_extraidos = []
    
            for linha in linhas_segmentos:
    
                linha = linha.strip()
    
                if not linha:
    
                    continue
    
                linha = re.sub(
                    r"^\s*[-•*]\s*",
                    "",
                    linha
                )
    
                linha = re.sub(
                    r"^\s*\d+[\.\)\-:]\s*",
                    "",
                    linha
                )
    
                linha = linha.strip()
    
                if linha:
    
                    segmentos_extraidos.append(
                        linha
                    )
    
            segmentos_unicos = []
    
            vistos_segmentos = set()
    
            for segmento in segmentos_extraidos:
    
                try:
    
                    chave_segmento = normalizar_tema_chave(
                        segmento
                    )
    
                except Exception:
    
                    chave_segmento = segmento.lower()
    
                if chave_segmento in vistos_segmentos:
    
                    continue
    
                vistos_segmentos.add(
                    chave_segmento
                )
    
                segmentos_unicos.append(
                    segmento
                )
    
            segmentos = segmentos_unicos[:12]
    
            # Atualizar a lista textual também.
            segmentos_textuais = segmentos
    
        # ----------------------------------------------------
        # TAGS
        # ----------------------------------------------------
    
        padroes_tags = [
    
            r"(?is)"
            r"(?:^|\n)"
            r"\s*(?:#+\s*)?"
            r"TAGS?\s*:?\s*"
            r"\n?"
            r"(.*?)(?="
            r"\n\s*(?:#+\s*)?"
            r"(?:SEGMENTOS?|BLOCO|FIM|$)"
            r")",
    
            r"(?is)"
            r"\[\s*TAGS?\s*\]"
            r"\s*(.*?)(?="
            r"\[\s*SEGMENTOS?\s*\]|$)"
        ]
    
        bloco_tags = None
    
        for padrao in padroes_tags:
    
            resultado = re.search(
                padrao,
                conteudo
            )
    
            if resultado:
    
                bloco_tags = resultado.group(
                    1
                ).strip()
    
                break
    
        if bloco_tags:
    
            tags_extraidas = []
    
            for parte in re.split(
                r"[,;\n]+",
                bloco_tags
            ):
    
                tag = re.sub(
                    r"^\s*[-•*]\s*",
                    "",
                    parte
                ).strip()
    
                tag = re.sub(
                    r"^\s*\d+[\.\)\-:]\s*",
                    "",
                    tag
                ).strip()
    
                if tag:
    
                    tags_extraidas.append(
                        tag
                    )
    
            tags = tags_extraidas
            recebeu_tags = True
    
        pagina[
            "status"
        ] = (
            "gerada"
            if conteudo
            else "erro"
        )
    
    # ========================================================
    # CATEGORIA: ARQUIVO DE ORIGEM
    # ========================================================
    
    elif categoria == "arquivo_origem":
    
        pagina[
            "arquivo_origem"
        ] = str(
            texto or ""
        ).strip()
    
    # ========================================================
    # CATEGORIA: TÍTULO
    # ========================================================
    
    elif categoria == "titulo":
    
        pagina[
            "titulo"
        ] = str(
            texto or ""
        ).strip()
    
    # ========================================================
    # CATEGORIA: SUBTÍTULO
    # ========================================================
    
    elif categoria == "subtitulo":
    
        pagina[
            "subtitulo"
        ] = str(
            texto or ""
        ).strip()
    
    # ========================================================
    # CATEGORIA: TAGS
    # ========================================================
    
    elif categoria == "tags":
    
        tags = normalizar_lista_local(
            texto
        )
    
        recebeu_tags = True
    
    # ========================================================
    # CATEGORIA: SEGMENTOS
    # ========================================================
    
    elif categoria == "segmentos":
    
        segmentos = normalizar_lista_local(
            texto
        )
    
        recebeu_segmentos = True
    
    # ========================================================
    # CATEGORIA: BLOCOS
    # ========================================================
    
    elif categoria == "blocos":
    
        if isinstance(
            texto,
            list
        ):
    
            for bloco in texto:
    
                if not isinstance(
                    bloco,
                    dict
                ):
    
                    continue
    
                try:
    
                    numero = int(
                        bloco.get(
                            "numero",
                            0
                        )
                    )
    
                except Exception:
    
                    numero = 0
    
                atualizar_bloco(
                    numero,
                    bloco
                )
    
    # ========================================================
    # CATEGORIA: MAPA MEAD
    # ========================================================
    
    elif categoria == "mapa_mead":
    
        valor = str(
            texto or ""
        ).strip()
    
        dados_tema[
            "mapa_mead"
        ] = {
    
            "status":
                "gerado"
                if valor
                else "vazio",
    
            "texto":
                valor
        }
    
    # ========================================================
    # CATEGORIA: ESTRUTURA
    # ========================================================
    
    elif categoria == "estrutura":
    
        pass
    
    
    # ========================================================
    # NORMALIZAR TAGS FINAIS
    #
    # NUNCA substituir tags existentes por [] apenas porque
    # esta chamada não recebeu tags.
    # ========================================================
    
    tags_atuais = dados_tema.get(
        "tags",
        []
    )
    
    if not isinstance(
        tags_atuais,
        list
    ):
    
        tags_atuais = []
    
    tags_finais = []
    
    vistos_tags = set()
    
    for tag in normalizar_lista_local(
        tags_atuais
    ):
    
        try:
    
            chave_tag = normalizar_tema_chave(
                tag
            )
    
        except Exception:
    
            chave_tag = tag.lower()
    
        if chave_tag in vistos_tags:
    
            continue
    
        vistos_tags.add(
            chave_tag
        )
    
        tags_finais.append(
            tag
        )
    
    dados_tema[
        "tags"
    ] = tags_finais[:30]
    
    # ========================================================
    # NORMALIZAR 12 SEGMENTOS
    # ========================================================
    
    segmentos_finais = []
    
    vistos_segmentos = set()
    
    for segmento in normalizar_lista_local(
        segmentos_textuais
    ):
    
        try:
    
            chave_segmento = normalizar_tema_chave(
                segmento
            )
    
        except Exception:
    
            chave_segmento = segmento.lower()
    
        if chave_segmento in vistos_segmentos:
    
            continue
    
        vistos_segmentos.add(
            chave_segmento
        )
    
        segmentos_finais.append(
            segmento
        )
    
    segmentos_finais = segmentos_finais[:12]
    
    # ========================================================
    # GARANTIR E GRAVAR 12 SEGMENTOS
    # ========================================================
    
    segmentos_listas_existentes = pagina.get(
        "segmentos_listas",
        {}
    )
    
    if not isinstance(
        segmentos_listas_existentes,
        dict
    ):
    
        segmentos_listas_existentes = {}
    
    segmentos_listas_finais = {}
    
    # --------------------------------------------------------
    # Preservar listas existentes.
    # --------------------------------------------------------
    
    for numero in range(
        1,
        13
    ):
    
        chave_segmento = (
            f"segmento_{numero}"
        )
    
        lista_existente = (
            segmentos_listas_existentes.get(
                chave_segmento,
                []
            )
        )
    
        if not isinstance(
            lista_existente,
            list
        ):
    
            lista_existente = []
    
        segmentos_listas_finais[
            chave_segmento
        ] = list(
            lista_existente
        )
    
    # ========================================================
    # GRAVAR OS 12 SEGMENTOS RECEBIDOS PELO PYTHON
    # ========================================================

    if segmentos_finais:

        for numero, segmento in enumerate(
            segmentos_finais,
            start=1
        ):

            if numero > 12:
                break

            segmentos_listas_finais[
                f"segmento_{numero}"
            ] = [
                segmento
            ]

    pagina[
        "segmentos_listas"
    ] = segmentos_listas_finais
    

    
    # ========================================================
    # POSICIONAMENTO DAS LISTAS
    # ========================================================
    
    posicionamento = pagina.get(
        "posicionamento_listas",
        {}
    )
    
    if not isinstance(
        posicionamento,
        dict
    ):
    
        posicionamento = {}
    
    pagina[
        "posicionamento_listas"
    ] = {
    
        "bloco":
            posicionamento.get(
                "bloco",
                None
            )
    }
    
    # ========================================================
    # GARANTIR 6 IMAGENS
    # ========================================================
    
    imagens = pagina.get(
        "imagens",
        {}
    )
    
    if not isinstance(
        imagens,
        dict
    ):
    
        imagens = {}
    
    imagens_oficiais = {}
    
    for numero in range(
        1,
        7
    ):
    
        chave_imagem = (
            f"imagem_{numero}"
        )
    
        imagem = imagens.get(
            chave_imagem,
            {}
        )
    
        if not isinstance(
            imagem,
            dict
        ):
    
            imagem = {}
    
        imagens_oficiais[
            chave_imagem
        ] = {
    
            "url":
                str(
                    imagem.get(
                        "url",
                        ""
                    )
                    or ""
                ).strip(),
    
            "arquivo":
                str(
                    imagem.get(
                        "arquivo",
                        ""
                    )
                    or ""
                ).strip(),
    
            "alt":
                str(
                    imagem.get(
                        "alt",
                        ""
                    )
                    or ""
                ).strip(),
    
            "descricao":
                str(
                    imagem.get(
                        "descricao",
                        ""
                    )
                    or ""
                ).strip()
        }
    
    pagina[
        "imagens"
    ] = imagens_oficiais
    
    # ========================================================
    # DADOS BÁSICOS
    # ========================================================
    
    pagina[
        "tema"
    ] = tema_original
    
    pagina[
        "arquivo_origem"
    ] = str(
        pagina.get(
            "arquivo_origem",
            ""
        )
        or ""
    ).strip()
    
    pagina[
        "h1"
    ] = tema_original
    
    pagina[
        "titulo"
    ] = tema_original
    
    pagina[
        "subtitulo"
    ] = str(
        pagina.get(
            "subtitulo",
            ""
        )
        or ""
    ).strip()
    
    pagina[
        "descricao"
    ] = str(
        pagina.get(
            "descricao",
            ""
        )
        or ""
    ).strip()
    
    # ========================================================
    # GARANTIR 5 BLOCOS
    # ========================================================
    
    for numero in range(
        1,
        6
    ):
    
        chave_bloco = (
            f"bloco_{numero}"
        )
    
        bloco = pagina.get(
            chave_bloco
        )
    
        if not isinstance(
            bloco,
            dict
        ):
    
            bloco = criar_bloco_vazio(
                numero
            )
    
        bloco[
            "id"
        ] = str(
            bloco.get(
                "id",
                f"bloco_{numero}"
            )
            or f"bloco_{numero}"
        ).strip()
    
        bloco[
            "informacoes_relevantes"
        ] = str(
            bloco.get(
                "informacoes_relevantes",
                ""
            )
            or ""
        ).strip()
    
        bloco[
            "titulo"
        ] = str(
            bloco.get(
                "titulo",
                ""
            )
            or ""
        ).strip()
    
        paragrafos = bloco.get(
            "paragrafos",
            []
        )
    
        if not isinstance(
            paragrafos,
            list
        ):
    
            paragrafos = []
    
        paragrafos_limpos = []
    
        for paragrafo in paragrafos[:3]:
    
            paragrafos_limpos.append(
                str(
                    paragrafo or ""
                ).strip()
            )
    
        bloco[
            "paragrafos"
        ] = (
            paragrafos_limpos
            + [
                "",
                "",
                ""
            ]
        )[:3]
    
        conteudo_hash = (
            bloco.get(
                "informacoes_relevantes",
                ""
            )
            + "|"
            + bloco.get(
                "titulo",
                ""
            )
            + "|"
            + "|".join(
                bloco.get(
                    "paragrafos",
                    []
                )[:3]
            )
        )
    
        bloco[
            "hash"
        ] = hashlib.sha256(
            conteudo_hash.encode(
                "utf-8"
            )
        ).hexdigest()
    
        pagina[
            chave_bloco
        ] = bloco
    
    # ========================================================
    # CONTAGEM DE REPETIÇÕES
    # ========================================================
    
    texto_para_contagem = []
    
    for numero in range(
        1,
        6
    ):
    
        bloco = pagina.get(
            f"bloco_{numero}",
            {}
        )
    
        if not isinstance(
            bloco,
            dict
        ):
    
            continue
    
        titulo_bloco = str(
            bloco.get(
                "titulo",
                ""
            )
            or ""
        ).strip()
    
        if titulo_bloco:
    
            texto_para_contagem.append(
                titulo_bloco
            )
    
        info_bloco = str(
            bloco.get(
                "informacoes_relevantes",
                ""
            )
            or ""
        ).strip()
    
        if info_bloco:
    
            texto_para_contagem.append(
                info_bloco
            )
    
        paragrafos = bloco.get(
            "paragrafos",
            []
        )
    
        if isinstance(
            paragrafos,
            list
        ):
    
            texto_para_contagem.extend(
                str(
                    paragrafo or ""
                ).strip()
                for paragrafo in paragrafos[:3]
                if str(
                    paragrafo or ""
                ).strip()
            )
    
    texto_para_contagem_final = " ".join(
        texto_para_contagem
    )
    
    try:
    
        palavra_chave = tema_original.strip()
    
        if palavra_chave:
    
            padrao_chave = re.escape(
                palavra_chave
            )
    
            repeticoes = len(
                re.findall(
                    rf"(?<!\w){padrao_chave}(?!\w)",
                    texto_para_contagem_final,
                    flags=re.IGNORECASE
                )
            )
    
        else:
    
            repeticoes = 0
    
    except Exception:
    
        repeticoes = 0
    
    dados_tema[
        "controle_repeticoes"
    ] = {
    
        "palavra_chave":
            tema_original,
    
        "meta_repeticoes":
            60,
    
        "repeticoes_realizadas":
            repeticoes,
    
        "repeticoes_faltantes":
            max(
                0,
                60 - repeticoes
            )
    }
    
    # ========================================================
    # CALCULAR CARACTERES
    # ========================================================
    
    textos_pagina = []
    
    for numero in range(
        1,
        6
    ):
    
        bloco = pagina.get(
            f"bloco_{numero}",
            {}
        )
    
        if not isinstance(
            bloco,
            dict
        ):
    
            continue
    
        info_bloco = str(
            bloco.get(
                "informacoes_relevantes",
                ""
            )
            or ""
        ).strip()
    
        if info_bloco:
    
            textos_pagina.append(
                info_bloco
            )
    
        titulo_bloco = str(
            bloco.get(
                "titulo",
                ""
            )
            or ""
        ).strip()
    
        if titulo_bloco:
    
            textos_pagina.append(
                titulo_bloco
            )
    
        paragrafos = bloco.get(
            "paragrafos",
            []
        )
    
        if isinstance(
            paragrafos,
            list
        ):
    
            for paragrafo in paragrafos[:3]:
    
                paragrafo = str(
                    paragrafo or ""
                ).strip()
    
                if paragrafo:
    
                    textos_pagina.append(
                        paragrafo
                    )
    
    pagina[
        "caracteres"
    ] = len(
        "\n\n".join(
            textos_pagina
        )
    )
    
    # ========================================================
    # STATUS
    # ========================================================
    
    if categoria == "conteudo_completo":
    
        pagina[
            "status"
        ] = (
            "gerada"
            if str(
                texto or ""
            ).strip()
            else "erro"
        )
    
    elif categoria == "pagina":
    
        pagina[
            "status"
        ] = str(
            pagina.get(
                "status",
                "em_construcao"
            )
            or "em_construcao"
        ).strip()
    
    else:
    
        if not pagina.get(
            "status"
        ):
    
            pagina[
                "status"
            ] = "em_construcao"
    
    # ========================================================
    # REMOVER INFORMACOES_RELEVANTES DA PÁGINA
    # ========================================================
    
    pagina.pop(
        "informacoes_relevantes",
        None
    )
    
    # ========================================================
    # MONTAR PÁGINA FINAL
    # ========================================================
    
    pagina_final = {
    
        "tema":
            pagina.get(
                "tema",
                tema_original
            ),
    
        "arquivo_origem":
            pagina.get(
                "arquivo_origem",
                ""
            ),
    
        "h1":
            pagina.get(
                "h1",
                tema_original
            ),
    
        "titulo":
            pagina.get(
                "titulo",
                tema_original
            ),
    
        "subtitulo":
            pagina.get(
                "subtitulo",
                ""
            ),
    
        "descricao":
            pagina.get(
                "descricao",
                ""
            )
    }
    
    # ========================================================
    # 5 BLOCOS
    # ========================================================
    
    for numero in range(
        1,
        6
    ):
    
        chave_bloco = (
            f"bloco_{numero}"
        )
    
        pagina_final[
            chave_bloco
        ] = pagina[
            chave_bloco
        ]
    
    # ========================================================
    # 12 SEGMENTOS
    # ========================================================
    
    pagina_final[
        "segmentos_listas"
    ] = pagina[
        "segmentos_listas"
    ]
    
    # ========================================================
    # POSICIONAMENTO
    # ========================================================
    
    pagina_final[
        "posicionamento_listas"
    ] = pagina[
        "posicionamento_listas"
    ]
    
    # ========================================================
    # 6 IMAGENS
    # ========================================================
    
    pagina_final[
        "imagens"
    ] = pagina[
        "imagens"
    ]
    
    # ========================================================
    # CARACTERES
    # ========================================================
    
    pagina_final[
        "caracteres"
    ] = pagina.get(
        "caracteres",
        0
    )
    
    # ========================================================
    # STATUS
    # ========================================================
    
    pagina_final[
        "status"
    ] = pagina.get(
        "status",
        "em_construcao"
    )
    
    # ========================================================
    # DADOS FINAIS DO TEMA
    # ========================================================
    
    dados_tema_final = {
        "tema": tema_original,
    
        "grupo": dados_tema.get(
            "grupo",
            ""
        ),
    
        "tipo": dados_tema.get(
            "tipo",
            ""
        ),
    
        "tags": dados_tema.get(
            "tags",
            []
        )[:30],
    
        "controle_repeticoes": dados_tema.get(
            "controle_repeticoes",
            {}
        ),
    
        "mapa_mead": dados_tema.get(
            "mapa_mead",
            {
                "status": "",
                "texto": ""
            }
        ),
    
        "grupo_principal_projeto":
            grupo_principal_projeto,
    
        "pagina": pagina_final
    }
    
    # ========================================================
    # PROTEÇÃO FINAL ABSOLUTA
    # ========================================================
    
    dados_tema_final.pop(
        "informacoes_relevantes",
        None
    )
    
    pagina_final.pop(
        "informacoes_relevantes",
        None
    )
    
    # ========================================================
    # GARANTIR QUE NÃO EXISTE INFORMACOES_RELEVANTES
    # GLOBAL NO TEMA
    # ========================================================
    
    dados_tema_final.pop(
        "informacoes_relevantes",
        None
    )
    
    # ========================================================
    # SALVAR NA CHAVE DO TEMA
    # ========================================================
    
    banco[
        chave_existente
    ] = dados_tema_final
    
    # ========================================================
    # REMOVER CHAVES DUPLICADAS DO MESMO TEMA
    # ========================================================
    
    for chave in list(
        banco.keys()
    ):
    
        if chave == chave_existente:
    
            continue
    
        if not isinstance(
            chave,
            str
        ):
    
            continue
    
        try:
    
            chave_normalizada = normalizar_tema_chave(
                chave
            )
    
        except Exception:
    
            chave_normalizada = (
                chave
                .strip()
                .lower()
            )
    
        if chave_normalizada == tema_normalizado:
    
            del banco[
                chave
            ]
    
    # ========================================================
    # SALVAR JSON
    # ========================================================
    
    try:
    
        diretorio = os.path.dirname(
            ARQUIVO_BANCO
        )
    
        if diretorio:
    
            os.makedirs(
                diretorio,
                exist_ok=True
            )
    
        with open(
            ARQUIVO_BANCO,
            "w",
            encoding="utf-8"
        ) as arquivo:
    
            json.dump(
                banco,
                arquivo,
                ensure_ascii=False,
                indent=4
            )
    
        # ====================================================
        # CONFERÊNCIA FINAL
        # ====================================================
    
        total_blocos = 0
        total_paragrafos = 0
        total_fragmentos = 0
    
        for numero in range(
            1,
            6
        ):
    
            bloco = pagina_final.get(
                f"bloco_{numero}",
                {}
            )
    
            if not isinstance(
                bloco,
                dict
            ):
    
                continue
    
            total_blocos += 1
    
            info_bloco = str(
                bloco.get(
                    "informacoes_relevantes",
                    ""
                )
                or ""
            ).strip()
    
            if info_bloco:
    
                total_fragmentos += 1
    
            paragrafos = bloco.get(
                "paragrafos",
                []
            )
    
            if isinstance(
                paragrafos,
                list
            ):
    
                total_paragrafos += len(
                    paragrafos[:3]
                )
    
        print()
        print(
            "=============================================="
        )
        print(
            "BANCO JSON OFICIAL SALVO"
        )
        print(
            "=============================================="
        )
        print(
            "TEMA:",
            tema_original
        )
        print(
            "CATEGORIA:",
            categoria
        )
        print(
            "ARQUIVO:",
            ARQUIVO_BANCO
        )
        print(
            "BLOCOS:",
            total_blocos
        )
        print(
            "PARÁGRAFOS:",
            total_paragrafos
        )
        print(
            "BLOCOS COM INFORMAÇÕES RELEVANTES:",
            total_fragmentos
        )
        print(
            "SEGMENTOS:",
            len(
                segmentos_finais
            )
        )
        print(
            "IMAGENS:",
            6
        )
        print(
            "TAGS:",
            len(
                dados_tema_final.get(
                    "tags",
                    []
                )
            )
        )
        print(
            "=============================================="
        )
    
        # ----------------------------------------------------
        # CHECK DOS 5 BLOCOS
        # ----------------------------------------------------
    
        print(
            "\nCHECK DOS 5 BLOCOS:"
        )
    
        for numero in range(
            1,
            6
        ):
    
            bloco = pagina_final.get(
                f"bloco_{numero}",
                {}
            )
    
            if not isinstance(
                bloco,
                dict
            ):
    
                print(
                    f"🔴 bloco_{numero}: inválido"
                )
    
                continue
    
            info = str(
                bloco.get(
                    "informacoes_relevantes",
                    ""
                )
                or ""
            ).strip()
    
            paragrafos = bloco.get(
                "paragrafos",
                []
            )
    
            quantidade_paragrafos = (
                len(
                    paragrafos[:3]
                )
                if isinstance(
                    paragrafos,
                    list
                )
                else 0
            )
    
            print(
                f"{'🟢' if info else '🔴'} "
                f"bloco_{numero}: "
                f"{len(info)} caracteres de informações relevantes | "
                f"{quantidade_paragrafos} parágrafos"
            )
    
        # ----------------------------------------------------
        # CHECK GLOBAL
        # ----------------------------------------------------
    
        print(
            "\nCHECK FINAL DE INTEGRIDADE:"
        )
    
        print(
            "informacoes_relevantes global:",
            "🔴 EXISTE"
            if "informacoes_relevantes"
            in dados_tema_final
            else "🟢 NÃO EXISTE"
        )
    
        print(
            "5 blocos:",
            "🟢"
            if all(
                isinstance(
                    pagina_final.get(
                        f"bloco_{numero}"
                    ),
                    dict
                )
                for numero in range(1, 6)
            )
            else "🔴"
        )
    
        print(
            "15 parágrafos:",
            "🟢"
            if total_paragrafos == 15
            else f"🔴 ({total_paragrafos})"
        )
    
        print(
            "12 segmentos:",
            "🟢"
            if len(
                segmentos_finais
            ) == 12
            else f"⚠️ ({len(segmentos_finais)})"
        )
    
        print(
            "30 tags:",
            "🟢"
            if len(
                dados_tema_final.get(
                    "tags",
                    []
                )
            ) == 30
            else f"⚠️ ({len(dados_tema_final.get('tags', []))})"
        )
    
        print(
            "=============================================="
        )
    
        return True
    
    except Exception as erro:
    
        print()
        print(
            "=============================================="
        )
        print(
            "ERRO GRAVANDO BANCO"
        )
        print(
            "=============================================="
        )
        print(
            repr(erro)
        )
        print(
            "=============================================="
        )
    
        return False




# ============================================================
# INTERFACE
# ============================================================

def selecionar_banco():

    arquivo = filedialog.askopenfilename(
        filetypes=[
            ("JSON", "*.json")
        ]
    )

    if arquivo:

        entrada_banco.delete(
            0,
            tk.END
        )

        entrada_banco.insert(
            0,
            arquivo
        )


def selecionar_modelo():

    arquivo = filedialog.askopenfilename(
        filetypes=[
            ("DOCX", "*.docx")
        ]
    )

    if arquivo:

        entrada_modelo.delete(
            0,
            tk.END
        )

        entrada_modelo.insert(
            0,
            arquivo
        )



# ============================================================
# NOME DO SITE
# ============================================================

def atualizar_banco_automatico(event=None):

    nome = entrada_site.get().strip()

    if not nome:
        return

    nome_arquivo = (
        nome.lower()
        .replace(" ", "_")
        .replace("-", "_")
    )

    caminho = (
        rf"C:\Python\gerador-conteudo\banco\{nome_arquivo}.json"
    )

    entrada_banco.delete(0, tk.END)
    entrada_banco.insert(0, caminho)
    

# ============================================================
# GERAR MATERIAL INTERFACE
# ============================================================

def atualizar_progresso(valor, texto):

    janela.after(
        0,
        lambda: barra_progresso.config(
            value=valor
        )
    )

    janela.after(
        0,
        lambda: status_progresso.config(
            text=texto
        )
    )

def iniciar_geracao():

    global IA_PROCESSANDO

    IA_PROCESSANDO = True

    thread = threading.Thread(
        target=gerar_material_interface,
        daemon=True
    )

    thread.start() 
    
    
 

# ============================================================
# INTERFACE - TEMPO ESTIMADO
# ============================================================

def formatar_tempo(segundos):

    minutos = int(segundos // 60)

    segundos = int(segundos % 60)

    return f"{minutos:02d}:{segundos:02d}"



# ============================================================
# TEMPO EM TEMPO REAL
# ============================================================

def atualizar_tempo_execucao(inicio, tema):

    while IA_PROCESSANDO:

        tempo = time.time() - inicio

        texto = (
            f"Tema: {tema} | "
            f"Tempo: {formatar_tempo(tempo)}"
        )

        janela.after(
            0,
            lambda: status_tempo.config(
                text=texto
            )
        )

        time.sleep(1)
        
        

# ============================================================
# COMPATIBILIDADE TEMPO TOTAL
# ============================================================

def atualizar_tempo_total(inicio, tema):

    print("FUNÇÃO atualizar_tempo_total CARREGADA")

    return atualizar_tempo_execucao(
        inicio,
        tema
    )



# ============================================================
# MONTAR PÁGINA OFICIAL DO JSON
# ============================================================

def montar_pagina_json(
        tema,
        grupo="",
        tipo="",
        tags=None,
        controle_repeticoes=None,
        mapa_mead=None,
        informacoes_relevantes=None,
        dados_pagina=None,
        grupo_principal_projeto=""
    ):

    """
    Monta a estrutura oficial da página que será gravada
    posteriormente no conteudo-site.json.

    Esta função:

    - cria a estrutura oficial;
    - recebe somente os dados pertencentes à nova estrutura;
    - preserva os campos existentes;
    - garante os 5 blocos;
    - garante os 12 segmentos;
    - garante as 6 imagens;
    - prepara tags;
    - prepara mapa_mead;
    - prepara informações relevantes por bloco;
    - prepara grupo_principal_projeto;
    - calcula os caracteres;
    - deixa a página pronta para gravação.

    IMPORTANTE:
    - Não pesquisa.
    - Não seleciona textos.
    - Não chama Ollama.
    - Não inventa informações.
    - Não grava o arquivo.
    - NÃO cria informacoes_adicionais.
    - NÃO cria categorias.
    """

    global pagina

    try:

        # ----------------------------------------------------
        # NORMALIZAR TEMA
        # ----------------------------------------------------

        tema = str(
            tema or ""
        ).strip()

        if not tema:

            print(
                "❌ Não foi possível montar JSON: "
                "tema vazio."
            )

            return None

        # ----------------------------------------------------
        # NORMALIZAR ENTRADAS
        # ----------------------------------------------------

        if not isinstance(
            tags,
            list
        ):

            tags = []

        if not isinstance(
            controle_repeticoes,
            dict
        ):

            controle_repeticoes = {}

        if not isinstance(
            mapa_mead,
            dict
        ):

            mapa_mead = {}

        if not isinstance(
            informacoes_relevantes,
            dict
        ):

            informacoes_relevantes = {}

        if not isinstance(
            dados_pagina,
            dict
        ):

            dados_pagina = {}

        grupo_principal_projeto = str(
            grupo_principal_projeto or ""
        ).strip()

        # ----------------------------------------------------
        # CRIAR ESTRUTURA OFICIAL
        # ----------------------------------------------------

        estrutura = criar_estrutura_json_pagina(
            tema
        )

        if not isinstance(
            estrutura,
            dict
        ):

            print(
                "❌ criar_estrutura_json_pagina() "
                "não retornou um dicionário."
            )

            return None

        if tema not in estrutura:

            print(
                "❌ Tema não encontrado na estrutura "
                "criada para o JSON."
            )

            return None

        pagina = estrutura[
            tema
        ]

        # ----------------------------------------------------
        # TEMA / GRUPO / TIPO
        # ----------------------------------------------------

        pagina[
            "tema"
        ] = tema

        pagina[
            "grupo"
        ] = str(
            grupo or ""
        ).strip()

        pagina[
            "tipo"
        ] = str(
            tipo or ""
        ).strip()

        # ----------------------------------------------------
        # GRUPO PRINCIPAL DO PROJETO
        #
        # VEM DA INTERFACE E FICA DIRETAMENTE NO TEMA.
        # NÃO FICA DENTRO DE informacoes_adicionais.
        # ----------------------------------------------------

        pagina[
            "grupo_principal_projeto"
        ] = grupo_principal_projeto

        # ----------------------------------------------------
        # TAGS
        # ----------------------------------------------------

        tags_finais = []

        for tag in tags:

            tag_limpa = str(
                tag or ""
            ).strip()

            if not tag_limpa:
                continue

            if tag_limpa in tags_finais:
                continue

            tags_finais.append(
                tag_limpa
            )

            if len(
                tags_finais
            ) >= 30:

                break

        pagina[
            "tags"
        ] = tags_finais

        # ----------------------------------------------------
        # CONTROLE DE REPETIÇÕES
        # ----------------------------------------------------

        palavra_controle = str(
            controle_repeticoes.get(
                "palavra_chave",
                tema
            )
            or tema
        ).strip()

        meta_repeticoes = controle_repeticoes.get(
            "meta_repeticoes",
            60
        )

        repeticoes_realizadas = controle_repeticoes.get(
            "repeticoes_realizadas",
            0
        )

        repeticoes_faltantes = controle_repeticoes.get(
            "repeticoes_faltantes",
            60
        )

        pagina[
            "controle_repeticoes"
        ] = {

            "palavra_chave":
                palavra_controle,

            "meta_repeticoes":
                meta_repeticoes,

            "repeticoes_realizadas":
                repeticoes_realizadas,

            "repeticoes_faltantes":
                repeticoes_faltantes
        }

        # ----------------------------------------------------
        # MAPA MEAD
        # ----------------------------------------------------

        pagina[
            "mapa_mead"
        ] = {

            "status":
                str(
                    mapa_mead.get(
                        "status",
                        ""
                    )
                    or ""
                ).strip(),

            "texto":
                str(
                    mapa_mead.get(
                        "texto",
                        ""
                    )
                    or ""
                ).strip()
        }

        # ----------------------------------------------------
        # INFORMAÇÕES RELEVANTES
        #
        # OS 15 FRAGMENTOS SELECIONADOS PELO PYTHON
        # FICAM DISTRIBUÍDOS NOS 5 BLOCOS.
        #
        # NÃO EXISTE MAIS CAMPO GLOBAL.
        # ----------------------------------------------------

        fragmentos_blocos = (
            informacoes_relevantes.get(
                "blocos",
                {}
            )
        )

        if not isinstance(
            fragmentos_blocos,
            dict
        ):

            fragmentos_blocos = {}

        for numero_bloco in range(
            1,
            6
        ):

            chave_bloco = (
                f"bloco_{numero_bloco}"
            )

            informacoes_bloco = (
                fragmentos_blocos.get(
                    chave_bloco,
                    []
                )
            )

            # -----------------------------------------------
            # NORMALIZAR FRAGMENTOS DO BLOCO
            # -----------------------------------------------

            if isinstance(
                informacoes_bloco,
                list
            ):

                fragmentos_limpos = []

                for fragmento in informacoes_bloco:

                    fragmento_limpo = str(
                        fragmento or ""
                    ).strip()

                    if fragmento_limpo:

                        fragmentos_limpos.append(
                            fragmento_limpo
                        )

                informacoes_bloco = (
                    "\n\n".join(
                        fragmentos_limpos
                    )
                )

            else:

                informacoes_bloco = str(
                    informacoes_bloco or ""
                ).strip()

            # -----------------------------------------------
            # GARANTIR BLOCO
            # -----------------------------------------------

            if chave_bloco not in pagina:

                pagina[
                    chave_bloco
                ] = {

                    "id":
                        chave_bloco,

                    "hash":
                        "",

                    "informacoes_relevantes":
                        "",

                    "titulo":
                        "",

                    "paragrafos":
                        []
                }

            # -----------------------------------------------
            # GRAVAR INFORMAÇÕES RELEVANTES
            # DIRETAMENTE NO BLOCO
            # -----------------------------------------------

            pagina[
                chave_bloco
            ][
                "informacoes_relevantes"
            ] = informacoes_bloco

        # ----------------------------------------------------
        # DADOS GERAIS DA PÁGINA
        # ----------------------------------------------------

        dados_gerais = [

            "tema",
            "arquivo_origem",
            "h1",
            "titulo",
            "subtitulo",
            "descricao"

        ]

        for campo in dados_gerais:

            if campo in dados_pagina:

                valor = dados_pagina.get(
                    campo,
                    ""
                )

                pagina[
                    "pagina"
                ][
                    campo
                ] = str(
                    valor or ""
                ).strip()

        # ----------------------------------------------------
        # GARANTIR TEMA DA PÁGINA
        # ----------------------------------------------------

        if not pagina[
            "pagina"
        ].get(
            "tema"
        ):

            pagina[
                "pagina"
            ][
                "tema"
            ] = tema

        # ----------------------------------------------------
        # ATUALIZAR OS 5 BLOCOS
        #
        # IMPORTANTE:
        # atualizar_bloco() recebe os dados produzidos
        # pelo gerador de conteúdo.
        # ----------------------------------------------------

        for numero in range(
            1,
            6
        ):

            chave_bloco = (
                f"bloco_{numero}"
            )

            dados_bloco = dados_pagina.get(
                chave_bloco,
                {}
            )

            if not isinstance(
                dados_bloco,
                dict
            ):

                dados_bloco = {}

            atualizar_bloco(
                numero,
                dados_bloco
            )

            # ------------------------------------------------
            # GARANTIR QUE A INFORMAÇÃO RELEVANTE
            # NÃO SEJA APAGADA PELO atualizar_bloco()
            # ------------------------------------------------

            informacao_selecionada = (
                pagina[
                    chave_bloco
                ].get(
                    "informacoes_relevantes",
                    ""
                )
            )

            if not informacao_selecionada:

                informacao_selecionada = (
                    fragmentos_blocos.get(
                        chave_bloco,
                        []
                    )
                )

                if isinstance(
                    informacao_selecionada,
                    list
                ):

                    informacao_selecionada = (
                        "\n\n".join(
                            str(item).strip()
                            for item in informacao_selecionada
                            if str(item).strip()
                        )
                    )

                else:

                    informacao_selecionada = str(
                        informacao_selecionada or ""
                    ).strip()

                pagina[
                    chave_bloco
                ][
                    "informacoes_relevantes"
                ] = informacao_selecionada

        # ----------------------------------------------------
        # SEGMENTOS
        # ----------------------------------------------------

        segmentos_recebidos = dados_pagina.get(
            "segmentos_listas",
            {}
        )

        if not isinstance(
            segmentos_recebidos,
            dict
        ):

            segmentos_recebidos = {}

        segmentos_oficiais = {}

        for numero in range(
            1,
            13
        ):

            chave_segmento = (
                f"segmento_{numero}"
            )

            valor = segmentos_recebidos.get(
                chave_segmento,
                []
            )

            if not isinstance(
                valor,
                list
            ):

                valor = []

            segmentos_oficiais[
                chave_segmento
            ] = [

                str(
                    item or ""
                ).strip()

                for item in valor

                if str(
                    item or ""
                ).strip()
            ]

        pagina[
            "pagina"
        ][
            "segmentos_listas"
        ] = segmentos_oficiais

        # ----------------------------------------------------
        # POSICIONAMENTO DAS LISTAS
        # ----------------------------------------------------

        posicionamento = dados_pagina.get(
            "posicionamento_listas",
            {}
        )

        if not isinstance(
            posicionamento,
            dict
        ):

            posicionamento = {}

        pagina[
            "pagina"
        ][
            "posicionamento_listas"
        ][
            "bloco"
        ] = posicionamento.get(
            "bloco"
        )

        # ----------------------------------------------------
        # IMAGENS
        # ----------------------------------------------------

        imagens_recebidas = dados_pagina.get(
            "imagens",
            {}
        )

        if not isinstance(
            imagens_recebidas,
            dict
        ):

            imagens_recebidas = {}

        for numero in range(
            1,
            7
        ):

            chave_imagem = (
                f"imagem_{numero}"
            )

            imagem = imagens_recebidas.get(
                chave_imagem,
                {}
            )

            if not isinstance(
                imagem,
                dict
            ):

                imagem = {}

            pagina[
                "pagina"
            ][
                "imagens"
            ][
                chave_imagem
            ] = {

                "url":
                    str(
                        imagem.get(
                            "url",
                            ""
                        )
                        or ""
                    ).strip(),

                "arquivo":
                    str(
                        imagem.get(
                            "arquivo",
                            ""
                        )
                        or ""
                    ).strip(),

                "alt":
                    str(
                        imagem.get(
                            "alt",
                            ""
                        )
                        or ""
                    ).strip(),

                "descricao":
                    str(
                        imagem.get(
                            "descricao",
                            ""
                        )
                        or ""
                    ).strip()
            }

        # ----------------------------------------------------
        # GARANTIR QUE AS ESTRUTURAS ANTIGAS NÃO EXISTAM
        # ----------------------------------------------------

        pagina.pop(
            "informacoes_adicionais",
            None
        )

        pagina.pop(
            "categorias",
            None
        )

        pagina.get(
            "pagina",
            {}
        ).pop(
            "informacoes_adicionais",
            None
        )

        pagina.get(
            "pagina",
            {}
        ).pop(
            "categorias",
            None
        )

        # ----------------------------------------------------
        # CALCULAR CARACTERES DA PÁGINA
        # ----------------------------------------------------

        caracteres = 0

        pagina_json = pagina.get(
            "pagina",
            {}
        )

        for numero in range(
            1,
            6
        ):

            bloco = pagina_json.get(
                f"bloco_{numero}",
                {}
            )

            if not isinstance(
                bloco,
                dict
            ):

                continue

            caracteres += len(
                str(
                    bloco.get(
                        "informacoes_relevantes",
                        ""
                    )
                    or ""
                ).strip()
            )

            caracteres += len(
                str(
                    bloco.get(
                        "titulo",
                        ""
                    )
                    or ""
                ).strip()
            )

            for paragrafo in bloco.get(
                "paragrafos",
                []
            )[:3]:

                caracteres += len(
                    str(
                        paragrafo or ""
                    ).strip()
                )

        pagina_json[
            "caracteres"
        ] = caracteres

        # ----------------------------------------------------
        # STATUS
        # ----------------------------------------------------

        pagina_json[
            "status"
        ] = "pronta_para_gravacao"

        # ----------------------------------------------------
        # RETORNO
        # ----------------------------------------------------

        print()
        print(
            "=========================================="
        )
        print(
            "✅ PÁGINA OFICIAL MONTADA"
        )
        print(
            "=========================================="
        )

        print(
            f"   Tema: {tema}"
        )

        print(
            f"   Grupo: {grupo}"
        )

        print(
            f"   Grupo principal: "
            f"{grupo_principal_projeto}"
        )

        print(
            f"   Tags: {len(tags_finais)}"
        )

        print(
            "   Blocos: 5"
        )

        blocos_com_info = 0

        total_fragmentos = 0

        for numero in range(
            1,
            6
        ):

            bloco = pagina.get(
                f"bloco_{numero}",
                {}
            )

            info = str(
                bloco.get(
                    "informacoes_relevantes",
                    ""
                )
                or ""
            ).strip()

            if info:

                blocos_com_info += 1

                total_fragmentos += len(
                    [
                        x
                        for x in info.split(
                            "\n\n"
                        )
                        if x.strip()
                    ]
                )

        print(
            f"   Blocos com informações: "
            f"{blocos_com_info}/5"
        )

        print(
            f"   Fragmentos: "
            f"{total_fragmentos}/15"
        )

        print(
            "   Segmentos: 12"
        )

        print(
            "   Imagens: 6"
        )

        print(
            f"   Caracteres: {caracteres}"
        )

        print(
            "   Status: pronta_para_gravacao"
        )

        print(
            "=========================================="
        )

        return pagina

    except Exception as erro:

        print()
        print(
            "❌ ERRO ao montar página oficial:"
        )

        print(
            erro
        )

        return None




# ============================================================
# INTERFACE - SALVAR PROGRESSO SEGURO
# ============================================================

def salvar_progresso(
    tema,
    indice_tema,
    categoria,
    indice_categoria,
    status="processando",
    temas=None
):

    arquivo = ARQUIVO_PROGRESSO


    dados = {

        "temas": temas or [],

        "tema_atual": tema,

        "indice_tema": indice_tema,

        "categoria_atual": categoria,

        "indice_categoria": indice_categoria,

        "status": status,

        "ultima_atualizacao": time.strftime(
            "%d/%m/%Y %H:%M:%S"
        )

    }


    try:
    
        os.makedirs(
            os.path.dirname(arquivo),
            exist_ok=True
        )


        arquivo_temp = arquivo + ".tmp"


        with open(
            arquivo_temp,
            "w",
            encoding="utf-8"
        ) as f:

            json.dump(
                dados,
                f,
                ensure_ascii=False,
                indent=4
            )


        os.replace(
            arquivo_temp,
            arquivo
        )


        print()
        print("==============================")
        print("PROGRESSO SALVO")
        print("==============================")
        print(dados)


    except Exception as e:

        print()
        print("ERRO AO SALVAR PROGRESSO:")
        print(e)


# ============================================================
# LIMPAR PROGRESSO ANTIGO
# ============================================================

def limpar_progresso():

    arquivo = ARQUIVO_PROGRESSO


    if os.path.exists(arquivo):

        os.remove(arquivo)

        print()
        print("==============================")
        print("PROGRESSO ANTIGO REMOVIDO")
        print("==============================")


    else:

        print()
        print("==============================")
        print("NENHUM PROGRESSO ANTIGO")
        print("==============================")
        

# ============================================================
# VERIFICAR NOVO CONJUNTO DE TEMAS
# ============================================================

def verificar_novo_conjunto(temas_atuais):

    arquivo = ARQUIVO_PROGRESSO


    if not os.path.exists(arquivo):

        print()
        print("==============================")
        print("PRIMEIRO PROCESSAMENTO")
        print("==============================")

        return True


    try:

        with open(
            arquivo,
            "r",
            encoding="utf-8"
        ) as f:

            progresso = json.load(f)


        temas_salvos = progresso.get(
            "temas",
            []
        )


        if temas_salvos != temas_atuais:

            print()
            print("==============================")
            print("NOVO CONJUNTO DE TEMAS")
            print("==============================")


            limpar_progresso()


            print()
            print("==============================")
            print("INICIANDO NOVO PROCESSAMENTO")
            print("==============================")


            return True



        print()
        print("==============================")
        print("MESMO CONJUNTO IDENTIFICADO")
        print("RETOMANDO PROCESSAMENTO")
        print("==============================")


        return False



    except Exception as e:

        print()
        print("ERRO AO VERIFICAR CONJUNTO")
        print(e)

        return True


# ============================================================
# INTERFACE - SALVAR PROGRESSO RETORNAR DE ONDE PAROU
# ============================================================

def carregar_progresso():

    arquivo = r"C:\Python\gerador-conteudo\banco\progresso.json"


    if not os.path.exists(arquivo):

        return None


    with open(
        arquivo,
        "r",
        encoding="utf-8"
    ) as f:

        return json.load(f)
        

# ============================================================
# VALIDAR CONTEÚDO RELEVANTE
# ============================================================

def validar_conteudo_relevante(
    texto,
    tema
):

    if not texto:
        return False


    # ========================================================
    # 01. NORMALIZA TEXTO
    # ========================================================

    def normalizar(valor):

        valor = valor.lower()

        valor = unicodedata.normalize(
            "NFKD",
            valor
        ).encode(
            "ASCII",
            "ignore"
        ).decode(
            "utf-8"
        )

        return valor



    texto_limpo = normalizar(texto)

    tema_limpo = normalizar(tema)



    # ========================================================
    # 02. PALAVRAS PRINCIPAIS DO TEMA
    # ========================================================
    
    tema_limpo = (
        tema_limpo
        .replace("_", " ")
        .replace("-", " ")
    )
    
    
    palavras_tema = tema_limpo.split()
    
    
    encontrados_tema = 0
    
    
    for palavra in palavras_tema:
    
        if palavra in texto_limpo:
    
            encontrados_tema += 1



    # ========================================================
    # 03. VALIDACAO FLEXIVEL DO TEMA
    # ========================================================

    minimo_necessario = 1


    if len(palavras_tema) >= 3:

        minimo_necessario = 2



    if encontrados_tema < minimo_necessario:


        print(
            "CONTEUDO REJEITADO - TEMA AUSENTE:",
            tema
        )


        return False

    print(
        "CONTEUDO RELEVANTE APROVADO:",
        tema
    )


    return True



# ============================================================
# VALIDAR CONTEÚDO TÉCNICO
# ============================================================

def validar_conteudo_tecnico(
    texto,
    tema
):

    # ========================================================
    # 01. VERIFICAR CONTEÚDO
    # ========================================================

    if not texto:

        print()
        print("==============================")
        print("CONTEUDO TECNICO REJEITADO")
        print("==============================")
        print("MOTIVO: CONTEUDO VAZIO")

        return False


    # ========================================================
    # 02. NORMALIZAR TEXTO
    # ========================================================

    texto = str(
        texto
    ).strip()


    if not texto:

        print()
        print("==============================")
        print("CONTEUDO TECNICO REJEITADO")
        print("==============================")
        print("MOTIVO: CONTEUDO VAZIO")

        return False


    # ========================================================
    # 03. LIMITAR TEXTO ENVIADO À IA
    # ========================================================

    texto_validacao = texto[
        :5000
    ]


    # ========================================================
    # 04. PROMPT SEMÂNTICO
    # ========================================================

    prompt = f"""
Analise o conteúdo abaixo exclusivamente para determinar
se ele possui informação técnica útil e relacionada ao tema.

TEMA:
{tema}

CONTEÚDO:
{texto_validacao}

Avalie semanticamente o conteúdo.

Não utilize listas fixas de palavras.
Não conte palavras técnicas.
Não utilize regras por segmento.
Não exija palavras específicas.
Não rejeite o conteúdo apenas porque determinado termo técnico
não aparece.

Considere o significado e o contexto geral do conteúdo.

Determine:

1. se o conteúdo é relevante para o tema;
2. se contém informação tecnicamente útil;
3. se pode contribuir para a construção de uma página MEAD;
4. o nível técnico de 0 a 10;
5. o motivo da decisão.

Responda SOMENTE em JSON válido:

{{
  "relevante": true,
  "tecnico": true,
  "util_para_mead": true,
  "nivel_tecnico": 0,
  "motivo": "..."
}}
"""


    # ========================================================
    # 05. DEBUG
    # ========================================================

    print()
    print("==============================")
    print("VALIDANDO CONTEÚDO TÉCNICO COM IA")
    print("==============================")

    print(
        "TEMA:",
        tema
    )

    print(
        "TAMANHO ORIGINAL:",
        len(texto)
    )

    print(
        "TAMANHO ENVIADO:",
        len(texto_validacao)
    )

    print(
        "MODELO:",
        "qwen2.5:3b"
    )


    # ========================================================
    # 06. CHAMADA OLLAMA
    # ========================================================

    try:

        inicio_ollama = time.time()

        resposta = requests.post(

            "http://localhost:11434/api/generate",

            json={

                "model": "qwen2.5:3b",

                "prompt": prompt,

                "stream": False,

                "options": {

                    "num_predict": 100,

                    "temperature": 0.0,

                    "num_ctx": 4096,

                    "think": False

                }

            },

            timeout=(30, 900)

        )

        tempo_ollama = (
            time.time()
            - inicio_ollama
        )


    except requests.exceptions.Timeout:

        print()
        print("==============================")
        print("TIMEOUT NA VALIDAÇÃO TÉCNICA")
        print("==============================")

        return False


    except requests.exceptions.RequestException as e:

        print()
        print("==============================")
        print("ERRO NA VALIDAÇÃO TÉCNICA")
        print("==============================")

        print(
            "ERRO:",
            e
        )

        return False


    except Exception as e:

        print()
        print("==============================")
        print("ERRO INESPERADO NA VALIDAÇÃO")
        print("==============================")

        print(
            "ERRO:",
            e
        )

        return False


    # ========================================================
    # 07. VERIFICAR RESPOSTA HTTP
    # ========================================================

    if resposta.status_code != 200:

        print()
        print("==============================")
        print("OLLAMA RETORNOU ERRO")
        print("==============================")

        print(
            "STATUS:",
            resposta.status_code
        )

        return False


    # ========================================================
    # 08. EXTRAIR RESPOSTA
    # ========================================================

    try:

        dados = resposta.json()

        resposta_ia = dados.get(
            "response",
            ""
        )

    except Exception as e:

        print()
        print("==============================")
        print("ERRO LENDO RESPOSTA OLLAMA")
        print("==============================")

        print(
            "ERRO:",
            e
        )

        return False


    if not resposta_ia:

        print()
        print("==============================")
        print("OLLAMA NÃO RETORNOU CONTEÚDO")
        print("==============================")

        return False


    resposta_ia = resposta_ia.strip()


    # ========================================================
    # 09. LIMPAR POSSÍVEL MARKDOWN
    # ========================================================

    if resposta_ia.startswith(
        "```"
    ):

        resposta_ia = resposta_ia.replace(
            "```json",
            ""
        )

        resposta_ia = resposta_ia.replace(
            "```",
            ""
        )

        resposta_ia = resposta_ia.strip()


    # ========================================================
    # 10. INTERPRETAR JSON
    # ========================================================

    try:

        resultado = json.loads(
            resposta_ia
        )

    except Exception:

        print()
        print("==============================")
        print("RESPOSTA IA NÃO É JSON VÁLIDO")
        print("==============================")

        print(
            resposta_ia[:1000]
        )

        return False


    # ========================================================
    # 11. EXTRAIR RESULTADOS
    # ========================================================

    relevante = bool(
        resultado.get(
            "relevante",
            False
        )
    )

    tecnico = bool(
        resultado.get(
            "tecnico",
            False
        )
    )

    util_para_mead = bool(
        resultado.get(
            "util_para_mead",
            False
        )
    )

    nivel_tecnico = resultado.get(
        "nivel_tecnico",
        0
    )

    motivo = resultado.get(
        "motivo",
        ""
    )


    # ========================================================
    # 12. DEBUG RESULTADO IA
    # ========================================================

    print()
    print("==============================")
    print("RESULTADO VALIDAÇÃO TÉCNICA")
    print("==============================")

    print(
        "RELEVANTE:",
        relevante
    )

    print(
        "TÉCNICO:",
        tecnico
    )

    print(
        "ÚTIL PARA MEAD:",
        util_para_mead
    )

    print(
        "NÍVEL TÉCNICO:",
        nivel_tecnico
    )

    print(
        "MOTIVO:",
        motivo
    )

    print(
        "TEMPO OLLAMA:",
        round(
            tempo_ollama,
            2
        ),
        "segundos"
    )


    # ========================================================
    # 13. DECISÃO FINAL
    # ========================================================

    if (
        relevante
        and
        tecnico
        and
        util_para_mead
    ):

        print()
        print("==============================")
        print("CONTEUDO TECNICO APROVADO")
        print("==============================")

        return True


    print()
    print("==============================")
    print("CONTEUDO TECNICO REJEITADO")
    print("==============================")

    return False

    

# ============================================================
# NORMALIZAR TEXTO PARA VALIDAÇÃO MEAD
# ============================================================

def normalizar_texto(texto):

    texto = texto.lower()


    texto = unicodedata.normalize(
        "NFD",
        texto
    )


    texto = "".join(
        c for c in texto
        if unicodedata.category(c) != "Mn"
    )


    return texto




# ============================================================
# VALIDAR PROTAGONISTA MAPA MEAD
# ============================================================

def validar_protagonista_mead(
    resposta,
    tema
):

    resposta_limpa = normalizar_texto(
        resposta
    )


    tema_limpo = normalizar_texto(
        tema
    )


    if "protagonista" not in resposta_limpa:

        return False


    inicio = resposta_limpa.find(
        "protagonista"
    )


    trecho = resposta_limpa[
        inicio:inicio+300
    ]


    palavras_tema = tema_limpo.split()


    for palavra in palavras_tema:

        if palavra not in trecho:

            return False


    return True


# ============================================================
# VALIDAR CONTEXTO TÉCNICO MAPA MEAD
# ============================================================

def validar_contexto_tecnico_mead(
    resposta,
    tema
):


    resposta_limpa = normalizar_texto(
        resposta
    )


    bloqueados = {


        "valvulas macho": [

            "thread macho",
            "rosca macho",
            "rosca femea",
            "conexao macho",
            "conexao femea",
            "roscamento",
            "encaixa em thread"

        ]

    }


    tema_normalizado = normalizar_texto(
        tema
    )


    if tema_normalizado in bloqueados:


        for termo in bloqueados[tema_normalizado]:


            if termo in resposta_limpa:


                print()
                print("==============================")
                print("MAPA MEAD BLOQUEADO")
                print("==============================")

                print(
                    "CONTEXTO INCOMPATÍVEL:",
                    termo
                )


                return False



    return True





# ============================================================
# GERAR E SALVAR MAPA MEAD
# ============================================================

def gerar_e_salvar_mapa_mead(
    tema,
    textos
):

    # =====================================
    # NORMALIZAR TEXTOS RECEBIDOS
    # PRESERVANDO METADADOS
    # =====================================
    
    textos_normalizados = []

    for item in textos:

        if isinstance(item, dict):

            textos_normalizados.append(
                {
                    "texto": item.get("texto", ""),
                    "url": item.get("url", ""),
                    "tipo": item.get("tipo", "texto")
                }
            )

        else:

            textos_normalizados.append(
                {
                    "texto": str(item),
                    "url": "",
                    "tipo": "texto"
                }
            )

    textos = textos_normalizados


    print()
    print("==============================")
    print("GERANDO MAPA MEAD")
    print("==============================")


    if len(textos) < 3:

        print()
        print("==============================")
        print("COLETA INSUFICIENTE PARA MAPA MEAD")
        print("==============================")

        return ""


    mapa_mead = gerar_mapa_mead(
        tema,
        textos
    )
    
    
    print()
    print("==============================")
    print("MAPA MEAD RECEBIDO")
    print("==============================")
    
    print(
        str(mapa_mead)[:1000]
    )


    if mapa_mead:
    
    
        print()
        print("==============================")
        print("MAPA MEAD RECEBIDO PARA VALIDAÇÃO")
        print("==============================")
    
        print(
            "TAMANHO:",
            len(str(mapa_mead))
        )
    
    
    
        resultado_protagonista = validar_protagonista_mead(
            mapa_mead,
            tema
        )
    
    
        print(
            "VALIDAÇÃO PROTAGONISTA:",
            resultado_protagonista
        )
    
    
    
        if not resultado_protagonista:
        
            print()
            print("==============================")
            print("MAPA MEAD BLOQUEADO")
            print("PROTAGONISTA NÃO CONFERE")
            print("==============================")
        
            return ""
    
    
    
    
        resultado_contexto = validar_contexto_tecnico_mead(
            mapa_mead,
            tema
        )
    
    
        print(
            "VALIDAÇÃO CONTEXTO:",
            resultado_contexto
        )
    
    
    
        if not resultado_contexto:
        
            print()
            print("==============================")
            print("MAPA MEAD BLOQUEADO")
            print("CONTEXTO TÉCNICO INVÁLIDO")
            print("==============================")
        
            return ""
    
    
    
    
        print()
        print("==============================")
        print("CHAMANDO SALVAR BANCO")
        print("==============================")
    
    
    
        salvar_banco(
            tema,
            "mapa_mead",
            mapa_mead
        )
    
    
    
        print()
        print("==============================")
        print("MAPA MEAD SALVO")
        print("==============================")
    
    
        return mapa_mead




    print()
    print("==============================")
    print("MAPA MEAD INVALIDO")
    print("==============================")


    return ""



# ============================================================
# BUSCAR TEMAS RELACIONADOS NO BANCO
# ============================================================

def buscar_temas_relacionados(
    tema,
    mapa_mead=None
):


    relacionados = []


    # ========================================================
    # 01. PALAVRAS BASE DA BUSCA
    # ========================================================

    palavras_tema = set(
        tema.lower().split()
    )


    # ========================================================
    # 02. USAR PROTAGONISTA DO MAPA MEAD
    # ========================================================

    if mapa_mead:


        texto_mapa = mapa_mead.lower()


        if "protagonista:" in texto_mapa:


            try:

                protagonista = texto_mapa.split(
                    "protagonista:",
                    1
                )[1]


                protagonista = protagonista.split(
                    "\n",
                    1
                )[0].strip()


                if protagonista:


                    palavras_tema.update(
                        protagonista.split()
                    )


                    print()
                    print("==============================")
                    print("PROTAGONISTA MEAD USADO NA BUSCA")
                    print("==============================")

                    print(
                        protagonista
                    )


            except Exception:


                pass



    if not os.path.exists(
        ARQUIVO_BANCO
    ):

        return relacionados



    try:

        with open(
            ARQUIVO_BANCO,
            "r",
            encoding="utf-8"
        ) as arquivo:

            banco = json.load(
                arquivo
            )


    except:


        return relacionados



    # ========================================================
    # 03. BANCO PRECISA SER DICIONÁRIO
    # ========================================================

    if not isinstance(
        banco,
        dict
    ):


        return relacionados



    for tema_banco, dados in banco.items():


        if not isinstance(
            tema_banco,
            str
        ):

            continue



        tema_banco = tema_banco.lower().strip()



        if not tema_banco:

            continue



    # ========================================================
    # 04. IGNORAR O PRÓPRIO TEMA
    # ========================================================

        if tema.lower().strip() == tema_banco:

            continue



        palavras_banco = set(
            tema_banco.split()
        )



        intersecao = palavras_tema.intersection(
            palavras_banco
        )



    # ========================================================
    # 05. COMPATIBILIDADE MÍNIMA
    # ========================================================

        if len(intersecao) >= 1:


            relacionados.append(
                tema_banco
            )



    return relacionados
    


# ============================================================
# PESQUISA MAPA NO BANCO
# ============================================================

def obter_dados_banco(tema):

    print()
    print("==============================")
    print("BUSCANDO MAPA NO BANCO")
    print("==============================")
    print("TEMA:", tema)

    resultado = {

        "existe": False,

        "mapa_mead": "",

        "conteudo": "",

        "bruto": []

    }


    # ========================================================
    # 01. VERIFICAR EXISTÊNCIA DO BANCO
    # ========================================================

    if not os.path.exists(
        ARQUIVO_BANCO
    ):

        print()
        print("==============================")
        print("BANCO NÃO ENCONTRADO")
        print("==============================")

        return resultado


    # ========================================================
    # 02. LER BANCO
    # ========================================================

    try:

        with open(
            ARQUIVO_BANCO,
            "r",
            encoding="utf-8"
        ) as arquivo:

            banco = json.load(
                arquivo
            )


    except Exception as e:

        print()
        print("==============================")
        print("ERRO LENDO BANCO")
        print("==============================")

        print(e)

        return resultado


    # ========================================================
    # 03. NORMALIZAR TEMA PROCURADO
    # ========================================================

    tema_normalizado = normalizar_tema_chave(
        tema
    )


    # ========================================================
    # 04. FUNÇÃO INTERNA PARA EXTRAIR MAPA
    # ========================================================

    def extrair_mapa_texto(mapa):

        if not mapa:

            return ""


    # ========================================================
    # 05. MAPA JÁ É TEXTO
    # ========================================================

        if isinstance(
            mapa,
            str
        ):

            return mapa.strip()


    # ========================================================
    # 06. MAPA É DICIONÁRIO
    # ========================================================

        if isinstance(
            mapa,
            dict
        ):

            # -----------------------------
            # CAMPO TEXTO
            # -----------------------------

            texto = mapa.get(
                "texto",
                ""
            )

            if isinstance(
                texto,
                str
            ):

                if texto.strip():

                    return texto.strip()


            # -----------------------------
            # CAMPO GERADO
            # -----------------------------

            gerado = mapa.get(
                "gerado",
                ""
            )


            if isinstance(
                gerado,
                str
            ):

                if gerado.strip():

                    return gerado.strip()


            # -----------------------------
            # GERADO COMO DICIONÁRIO
            # -----------------------------

            if isinstance(
                gerado,
                dict
            ):

                texto_gerado = gerado.get(
                    "texto",
                    ""
                )

                if isinstance(
                    texto_gerado,
                    str
                ):

                    if texto_gerado.strip():

                        return texto_gerado.strip()


            # -----------------------------
            # ÚLTIMO RECURSO
            # -----------------------------

            try:

                texto_dict = str(
                    mapa
                )

                if texto_dict.strip():

                    return texto_dict.strip()

            except Exception:

                pass


        return ""



    # ========================================================
    # 07. FUNÇÃO INTERNA PARA VALIDAR MAPA
    # ========================================================
    
    def validar_mapa_do_tema(
        mapa_texto
    ):
    
        if not mapa_texto:
    
            return False
    
    
    # ========================================================
    # 08. VALIDAR PROTAGONISTA
    # ========================================================
    
        valido_protagonista = validar_protagonista_mead(
            mapa_texto,
            tema
        )
    
    
        print()
        print("==============================")
        print("VALIDAÇÃO DO MAPA MEAD")
        print("==============================")
    
    
        print(
            "TEMA:",
            tema
        )
    
    
        print(
            "PROTAGONISTA VÁLIDO:",
            valido_protagonista
        )
    
    
    # ========================================================
    # 09. PROTAGONISTA INVÁLIDO
    # ========================================================
    
        if not valido_protagonista:
    
            print()
            print("==============================")
            print("MAPA MEAD IGNORADO")
            print("==============================")
    
    
            print(
                "MOTIVO: PROTAGONISTA NÃO CONFERE"
            )
    
    
            return False
    
    
    # ========================================================
    # 10. MAPA VALIDADO
    # ========================================================
    
        print()
        print("==============================")
        print("MAPA MEAD VALIDADO")
        print("==============================")
    
    
        print(
            "TEMA:",
            tema
        )
    
    
        print(
            "MAPA:",
            len(mapa_texto),
            "caracteres"
        )
    
    
        return True




    # ========================================================
    # 11. BANCO EM FORMATO LISTA
    # ========================================================

    if isinstance(
        banco,
        list
    ):

        for item in banco:

            if not isinstance(
                item,
                dict
            ):

                continue


            tema_banco = item.get(
                "tema",
                ""
            )


            if not isinstance(
                tema_banco,
                str
            ):

                continue


            tema_banco_normalizado = normalizar_tema_chave(
                tema_banco
            )


    # ========================================================
    # 12. LOCALIZAR TEMA
    # ========================================================

            if tema_banco_normalizado != tema_normalizado:

                continue


            resultado["existe"] = True


    # ========================================================
    # 13. CONTEÚDO
    # ========================================================

            resultado["conteudo"] = item.get(
                "conteudo_completo",
                ""
            )


    # ========================================================
    # 14. BRUTO
    # ========================================================

            resultado["bruto"] = item.get(
                "bruto",
                []
            )


    # ========================================================
    # 15. MAPA MEAD
    # ========================================================

            mapa_original = item.get(
                "mapa_mead",
                ""
            )


            print()
            print("==============================")
            print("MAPA BRUTO ENCONTRADO")
            print("==============================")

            print(
                "TIPO:",
                type(mapa_original)
            )


            mapa_texto = extrair_mapa_texto(
                mapa_original
            )


            print()
            print("==============================")
            print("MAPA MEAD EXTRAÍDO")
            print("==============================")

            print(
                mapa_texto[:1000]
            )


    # ========================================================
    # 16. VALIDAR MAPA
    # ========================================================

            if mapa_texto:

                mapa_valido = validar_mapa_do_tema(
                    mapa_texto
                )


                if mapa_valido:

                    resultado["mapa_mead"] = mapa_texto

                else:

                    resultado["mapa_mead"] = ""


            else:

                print()
                print("==============================")
                print("MAPA MEAD NÃO ENCONTRADO")
                print("==============================")


                resultado["mapa_mead"] = ""


    # ========================================================
    # 17. DEBUG FINAL
    # ========================================================

            print()
            print("==============================")
            print("MAPA FINAL RETORNADO PELO BANCO")
            print("==============================")

            print(
                resultado["mapa_mead"][:1000]
            )

            print(
                "TAMANHO MAPA:",
                len(
                    resultado["mapa_mead"]
                )
            )


            break


    # ========================================================
    # 18. BANCO EM FORMATO DICIONÁRIO
    # ========================================================

    elif isinstance(
        banco,
        dict
    ):


        dados = None

        tema_encontrado = None


        # =================================
        # PRIMEIRA TENTATIVA:
        # CHAVE EXATA
        # =================================

        if tema in banco:

            dados = banco.get(
                tema
            )

            tema_encontrado = tema


        # =================================
        # SEGUNDA TENTATIVA:
        # CHAVE NORMALIZADA
        # =================================

        else:

            for chave, valor in banco.items():

                if not isinstance(
                    chave,
                    str
                ):

                    continue


                chave_normalizada = normalizar_tema_chave(
                    chave
                )


                if chave_normalizada == tema_normalizado:

                    dados = valor

                    tema_encontrado = chave

                    break


    # ========================================================
    # 19. TEMA NÃO ENCONTRADO
    # ========================================================

        if not isinstance(
            dados,
            dict
        ):

            return resultado


        resultado["existe"] = True


    # ========================================================
    # 20. CONTEÚDO
    # ========================================================

        resultado["conteudo"] = dados.get(
            "conteudo_completo",
            ""
        )


        # =================================
        # COMPATIBILIDADE COM ESTRUTURAS
        # ANTIGAS DE CATEGORIAS
        # =================================

        if not resultado["conteudo"]:

            categorias = dados.get(
                "categorias",
                {}
            )


            if isinstance(
                categorias,
                dict
            ):

                resultado["conteudo"] = categorias.get(
                    "conteudo_completo",
                    ""
                )


    # ========================================================
    # 21. BRUTO
    # ========================================================

        resultado["bruto"] = dados.get(
            "bruto",
            []
        )


    # ========================================================
    # 22. MAPA MEAD
    # ========================================================

        mapa_original = dados.get(
            "mapa_mead",
            ""
        )


        print()
        print("==============================")
        print("MAPA BRUTO ENCONTRADO")
        print("==============================")

        print(
            "TEMA ENCONTRADO:",
            tema_encontrado
        )

        print(
            "TIPO:",
            type(mapa_original)
        )


        mapa_texto = extrair_mapa_texto(
            mapa_original
        )


        print()
        print("==============================")
        print("MAPA MEAD EXTRAÍDO")
        print("==============================")

        print(
            mapa_texto[:1000]
        )


    # ========================================================
    # 23. VALIDAR MAPA DO PRÓPRIO TEMA
    # ========================================================

        if mapa_texto:

            mapa_valido = validar_mapa_do_tema(
                mapa_texto
            )


            if mapa_valido:

                resultado["mapa_mead"] = mapa_texto

            else:

                resultado["mapa_mead"] = ""


        else:

            print()
            print("==============================")
            print("MAPA MEAD NÃO ENCONTRADO")
            print("==============================")


            resultado["mapa_mead"] = ""


    # ========================================================
    # 24. FORMATO DE BANCO DESCONHECIDO
    # ========================================================

    else:

        print()
        print("==============================")
        print("FORMATO DE BANCO NÃO SUPORTADO")
        print("==============================")

        print(
            "TIPO:",
            type(banco)
        )


    # ========================================================
    # 25. RESULTADO FINAL
    # ========================================================

    print()
    print("==============================")
    print("DADOS FINAIS DO BANCO")
    print("==============================")

    print(
        "TEMA:",
        tema
    )

    print(
        "EXISTE:",
        resultado["existe"]
    )

    print(
        "MAPA MEAD:",
        bool(
            resultado["mapa_mead"]
        )
    )

    print(
        "CONTEÚDO:",
        bool(
            resultado["conteudo"]
        )
    )

    print(
        "BRUTO:",
        len(
            resultado["bruto"]
        )
        if isinstance(
            resultado["bruto"],
            list
        )
        else type(
            resultado["bruto"]
        )
    )


    return resultado

    # ========================================================
    # 26. GRUPOS PRINCIPAIS
    # ========================================================

GRUPOS_PRINCIPAIS = {


    "limpeza_industrial":[

        "lavadora",
        "lavadora industrial",
        "alta pressão",
        "hidrojateadora",
        "limpeza industrial",
        "jato de água"

    ],


    "componentes_mecanicos":[

        "terminal rotular",
        "terminal de rótula",
        "rótula",
        "rotula",
        "rótulas esféricas",
        "rolamento esférico",
        "mancal",
        "bucha esférica",
        "spherical plain bearing",
        "rod end"

    ],


    "construcao":[

        "argamassa",
        "revestimento",
        "selagem",
        "selagem corta fogo",
        "corta fogo",
        "firestop",
        "passive fire protection"

    ],


    "eletrica":[

        "painel elétrico",
        "quadro elétrico",
        "cabo",
        "disjuntor",
        "automação"

    ],


    "hidraulica":[

        "válvula",
        "valvula",
        "bomba",
        "hidráulica",
        "hidraulico",
        "tubulação"

    ]

}



    # ========================================================
    # 27. SINAIS TÉCNICOS GERAIS
    # ========================================================

SINAIS_UNIVERSAIS = [

    "modelo",
    "tipo",
    "aplicação",
    "aplicacao",
    "instalação",
    "instalacao",
    "manutenção",
    "manutencao",
    "norma",
    "certificação",
    "certificacao",
    "especificação",
    "especificacao",
    "dimensão",
    "dimensao",
    "material",
    "aço",
    "aco",
    "inox",
    "temperatura",
    "pressão",
    "pressao",
    "capacidade",
    "resistência",
    "resistencia"

]




    # ========================================================
    # 28. INDICADORES TÉCNICOS
    # ========================================================

INDICADORES_TECNICOS = [

    "medida",
    "mm",
    "cm",
    "kg",
    "mpa",
    "bar",
    "°c",
    "norma",
    "componente",
    "componentes",
    "processo",
    "industrial"

]



# ============================================================
# GRUPO PRINCIPAL
# ============================================================

def pertence_ao_grupo_principal(
    tema,
    relacionado
):

    if not tema or not relacionado:

        return False


    tema_lower = tema.lower()

    relacionado_lower = relacionado.lower()


    grupo_tema = None


    # ========================================================
    # 01. IDENTIFICAR GRUPO DO TEMA
    # ========================================================

    for grupo, palavras in GRUPOS_PRINCIPAIS.items():

        for palavra in palavras:

            if palavra in tema_lower:

                grupo_tema = grupo

                break


        if grupo_tema:

            break


    # Tema sem grupo conhecido

    if not grupo_tema:

        return False


    # ========================================================
    # 02. RELAÇÃO DIRETA
    # ========================================================

    for palavra in GRUPOS_PRINCIPAIS[grupo_tema]:

        if palavra in relacionado_lower:

            return True


    # ========================================================
    # 03. RELAÇÃO TÉCNICA FLEXÍVEL
    # ========================================================

    sinais = 0


    for sinal in SINAIS_UNIVERSAIS:

        if sinal in relacionado_lower:

            sinais += 1


    for indicador in INDICADORES_TECNICOS:

        if indicador in relacionado_lower:

            sinais += 1


    # Se tiver sinais técnicos suficientes,
    # aceita mesmo sem palavra exata do grupo

    if sinais >= 2:

        return True


    return False



# ============================================================
# NORMALIZAR TEXTO SEM ACENTOS
# ============================================================

def normalizar_sem_acento(texto):

    if not texto:
        return ""


    texto = unicodedata.normalize(
        "NFD",
        str(texto)
    )


    texto = "".join(
        caractere
        for caractere in texto
        if unicodedata.category(caractere) != "Mn"
    )


    return texto.upper()
    
    

# ============================================================
# COLETAR SITES INFORMADOS PELO USUÁRIO
# ============================================================

def coletar_sites_consulta():

    resultados = []

    try:

        if "txt_sites" not in globals():
            return resultados

        sites = txt_sites.get(
            "1.0",
            tk.END
        ).strip()

        if not sites:
            return resultados

        print()
        print("==========================================")
        print("SITES INFORMADOS PELO USUÁRIO")
        print("==========================================")

        linhas = [
            linha.strip()
            for linha in sites.splitlines()
            if linha.strip()
        ]

        urls_vistas = set()

        for url in linhas:

            if not url.startswith(
                ("http://", "https://")
            ):

                url = "https://" + url

            url = url.strip()

            if url in urls_vistas:
                continue

            urls_vistas.add(url)

            print()
            print("SITE PARA CONSULTA:")
            print(url)

            resultado = coletar_pagina(
                url
            )

            if not resultado:
                print(
                    "SITE NÃO COLETADO:",
                    url
                )
                continue

            texto = str(
                resultado.get(
                    "texto",
                    ""
                )
            ).strip()

            if not texto:
                continue

            resultados.append(
                resultado
            )

            print(
                "SITE COLETADO:",
                url
            )

            print(
                "TIPO:",
                resultado.get(
                    "tipo",
                    ""
                )
            )

            print(
                "CARACTERES:",
                len(texto)
            )

        print()
        print(
            "TOTAL DE SITES UTILIZADOS:",
            len(resultados)
        )

        print("==========================================")

        return resultados

    except Exception as erro:

        print()
        print(
            "ERRO AO COLETAR SITES INFORMADOS:"
        )
        print(erro)

        return resultados


# ============================================================
# ANALISAR SITE INFORMADO COMO REFERÊNCIA PRINCIPAL
# ============================================================

def analisar_site_referencia(
    tema,
    sites_consulta
):

    resultado = {
        "disponivel": False,
        "urls": [],
        "conteudo": "",
        "contexto": ""
    }

    try:

        if not sites_consulta:
            return resultado

        textos_site = []

        for item in sites_consulta:

            if not isinstance(
                item,
                dict
            ):
                continue

            texto = str(
                item.get(
                    "texto",
                    ""
                )
            ).strip()

            url = str(
                item.get(
                    "url",
                    ""
                )
            ).strip()

            if not texto:
                continue

            if url:
                resultado["urls"].append(
                    url
                )

            textos_site.append(
                texto
            )

        if not textos_site:
            return resultado

        conteudo_site = "\n\n".join(
            textos_site
        )

        resultado["conteudo"] = (
            conteudo_site
        )

        resultado["disponivel"] = True

        contexto_maximo = 12000

        resultado["contexto"] = (
            conteudo_site[
                :contexto_maximo
            ]
        )

        print()
        print("==========================================")
        print("ANÁLISE DO SITE COMO REFERÊNCIA PRINCIPAL")
        print("==========================================")

        print(
            "TEMA:",
            tema
        )

        print(
            "SITES:",
            len(resultado["urls"])
        )

        print(
            "CARACTERES DO SITE:",
            len(conteudo_site)
        )

        print(
            "CONTEXTO DE REFERÊNCIA:",
            len(resultado["contexto"])
        )

        print("==========================================")

        return resultado

    except Exception as erro:

        print()
        print(
            "ERRO AO ANALISAR SITE DE REFERÊNCIA:"
        )
        print(erro)

        return resultado    

# ============================================================
# GERAR MATERIAL INTERFACE
# ============================================================

def gerar_material_interface():

    global ARQUIVO_BANCO
    global IA_PROCESSANDO
    global PROCESSAMENTO_ATIVO

    try:

        PROCESSAMENTO_ATIVO = True
        PAGINAS_EM_PROCESSAMENTO.clear()
        IA_PROCESSANDO = True

        progresso = carregar_progresso()

        inicio_total = time.time()

        threading.Thread(
            target=atualizar_tempo_execucao,
            args=(inicio_total, "Processamento geral"),
            daemon=True
        ).start()

        if progresso:

            print()
            print("==============================")
            print("RETOMANDO PROCESSAMENTO")
            print("==============================")

            print(progresso)

        atualizar_progresso(
            0,
            "Iniciando..."
        )

        ARQUIVO_BANCO = entrada_banco.get().strip()

        if not ARQUIVO_BANCO:

            ARQUIVO_BANCO = (
                r"C:\Python\gerador-conteudo\banco\conteudo-site.json"
            )

            print()
            print("==============================")
            print("BANCO PADRÃO UTILIZADO")
            print("==============================")

            print(ARQUIVO_BANCO)

            print()

            if os.path.exists(ARQUIVO_BANCO):
                print("BANCO EXISTE")
            else:
                print("BANCO NÃO EXISTE")

            print()

        palavras = txt_palavras.get(
            "1.0",
            tk.END
        ).strip()

        if not palavras:

            messagebox.showerror(
                "Erro",
                "Informe ao menos uma palavra-chave."
            )

            return

        lista_palavras = [
            x.strip()
            for x in palavras.splitlines()
            if x.strip()
        ]

        novo_conjunto = verificar_novo_conjunto(
            lista_palavras
        )

        if novo_conjunto:
            progresso = None

        # ========================================================
        # 01. RETOMAR DE ONDE PAROU
        # ========================================================

        inicio_lista = 0

        if progresso:

            inicio_lista = progresso.get(
                "indice_tema",
                0
            )

            if inicio_lista >= len(lista_palavras):

                print()
                print("==============================")
                print("TODOS OS TEMAS JÁ PROCESSADOS")
                print("==============================")

                atualizar_progresso(
                    100,
                    "Todos os temas já foram processados."
                )

                return

        # ========================================================
        # 02. PROCESSAR TEMAS
        # ========================================================

        for indice_tema, tema in enumerate(
            lista_palavras[inicio_lista:],
            start=inicio_lista
        ):

            inicio_tema = time.time()

            print()
            print("=" * 50)
            print(
                "PROCESSANDO:",
                tema
            )
            print("=" * 50)

            # ====================================================
            # 03. INICIALIZAR DADOS DO TEMA
            # ====================================================

            textos = []
            paginas = []
            urls = []
            dados_coleta = []
            paginas_aprovadas = 0

            # ====================================================
            # 04. IDENTIFICAR GRUPO DO TEMA
            # ====================================================

            print()
            print("==============================")
            print("IDENTIFICANDO GRUPO DO TEMA")
            print("==============================")

            print(
                "TEMA:",
                tema
            )

            grupo = identificar_grupo_tema(
                tema
            )

            print()
            print("==============================")
            print("GRUPO IDENTIFICADO")
            print("==============================")

            print(grupo)

            # ====================================================
            # 05. GERAR ENTENDIMENTO DO PRODUTO
            # ====================================================

            print()
            print("==============================")
            print("GERANDO ENTENDIMENTO DO PRODUTO")
            print("==============================")

            print(
                "TEMA:",
                tema
            )

            print(
                "GRUPO:",
                grupo
            )

            entendimento = gerar_entendimento_produto(
                tema,
                grupo
            )

            # ====================================================
            # 06. VALIDAR ENTENDIMENTO
            # ====================================================

            print()
            print("==============================")
            print("ENTENDIMENTO DO PRODUTO")
            print("==============================")

            if entendimento:

                print(entendimento)

                print()
                print(
                    "ENTENDIMENTO GERADO COM SUCESSO"
                )

                print(
                    "CARACTERES:",
                    len(
                        str(
                            entendimento
                        )
                    )
                )

            else:

                print(
                    "ENTENDIMENTO NÃO GERADO"
                )

                print(
                    "O FLUXO CONTINUARÁ SEM ENTENDIMENTO."
                )

            # ====================================================
            # 07. INICIO COLETA
            # ====================================================

            print()
            print("==============================")
            print("INICIO COLETA")
            print("==============================")

            print(
                "TEMA:",
                tema
            )

            # ====================================================
            # TRY INTERNO DO TEMA
            # ====================================================

            try:

                # =================================================
                # 08. OBTER DADOS EXISTENTES DO TEMA
                # =================================================

                dados_existentes = obter_dados_banco(
                    tema
                )

                # =================================================
                # 09. INICIALIZAR MAPA MEAD
                # =================================================

                mapa_mead = None

                print()
                print("==============================")
                print("VERIFICANDO MAPA MEAD")
                print("==============================")

                # =================================================
                # 10. VERIFICAR SE EXISTE MAPA NO BANCO
                # =================================================

                if dados_existentes.get("mapa_mead"):

                    print()
                    print("==============================")
                    print("MAPA MEAD ENCONTRADO NO BANCO")
                    print("==============================")

                    mapa_mead = dados_existentes.get(
                        "mapa_mead"
                    )

                    # =============================================
                    # 11. NORMALIZAR MAPA RECEBIDO
                    # =============================================

                    if isinstance(
                        mapa_mead,
                        dict
                    ):

                        mapa_mead = mapa_mead.get(
                            "texto",
                            ""
                        )

                    elif mapa_mead is None:

                        mapa_mead = ""

                    else:

                        mapa_mead = str(
                            mapa_mead
                        )

                    mapa_mead = mapa_mead.strip()

                    print()
                    print("MAPA MEAD CARREGADO:")

                    print(
                        mapa_mead[:500]
                    )

                else:

                    # =============================================
                    # 12. MAPA NÃO EXISTE
                    # =============================================

                    print()
                    print("==============================")
                    print("MAPA MEAD NÃO ENCONTRADO")
                    print("==============================")

                    mapa_mead = None

                # =================================================
                # 13. NORMALIZAR SEM ACENTO
                # =================================================

                def normalizar_sem_acento(texto):

                    if not texto:
                        return ""

                    texto = unicodedata.normalize(
                        "NFD",
                        str(texto)
                    )

                    texto = "".join(
                        caractere
                        for caractere in texto
                        if unicodedata.category(
                            caractere
                        ) != "Mn"
                    )

                    return texto.upper()

                # =================================================
                # 14. VALIDAR MAPA EXISTENTE
                # =================================================

                mapa_validado = False

                if mapa_mead:

                    mapa_maiusculo = normalizar_sem_acento(
                        mapa_mead
                    )

                    mapa_validado = True

                # =================================================
                # 15. ESTADO FINAL DO MAPA
                # =================================================

                print()
                print("==============================")
                print("ESTADO FINAL DO MAPA MEAD")
                print("==============================")

                print(
                    "MAPA EXISTE:",
                    bool(mapa_mead)
                )

                print(
                    "TAMANHO:",
                    len(mapa_mead)
                    if mapa_mead
                    else 0
                )

                if mapa_mead:

                    print()
                    print("MAPA MEAD FINAL:")

                    print(
                        mapa_mead[:500]
                    )

                else:

                    print()
                    print(
                        "NENHUM MAPA MEAD DISPONÍVEL"
                    )

                # =================================================
                # 16. GARANTIR TEXTOS
                # =================================================

                if not textos:

                    print()
                    print("==============================")
                    print("CARREGANDO DADOS BRUTOS PARA MAPA MEAD")
                    print("==============================")

                    textos = carregar_bruto(
                        tema
                    )

                    print(
                        "TEXTOS CARREGADOS:",
                        len(textos)
                    )

                # =================================================
                # 17. ADICIONAR SITES INFORMADOS PELO USUÁRIO
                # =================================================

                print()
                print("==============================")
                print("TEXTOS ANTES DA ANÁLISE")
                print("==============================")

                print(
                    "TOTAL:",
                    len(textos)
                )

                sites_consulta = coletar_sites_consulta()

                analise_site = analisar_site_referencia(
                    tema,
                    sites_consulta
                )

                if sites_consulta:

                    textos.extend(
                        sites_consulta
                    )

                    print()
                    print("==============================")
                    print("SITES ADICIONADOS AO PATRIMÔNIO")
                    print("==============================")

                    print(
                        "SITES:",
                        len(sites_consulta)
                    )

                    print(
                        "TOTAL DE TEXTOS:",
                        len(textos)
                    )

                # =================================================
                # 18. NORMALIZAR TEXTOS PARA MEAD
                # =================================================

                textos_normalizados = []

                for item in textos:

                    if isinstance(
                        item,
                        dict
                    ):

                        texto = item.get(
                            "texto",
                            ""
                        )

                        if texto and texto.strip():

                            textos_normalizados.append({
                                "url": item.get(
                                    "url",
                                    ""
                                ),
                                "tipo": item.get(
                                    "tipo",
                                    "texto"
                                ),
                                "texto": texto
                            })

                    elif isinstance(
                        item,
                        str
                    ):

                        if item.strip():

                            textos_normalizados.append({
                                "url": "",
                                "tipo": "texto",
                                "texto": item
                            })

                print()
                print("==============================")
                print("TEXTOS NORMALIZADOS PARA MEAD")
                print("==============================")

                print(
                    "QUANTIDADE:",
                    len(textos_normalizados)
                )

                # =================================================
                # 19. LIMPAR REFERÊNCIAS
                # =================================================

                textos_para_mapa = limpar_lista_referencias(
                    textos_normalizados,
                    tema
                )

                print()
                print("==============================")
                print("TEXTOS ENVIADOS AO MAPA MEAD")
                print("==============================")

                print(
                    "QUANTIDADE:",
                    len(textos_para_mapa)
                )

                # =================================================
                # 20. GARANTIR FORMATO FINAL
                # =================================================

                textos_para_mapa_corrigidos = []

                for item in textos_para_mapa:

                    if isinstance(
                        item,
                        dict
                    ):

                        texto = item.get(
                            "texto",
                            ""
                        )

                        if texto and texto.strip():

                            textos_para_mapa_corrigidos.append({
                                "url": item.get(
                                    "url",
                                    ""
                                ),
                                "tipo": item.get(
                                    "tipo",
                                    "texto"
                                ),
                                "texto": texto
                            })

                    elif isinstance(
                        item,
                        str
                    ):

                        if item.strip():

                            textos_para_mapa_corrigidos.append({
                                "url": "",
                                "tipo": "texto",
                                "texto": item
                            })

                # =================================================
                # 21. LIMPAR MAPA MEAD
                # =================================================

                def limpar_mapa_mead(texto):

                    if not texto:
                        return ""

                    substituicoes = {

                        "ATIVOS:":
                        "ATIVOS_NARRATIVOS:",

                        "EXPERIÊNCIA DO FABRICANTE":
                        "experiência técnica",

                        "experiência do fabricante":
                        "experiência técnica",

                        "garantia de alta resistência":
                        "desempenho técnico",

                        "cumprimento das normas":
                        "conformidade técnica",

                        "fabricante":
                        "especialista do segmento",

                        "Fabricante":
                        "especialista do segmento",
                    }

                    for antigo, novo in substituicoes.items():

                        texto = texto.replace(
                            antigo,
                            novo
                        )

                    remover = [

                        "IT 09",
                        "IT-09",
                        "CBPMESP",
                        "NBR",
                        "ISO",
                        "ABNT",
                        "certificação",
                        "certificado",
                        "marca",
                        "modelo",
                        "linha",
                        "série",
                    ]

                    for palavra in remover:

                        texto = texto.replace(
                            palavra,
                            ""
                        )

                    return texto.strip()

                if mapa_mead:

                    mapa_mead = limpar_mapa_mead(
                        mapa_mead
                    )

                # =================================================
                # 22. VERIFICAR CONTEÚDO RELACIONADO
                # =================================================

                textos_encontrados = False

                # =================================================
                # 23. BUSCAR TEMAS RELACIONADOS
                # =================================================

                if not dados_existentes.get("conteudo"):

                    relacionados = buscar_temas_relacionados(
                        tema
                    )

                    relacionados_filtrados = []

                    if relacionados:

                        for item in relacionados:

                            if pertence_ao_grupo_principal(
                                tema,
                                item
                            ):

                                relacionados_filtrados.append(
                                    item
                                )

                    relacionados = relacionados_filtrados

                    if relacionados:

                        relacionados = relacionados[:5]

                        print()
                        print("==============================")
                        print("TEMAS RELACIONADOS COMPATÍVEIS")
                        print("==============================")

                        for item in relacionados:

                            print(item)

                            dados_relacionado = obter_dados_banco(
                                item
                            )

                            conteudo_relacionado = False

                            # NÃO USAR MAPAS MEAD DE OUTROS TEMAS

                            if dados_relacionado.get(
                                "conteudo"
                            ):

                                textos.append(
                                    dados_relacionado["conteudo"]
                                )

                                conteudo_relacionado = True

                            if dados_relacionado.get(
                                "bruto"
                            ):

                                textos.extend(
                                    dados_relacionado["bruto"]
                                )

                                conteudo_relacionado = True

                            if conteudo_relacionado:

                                textos_encontrados = True

                            print()
                            print("==============================")
                            print("BASE RELACIONADA ANALISADA")
                            print("==============================")

                            print(
                                "TEXTOS ACUMULADOS:",
                                len(textos)
                            )

                            print(
                                "CONTEÚDO RELACIONADO VÁLIDO:",
                                conteudo_relacionado
                            )

                    else:

                        print()
                        print("==============================")
                        print("NENHUM TEMA RELACIONADO COMPATÍVEL")
                        print("==============================")

                # =================================================
                # 24. COMPLETAR COLETA QUANDO NECESSÁRIO
                # =================================================

                if not textos_encontrados:

                    textos_brutos = carregar_bruto(
                        tema
                    )

                    if textos_brutos:

                        print(
                            "USANDO DADOS BRUTOS EXISTENTES"
                        )

                        # Não apagar os sites já coletados.
                        textos = (
                            list(textos_brutos)
                            + list(sites_consulta)
                        )

                    else:

                        print(
                            "DADOS BRUTOS NÃO ENCONTRADOS - PESQUISANDO"
                        )

                        urls = pesquisar_completo(
                            tema
                        )

                        for url in urls:

                            print()
                            print("COLETANDO:")
                            print(url)

                            dados = coletar_pagina(
                                url
                            )

                            texto = ""

                            if dados:

                                texto = dados.get(
                                    "texto",
                                    ""
                                )

                            # =====================================
                            # 25. GARANTIR TIPO DA FONTE
                            # =====================================

                            tipo = "html"

                            if isinstance(
                                dados,
                                dict
                            ):

                                tipo = dados.get(
                                    "tipo",
                                    "html"
                                )

                            # =====================================
                            # 26. VALIDAR CONTEÚDO
                            # =====================================

                            if len(texto) > 500:

                                relevante = (
                                    validar_conteudo_relevante(
                                        texto,
                                        tema
                                    )
                                )

                                tecnico = (
                                    validar_conteudo_tecnico(
                                        texto,
                                        tema
                                    )
                                )

                                if relevante and tecnico:

                                    textos.append(
                                        {
                                            "url": url,
                                            "tipo": tipo,
                                            "texto": texto
                                        }
                                    )

                                    paginas_aprovadas += 1

                                    dados_coleta.append({
                                        "url": url,
                                        "tipo": tipo,
                                        "status": "aprovado",
                                        "motivo": "",
                                        "texto": texto
                                    })

                                else:

                                    dados_coleta.append({
                                        "url": url,
                                        "tipo": tipo,
                                        "status": "descartado",
                                        "motivo": "conteúdo não aprovado",
                                        "texto": texto
                                    })

                                    print(
                                        "CONTEUDO MANTIDO NO JSON COMO DESCARTADO"
                                    )

                        # =========================================
                        # 27. SALVAR DADOS BRUTOS
                        # =========================================

                        print()
                        print("==============================")
                        print("SALVANDO DADOS BRUTOS")
                        print("==============================")

                        print(
                            "TEMA:",
                            tema
                        )

                        print(
                            "PÁGINAS PARA SALVAR:",
                            len(dados_coleta)
                        )

                        if dados_coleta:

                            salvar_bruto(
                                tema,
                                dados_coleta
                            )

                        else:

                            print()
                            print("==============================")
                            print("DADOS COLETA VAZIOS")
                            print("BRUTO EXISTENTE SERÁ PRESERVADO")
                            print("==============================")

                        # =========================================
                        # 28. RECARREGAR BRUTO APÓS SALVAR
                        # =========================================

                        dados_coleta = carregar_bruto(
                            tema
                        )

                        textos = []

                        for item in dados_coleta:

                            if not isinstance(
                                item,
                                dict
                            ):
                                continue

                            texto = item.get(
                                "texto",
                                ""
                            )

                            if texto and texto.strip():

                                textos.append({
                                    "url": item.get(
                                        "url",
                                        ""
                                    ),
                                    "tipo": item.get(
                                        "tipo",
                                        ""
                                    ),
                                    "texto": texto
                                })

                        # Preservar os sites informados pelo usuário.

                        if sites_consulta:

                            textos.extend(
                                sites_consulta
                            )

                        print(
                            "TEXTOS PARA MAPA MEAD:",
                            len(textos)
                        )

                # =================================================
                # 29. NORMALIZAR TEXTOS PARA MAPA MEAD
                # =================================================

                textos_para_mapa = limpar_lista_referencias(
                    textos,
                    tema
                )

                print(
                    "TEXTOS APÓS LIMPEZA:",
                    len(textos_para_mapa)
                )

                # =================================================
                # 30. GARANTIR FORMATO FINAL
                # =================================================

                textos_para_mapa_corrigidos = []

                for item in textos_para_mapa:

                    if isinstance(
                        item,
                        str
                    ):

                        textos_para_mapa_corrigidos.append({
                            "url": "",
                            "tipo": "texto",
                            "texto": item
                        })

                    elif isinstance(
                        item,
                        dict
                    ):

                        texto = item.get(
                            "texto",
                            ""
                        )

                        if texto and texto.strip():

                            textos_para_mapa_corrigidos.append({
                                "url": item.get(
                                    "url",
                                    ""
                                ),
                                "tipo": item.get(
                                    "tipo",
                                    "texto"
                                ),
                                "texto": texto
                            })

                print(
                    "TEXTOS FINAIS PARA MAPA MEAD:",
                    len(textos_para_mapa_corrigidos)
                )

                # =================================================
                # 31. GERAR MAPA MEAD SE NECESSÁRIO
                # =================================================

                if not mapa_mead:

                    mapa_mead = gerar_e_salvar_mapa_mead(
                        tema,
                        textos_para_mapa_corrigidos
                    )

                if mapa_mead:

                    mapa_mead = limpar_mapa_mead(
                        mapa_mead
                    )

                if not mapa_mead:

                    print()
                    print("==============================")
                    print("MAPA MEAD NÃO GERADO")
                    print("==============================")

                    continue

                print()
                print("==============================")
                print("COLETA E MAPA MEAD FINALIZADOS")
                print("==============================")

                salvar_progresso(
                    tema,
                    indice_tema + 1,
                    "mapa_mead",
                    0,
                    "concluido",
                    lista_palavras
                )

                # =================================================
                # 32. VERIFICAR DADOS PARA IA
                # =================================================

                if not textos:

                    print()
                    print("==============================")
                    print("SEM DADOS PARA GERAR CONTEÚDO")
                    print("==============================")

                    salvar_progresso(
                        tema,
                        indice_tema,
                        "erro_coleta",
                        0,
                        "sem dados brutos",
                        lista_palavras
                    )

                    continue

                # =================================================
                # 33. BLOQUEIO MAPA MEAD OBRIGATÓRIO
                # =================================================

                if (
                    not mapa_mead
                    or not str(mapa_mead).strip()
                    or str(mapa_mead).strip() in [
                        "{}",
                        "[]"
                    ]
                ):

                    print()
                    print("==============================")
                    print("MAPA MEAD AUSENTE OU INVÁLIDO")
                    print("==============================")

                    print(
                        "NÃO É POSSÍVEL GERAR CONTEÚDO"
                    )

                    salvar_progresso(
                        tema,
                        indice_tema,
                        "erro_mapa_mead",
                        0,
                        "mapa mead ausente",
                        lista_palavras
                    )

                    continue

                # =================================================
                # 34. INICIAR GERAÇÃO DE CONTEÚDO
                # =================================================

                salvar_progresso(
                    tema,
                    indice_tema,
                    "conteudo_completo",
                    1,
                    "processando",
                    lista_palavras
                )

                print()
                print("==============================")
                print("CRIANDO CONTEÚDO COMPLETO")
                print("==============================")

                inicio_conteudo = time.time()

                # =================================================
                # 35. INICIAR MONITOR DE IA
                # =================================================

                IA_PROCESSANDO = True

                atualizar_progresso(
                    50,
                    f"Tema: {tema} | Gerando conteúdo com IA..."
                )

                total_textos = len(textos)

                total_caracteres = sum(
                    len(
                        item.get(
                            "texto",
                            ""
                        )
                        if isinstance(
                            item,
                            dict
                        )
                        else str(item)
                    )
                    for item in textos
                )

                media_caracteres = (
                    total_caracteres / total_textos
                    if total_textos
                    else 0
                )

                print()
                print("==============================")
                print("DADOS ENVIADOS PARA IA")
                print("==============================")

                print(
                    "TEMA:",
                    tema
                )

                print(
                    "TEXTOS:",
                    total_textos
                )

                print(
                    "CARACTERES:",
                    total_caracteres
                )

                print(
                    "MÉDIA POR TEXTO:",
                    int(media_caracteres)
                )

                if textos:

                    primeiro_texto = (

                        textos[0].get(
                            "texto",
                            ""
                        )

                        if isinstance(
                            textos[0],
                            dict
                        )

                        else str(
                            textos[0]
                        )
                    )

                    print(
                        "PRIMEIRO TEXTO:",
                        len(primeiro_texto),
                        "caracteres"
                    )

                print()
                print("==============================")
                print("MAPA MEAD ENVIADO AO CONTEÚDO")
                print("==============================")

                if mapa_mead:

                    print(
                        str(mapa_mead)[:1000]
                    )

                else:

                    print("MAPA VAZIO")

                print()
                print("==============================")
                print("CHAMANDO GERADOR DE CONTEÚDO")
                print("==============================")

                inicio_ia = time.time()

                print()
                print("==============================")
                print("INICIANDO IA - CONTEÚDO")
                print("==============================")

                print(
                    "TEMA:",
                    tema
                )

                print(
                    "TEXTOS:",
                    len(textos)
                )

                print(
                    "MAPA:",
                    len(str(mapa_mead))
                )

                print(
                    "CARACTERES DOS TEXTOS:",
                    sum(
                        len(
                            item.get(
                                "texto",
                                ""
                            )
                            if isinstance(
                                item,
                                dict
                            )
                            else str(item)
                        )
                        for item in textos
                    )
                )

                print(
                    "INÍCIO IA:",
                    time.strftime("%H:%M:%S")
                )

                print()
                print(
                    "AGUARDANDO RETORNO DO OLLAMA..."
                )
                print("==============================")

                estrutura_editorial_atual = (
                    obter_estrutura_editorial()
                )

                print()
                print("==============================")
                print("DEBUG CHECKBOXES EDITORIAIS")
                print("==============================")

                for chave, valor in (
                    estrutura_editorial_atual.items()
                ):

                    print(
                        f"{chave}: {valor}"
                    )

                print("==============================")

                print(
                    "TOTAL SELECIONADOS:",
                    sum(
                        1
                        for valor
                        in estrutura_editorial_atual.values()
                        if valor
                    )
                )

                print(
                    "TOTAL BLOCOS:",
                    len(
                        estrutura_editorial_atual
                    )
                )

                print("==============================")

                conteudo_completo = (
                    gerar_conteudo_completo(
                        tema,
                        textos,
                        mapa_mead,
                        estrutura_editorial_atual,
                        dados_coleta=dados_coleta
                    )
                )

                IA_PROCESSANDO = False

                fim_ia = time.time()

                print()
                print("==============================")
                print("OLLAMA RETORNOU")
                print("==============================")

                print(
                    "FIM IA:",
                    time.strftime("%H:%M:%S")
                )

                print(
                    "TEMPO IA:",
                    formatar_tempo(
                        fim_ia - inicio_ia
                    )
                )

                print(
                    "TIPO:",
                    type(conteudo_completo)
                )

                print(
                    "CARACTERES:",
                    len(
                        conteudo_completo or ""
                    )
                )

                tempo_total = (
                    time.time()
                    -
                    inicio_conteudo
                )

                tempo_ia = (
                    fim_ia
                    -
                    inicio_ia
                )

                tamanho_conteudo = len(
                    conteudo_completo or ""
                )

                print()
                print("==============================")
                print("RETORNOU DO GERADOR")
                print("==============================")

                print(
                    "TIPO:",
                    type(conteudo_completo)
                )

                print(
                    "TAMANHO:",
                    tamanho_conteudo
                )

                print()
                print("==============================")
                print("ESTATÍSTICAS DA IA")
                print("==============================")

                print(
                    "TEXTOS ENVIADOS:",
                    total_textos
                )

                print(
                    "CARACTERES ENVIADOS:",
                    total_caracteres
                )

                print(
                    "MÉDIA POR TEXTO:",
                    int(media_caracteres)
                )

                print(
                    "CARACTERES GERADOS:",
                    tamanho_conteudo
                )

                print(
                    "TEMPO IA:",
                    formatar_tempo(
                        tempo_ia
                    )
                )

                print(
                    "TEMPO TOTAL:",
                    formatar_tempo(
                        tempo_total
                    )
                )

                if (
                    tempo_ia > 0
                    and tamanho_conteudo > 0
                ):

                    print(
                        "VELOCIDADE:",
                        int(
                            tamanho_conteudo
                            /
                            tempo_ia
                        ),
                        "caracteres/seg"
                    )

                atualizar_progresso(
                    90,
                    (
                        f"Tema: {tema} | "
                        f"Conteúdo gerado | "
                        f"Tempo: {formatar_tempo(tempo_total)}"
                    )
                )

                # =================================================
                # 36. SALVAR CONTEÚDO
                # =================================================

                if (
                    conteudo_completo
                    and conteudo_completo.strip()
                ):

                    salvar_banco(
                        tema,
                        "conteudo_completo",
                        conteudo_completo
                    )

                    print()
                    print("==============================")
                    print("SALVANDO CONTEÚDO")
                    print("==============================")

                    print(
                        "CHAVE:",
                        tema
                    )

                    print(
                        "TAMANHO:",
                        tamanho_conteudo
                    )

                    print(
                        conteudo_completo[:500]
                    )

                    print()
                    print("==============================")
                    print("CONTEÚDO COMPLETO SALVO")
                    print("==============================")

                    salvar_progresso(
                        tema,
                        indice_tema + 1,
                        "finalizado",
                        0,
                        "concluido",
                        lista_palavras
                    )

                else:

                    IA_PROCESSANDO = False

                    print()
                    print("==============================")
                    print("FALHA AO GERAR CONTEÚDO COMPLETO")
                    print("==============================")

                    print(
                        "TEMA:",
                        tema
                    )

                    print(
                        "RETORNO:",
                        repr(conteudo_completo)
                    )

                    salvar_progresso(
                        tema,
                        indice_tema,
                        "mapa_mead",
                        0,
                        "concluido",
                        lista_palavras
                    )

            except Exception as e:

                IA_PROCESSANDO = False

                print()
                print("==============================")
                print("ERRO NO TEMA")
                print("==============================")

                print(
                    repr(e)
                )

                atualizar_progresso(
                    0,
                    f"Erro no tema {tema}: {e}"
                )

        # ========================================================
        # 37. FINAL DO PROCESSAMENTO GERAL
        # ========================================================

        PROCESSAMENTO_ATIVO = False
        IA_PROCESSANDO = False

        tempo_final = (
            time.time()
            -
            inicio_total
        )

        print()
        print("==============================")
        print("PROCESSAMENTO FINALIZADO")
        print("==============================")

        print(
            "TEMPO TOTAL:",
            formatar_tempo(
                tempo_final
            )
        )

        atualizar_progresso(
            100,
            (
                f"Concluído | "
                f"Tempo total: "
                f"{formatar_tempo(tempo_final)}"
            )
        )

    except Exception as e:

        PROCESSAMENTO_ATIVO = False
        IA_PROCESSANDO = False

        print()
        print("==============================")
        print("ERRO GERAL NO PROCESSAMENTO")
        print("==============================")

        print(
            repr(e)
        )

        atualizar_progresso(
            0,
            f"Erro geral: {e}"
        )
        
        

# ============================================================
# BOTÃO - APAGAR PROGRESSO MANUAL
# ============================================================

def apagar_progresso_manual():

    try:

        if os.path.exists(ARQUIVO_PROGRESSO):

            os.remove(
                ARQUIVO_PROGRESSO
            )

            messagebox.showinfo(
                "Progresso",
                "Progresso apagado com sucesso."
            )

        else:

            messagebox.showinfo(
                "Progresso",
                "Não existe progresso salvo."
            )


    except Exception as e:

        messagebox.showerror(
            "Erro",
            str(e)
        )
        

# ============================================================
# GERAR DOCX MEAD
# ============================================================

def gerar_docx_mead(
    titulo,
    conteudo,
    pasta_destino
):




    documento = Document()


    # TÍTULO

    documento.add_heading(
        titulo,
        level=1
    )


    # CONTEÚDO

    for linha in conteudo.split("\n"):

        linha = linha.strip()


        if not linha:

            continue


        documento.add_paragraph(
            linha
        )

    
    # ========================================================
    # 01. NORMALIZAR NOME DO ARQUIVO DOCX
    # ========================================================
    
    nome_arquivo = (
        titulo
        .lower()
        .strip()
        .replace(" ", "_")
        .replace("/", "-")
        .replace("\\", "-")
        + ".docx"
    )
    
    
    caminho = os.path.join(
        pasta_destino,
        nome_arquivo
    )
    
    
    documento.save(
        caminho
    )
    
    
    return caminho    




# ============================================================
# INTERFACE
# ============================================================
           

def gerar_docx_interface():



    pasta = entrada_pasta_docx.get().strip()

    if not pasta:

        messagebox.showwarning(
            "DOCX",
            "Selecione a pasta destino dos DOCX."
        )

        return


    caminho_banco = ARQUIVO_BANCO


    if not os.path.exists(caminho_banco):

        messagebox.showerror(
            "DOCX",
            "Banco de conteúdo não encontrado."
        )

        return


    try:

        with open(
            caminho_banco,
            "r",
            encoding="utf-8"
        ) as arquivo:

            banco = json.load(
                arquivo
            )

    except Exception as erro:

        messagebox.showerror(
            "DOCX",
            f"Erro ao abrir banco:\n{erro}"
        )

        return


    if not banco:

        messagebox.showwarning(
            "DOCX",
            "Banco vazio."
        )

        return


    total = 0

    print()
    print("==============================")
    print("GERANDO DOCX")
    print("==============================")
    print("TOTAL DE TEMAS:", len(banco))
    print("PASTA:", pasta)

    for tema, dados in banco.items():

        conteudo = dados.get(
            "conteudo_completo",
            ""
        )
        
        
        if not conteudo:
        
            categorias = dados.get(
                "categorias",
                {}
            )
        
            if isinstance(
                categorias,
                dict
            ):
        
                conteudo = categorias.get(
                    "conteudo_completo",
                    ""
                )

        if not conteudo.strip():

            print()
            print("==============================")
            print("SEM CONTEÚDO")
            print("==============================")
            print(tema)

            continue


        try:

            caminho = gerar_docx_mead(
                tema,
                conteudo,
                pasta
            )
        
            total += 1
        
            print()
            print("==============================")
            print("DOCX GERADO")
            print("==============================")
            print("TEMA:", tema)
            print("ARQUIVO:", caminho)

        except Exception as erro:

            print()
            print("==============================")
            print("ERRO AO GERAR DOCX")
            print("==============================")
            print(tema)
            print(erro)


    messagebox.showinfo(
        "DOCX",
        f"{total} arquivo(s) DOCX gerado(s)."
    )
    
    print()
    print("==============================")
    print("PROCESSO FINALIZADO")
    print("==============================")
    print("DOCX GERADOS:", total)



    # ========================================================
    # 01. INTERFACE PRINCIPAL
    # ========================================================


    # ========================================================
    # 02. ESTRUTURA EDITORIAL
    # ========================================================

ESTRUTURA_EDITORIAL = {
    "apresentacao": True,
    "funcionamento": True,
    "aplicacoes": True,
    "criterios": True,
    "comercial": True,
    "informacao_tecnica": True,
    "instalacao_execucao": True,
    "beneficios": True
}


def obter_estrutura_editorial():

    estrutura = {
        chave: bool(valor)
        for chave, valor
        in ESTRUTURA_EDITORIAL.items()
    }

    print()
    print("==============================")
    print("CHECKBOXES EDITORIAIS")
    print("==============================")

    for chave, valor in estrutura.items():

        print(
            f"{chave}: {valor}"
        )

    print("==============================")
    print("ESTRUTURA EDITORIAL ENVIADA")
    print("==============================")

    print(
        estrutura
    )

    print("==============================")

    return estrutura


def selecionar_todos_blocos_editoriais():

    for variavel in VARIAVEIS_EDITORIAIS.values():

        variavel.set(True)

    atualizar_estrutura_editorial()


def limpar_blocos_editoriais():

    for variavel in VARIAVEIS_EDITORIAIS.values():

        variavel.set(False)

    atualizar_estrutura_editorial()


def atualizar_estrutura_editorial():

    global ESTRUTURA_EDITORIAL

    for chave, variavel in VARIAVEIS_EDITORIAIS.items():

        ESTRUTURA_EDITORIAL[chave] = bool(
            variavel.get()
        )


    # ========================================================
    # 01. NOMES DOS BLOCOS EDITORIAIS
    # ========================================================

NOMES_BLOCOS_EDITORIAIS = {

    "apresentacao":
        "Apresentação e contexto",

    "funcionamento":
        "Funcionamento / Como funciona",

    "aplicacoes":
        "Aplicações / Situações atendidas",

    "criterios":
        "Critérios e diferenciais",

    "comercial":
        "Contexto comercial",

    "informacao_tecnica":
        "Informação técnica",

    "instalacao_execucao":
        "Instalação / Execução / Processo",

    "beneficios":
        "Benefícios"
}


    # ========================================================
    # 02. VARIÁVEIS DOS CHECKBOXES
    # ========================================================

VARIAVEIS_EDITORIAIS = {}



# ============================================================
# INTERFACE PRINCIPAL
# ============================================================

def abrir_interface():

    global janela
    global entrada_site
    global entrada_grupo
    global entrada_banco
    global entrada_modelo
    global txt_sites
    global txt_palavras
    global barra_progresso
    global status_progresso
    global entrada_pasta_docx
    global status_tempo

    janela = tk.Tk()

    janela.title(
        "Gerador SEO MEAD"
    )

    janela.geometry(
        "1200x800"
    )


    # =====================================
    # PAINEL LATERAL
    # ESTRUTURA EDITORIAL
    # =====================================

    frame_editorial = tk.LabelFrame(
        janela,
        text="ESTRUTURA EDITORIAL",
        padx=12,
        pady=12
    )

    frame_editorial.pack(
        side="right",
        fill="y",
        padx=10,
        pady=10
    )


    # ========================================================
    # 01. TÍTULO DO PAINEL
    # ========================================================

    tk.Label(
        frame_editorial,
        text="Selecione os conteúdos\ndesta página:",
        font=("Arial", 10, "bold"),
        justify="left"
    ).pack(
        anchor="w",
        pady=(0, 10)
    )


    # ========================================================
    # 02. CHECKBOXES
    # ========================================================

    for chave, nome in NOMES_BLOCOS_EDITORIAIS.items():

        variavel = tk.BooleanVar(
            value=ESTRUTURA_EDITORIAL.get(
                chave,
                True
            )
        )

        VARIAVEIS_EDITORIAIS[chave] = variavel

        tk.Checkbutton(
            frame_editorial,
            text=nome,
            variable=variavel,
            anchor="w",
            justify="left",
            command=atualizar_estrutura_editorial
        ).pack(
            anchor="w",
            fill="x",
            pady=2
        )


    # ========================================================
    # 03. SEPARADOR
    # ========================================================

    tk.Frame(
        frame_editorial,
        height=2
    ).pack(
        fill="x",
        pady=10
    )


    # ========================================================
    # 04. SELECIONAR TODOS
    # ========================================================

    tk.Button(
        frame_editorial,
        text="SELECIONAR TODOS",
        command=selecionar_todos_blocos_editoriais
    ).pack(
        fill="x",
        pady=3
    )


    # ========================================================
    # 05. LIMPAR SELEÇÃO
    # ========================================================

    tk.Button(
        frame_editorial,
        text="LIMPAR SELEÇÃO",
        command=limpar_blocos_editoriais
    ).pack(
        fill="x",
        pady=3
    )


    # ========================================================
    # 06. NOME DO SITE
    # ========================================================

    tk.Label(
        janela,
        text="Nome do Site"
    ).pack()

    entrada_site = tk.Entry(
        janela,
        width=80
    )

    entrada_site.pack()

    entrada_site.bind(
        "<KeyRelease>",
        atualizar_banco_automatico
    )


    # ========================================================
    # 07. GRUPO
    # ========================================================

    tk.Label(
        janela,
        text="Grupo Principal do Projeto"
    ).pack()

    entrada_grupo = tk.Entry(
        janela,
        width=80
    )

    entrada_grupo.pack()


    # ========================================================
    # 08. BANCO
    # ========================================================

    tk.Label(
        janela,
        text="Arquivo Banco"
    ).pack()

    frame_banco = tk.Frame(
        janela
    )

    frame_banco.pack()

    entrada_banco = tk.Entry(
        frame_banco,
        width=70
    )

    entrada_banco.pack(
        side="left"
    )

    tk.Button(
        frame_banco,
        text="Selecionar",
        command=selecionar_banco
    ).pack(
        side="left"
    )


    # ========================================================
    # 09. MODELO DOCX
    # ========================================================

    tk.Label(
        janela,
        text="Modelo DOCX"
    ).pack()

    frame_modelo = tk.Frame(
        janela
    )

    frame_modelo.pack()

    entrada_modelo = tk.Entry(
        frame_modelo,
        width=70
    )

    entrada_modelo.pack(
        side="left"
    )

    tk.Button(
        frame_modelo,
        text="Selecionar",
        command=selecionar_modelo
    ).pack(
        side="left"
    )


    # ========================================================
    # 10. SITES
    # ========================================================

    tk.Label(
        janela,
        text="Sites para Consulta"
    ).pack()

    txt_sites = tk.Text(
        janela,
        height=8,
        width=100
    )

    txt_sites.pack()


    # ========================================================
    # 11. PALAVRAS
    # ========================================================

    tk.Label(
        janela,
        text="Palavras-chave"
    ).pack()

    txt_palavras = tk.Text(
        janela,
        height=15,
        width=100
    )

    txt_palavras.pack()


    # ========================================================
    # 12. STATUS
    # ========================================================

    status_progresso = tk.Label(
        janela,
        text="Aguardando..."
    )

    status_progresso.pack(
        pady=5
    )

    status_tempo = tk.Label(
        janela,
        text=""
    )

    status_tempo.pack(
        pady=2
    )

    barra_progresso = ttk.Progressbar(
        janela,
        length=700,
        mode="determinate"
    )

    barra_progresso.pack(
        pady=5
    )


    # ========================================================
    # 13. PASTA DOCX
    # ========================================================

    tk.Label(
        janela,
        text="Pasta destino dos DOCX:"
    ).pack()

    entrada_pasta_docx = tk.Entry(
        janela,
        width=60
    )

    entrada_pasta_docx.pack(
        pady=5
    )


    def selecionar_pasta_docx():

        pasta = filedialog.askdirectory()

        if pasta:

            entrada_pasta_docx.delete(
                0,
                tk.END
            )

            entrada_pasta_docx.insert(
                0,
                pasta
            )


    tk.Button(
        janela,
        text="SELECIONAR PASTA DOCX",
        command=selecionar_pasta_docx
    ).pack(
        pady=5
    )


    # ========================================================
    # 14. BOTÕES
    # ========================================================
    
    frame_botoes = tk.Frame(
        janela
    )
    
    frame_botoes.pack(
        pady=10
    )
    
    
    # ========================================================
    # 15. APAGAR PROGRESSO
    # ========================================================
    
    tk.Button(
        frame_botoes,
        text="APAGAR PROGRESSO",
        height=2,
        command=apagar_progresso_manual
    ).pack(
        side="left",
        padx=5
    )
    
    
    # ========================================================
    # 16. GERAR MATERIAL
    # ========================================================
    
    tk.Button(
        frame_botoes,
        text="GERAR MATERIAL",
        height=2,
        command=iniciar_geracao
    ).pack(
        side="left",
        padx=5
    )
    
    
    # ========================================================
    # 17. GERAR DOCX
    # ========================================================
    
    tk.Button(
        frame_botoes,
        text="GERAR DOCX",
        height=2,
        command=gerar_docx_interface
    ).pack(
        side="left",
        padx=5
    )


    janela.mainloop()



# ============================================================
# EXECUÇÃO
# ============================================================

def executar():

    global etapa_atual
    global inicio_geracao
    global tempos_etapas
    global etapas_total

    inicio_geracao = time.time()

    etapa_atual = 0
    tempos_etapas = []
    etapas_total = 4
    PAGINAS_EM_PROCESSAMENTO.clear()

    # ========================================================
    # 01. TEMA
    # ========================================================

    tema = input(
        "Tema: "
    ).strip()

    if not tema:

        print()
        print("==============================")
        print("TEMA NÃO INFORMADO")
        print("==============================")

        return

    # ========================================================
    # 02. IDENTIFICAR GRUPO
    # ========================================================

    grupo = identificar_grupo_tema(
        tema
    )

    print()
    print("==============================")
    print("GRUPO IDENTIFICADO")
    print("==============================")

    print(
        grupo
    )

    # ========================================================
    # 03. GERAR ENTENDIMENTO DO PRODUTO
    # ========================================================

    print()
    print("==============================")
    print("GERANDO ENTENDIMENTO DO PRODUTO")
    print("==============================")

    print(
        "TEMA:",
        tema
    )

    print(
        "GRUPO:",
        grupo
    )

    print()
    print("CHAMANDO:")
    print(
        "gerar_entendimento_produto()"
    )

    inicio_entendimento = time.time()

    entendimento = gerar_entendimento_produto(
        tema,
        grupo
    )

    fim_entendimento = time.time()

    # ========================================================
    # 04. RESULTADO DO ENTENDIMENTO
    # ========================================================

    print()
    print("==============================")
    print("ENTENDIMENTO DO PRODUTO RETORNOU")
    print("==============================")

    print(
        "TIPO:",
        type(entendimento)
    )

    print(
        "CARACTERES:",
        len(
            str(
                entendimento or ""
            )
        )
    )

    print(
        "TEMPO:",
        formatar_tempo(
            fim_entendimento
            -
            inicio_entendimento
        )
    )

    print()

    if entendimento:

        print(
            "ENTENDIMENTO:"
        )

        print(
            entendimento
        )

    else:

        print(
            "ENTENDIMENTO VAZIO"
        )

    print()
    print("==============================")
    print("FIM ENTENDIMENTO DO PRODUTO")
    print("==============================")

    # ========================================================
    # 05. PESQUISA
    # ========================================================

    print()
    print("==============================")
    print("PESQUISANDO")
    print("==============================")

    urls = pesquisar_completo(
        tema
    )

    etapa_atual += 1

    tempo = time.time() - inicio_geracao

    tempos_etapas.append(
        tempo
    )

    media = (
        sum(tempos_etapas)
        /
        len(tempos_etapas)
    )

    restante = media * (
        etapas_total
        -
        etapa_atual
    )

    atualizar_progresso(
        25,
        f"Pesquisa concluída | Restante: {formatar_tempo(restante)}"
    )

    # ========================================================
    # 06. FONTES ENCONTRADAS
    # ========================================================

    print()
    print("==============================")
    print("FONTES ENCONTRADAS")
    print("==============================")

    print(
        len(urls)
    )

    paginas = []

    textos = []

    # ========================================================
    # 07. COLETAR FONTES
    # ========================================================

    for i, url in enumerate(
        urls,
        start=1
    ):

        print()

        print(
            f"{i}/{len(urls)}",
            url
        )

        dados = coletar_pagina(
            url
        )

    # ========================================================
    # 08. NORMALIZAR RETORNO DA COLETA
    # ========================================================

        texto = ""

        tipo = "html"

        if isinstance(
            dados,
            dict
        ):

            texto = dados.get(
                "texto",
                ""
            )

            tipo = dados.get(
                "tipo",
                "html"
            )

        elif isinstance(
            dados,
            str
        ):

            texto = dados

        if not texto:

            texto = ""

        print(
            "CARACTERES:",
            len(texto)
        )

    # ========================================================
    # 09. VALIDAR RELEVÂNCIA DA FONTE
    # ========================================================

        texto_lower = texto.lower()

        palavras_tema = [

            p

            for p in tema.lower().split()

            if len(p) > 3

        ]

        ocorrencias = 0

        for palavra in palavras_tema:

            if palavra in texto_lower:

                ocorrencias += 1

    # ========================================================
    # 10. TERMOS TÉCNICOS POR GRUPO
    # ========================================================

        TERMOS_GRUPOS = {

            "hidraulica": [

                "hidraulica",
                "hidráulica",
                "valvula",
                "válvula",
                "pressao",
                "pressão",
                "atuador",
                "fluido",
                "bomba",
                "vazao",
                "vazão",
                "cilindro"

            ],

            "blindagem": [

                "blindagem",
                "balistico",
                "balístico",
                "protecao",
                "proteção",
                "vidro blindado",
                "nivel iii",
                "nível iii"

            ],

            "construcao": [

                "concreto",
                "argamassa",
                "cimento",
                "estrutura",
                "obra",
                "fundacao",
                "fundação"

            ]

        }

        termos_grupo = TERMOS_GRUPOS.get(
            grupo,
            []
        )

        pontos_grupo = 0

        for termo in termos_grupo:

            if termo in texto_lower:

                pontos_grupo += 1

    # ========================================================
    # 11. APROVAÇÃO FINAL
    # ========================================================

        print()
        print("==============================")
        print("ANTES APROVAÇÃO FINAL")
        print("==============================")

        print(
            "TIPO:",
            tipo
        )

        print(
            "TEXTO:",
            len(texto)
        )

        print(
            "OCORRENCIAS:",
            ocorrencias
        )

        print(
            "TERMOS GRUPO:",
            termos_grupo
        )

        print(
            "PONTOS GRUPO:",
            pontos_grupo
        )

        if (

            len(texto) > 1500

            and ocorrencias >= max(
                2,
                len(palavras_tema) // 2
            )

            and (

                len(termos_grupo) == 0

                or

                pontos_grupo >= 3

            )

        ):

            print()
            print("==============================")
            print("ENTROU NA APROVAÇÃO")
            print("==============================")

            print(
                "URL:",
                url
            )

            print(
                "TIPO:",
                tipo
            )

            print(
                "CARACTERES:",
                len(texto)
            )

            paginas.append({

                "url":
                    url,

                "tipo":
                    tipo,

                "caracteres":
                    len(texto),

                "texto":
                    texto

            })

            textos.append({

                "url":
                    url,

                "tipo":
                    tipo,

                "texto":
                    texto[:30000]

            })

            print()
            print("APROVADO")

            print(
                "CARACTERES:",
                len(texto)
            )

            print(
                "PALAVRAS TEMA:",
                ocorrencias
            )

            print(
                "PONTOS GRUPO:",
                pontos_grupo
            )

        else:

            print()
            print("DESCARTADO")

            print(
                "URL:",
                url
            )

            print(
                "CARACTERES:",
                len(texto)
            )

            print(
                "PALAVRAS TEMA:",
                ocorrencias
            )

            print(
                "PONTOS GRUPO:",
                pontos_grupo
            )

    # ========================================================
    # 12. SALVAR NOVO BRUTO
    # ========================================================

    if paginas:

        print()
        print("==============================")
        print("SALVANDO NOVO BRUTO")
        print("==============================")

        salvar_bruto(
            tema,
            paginas
        )

    else:

        print()
        print("==============================")
        print("BRUTO EXISTENTE PRESERVADO")
        print("==============================")

    # ========================================================
    # 13. ATUALIZAR PROGRESSO
    # ========================================================

    etapa_atual += 1

    tempo = time.time() - inicio_geracao

    tempos_etapas.append(
        tempo
    )

    media = (
        sum(tempos_etapas)
        /
        len(tempos_etapas)
    )

    restante = media * (
        etapas_total
        -
        etapa_atual
    )

    atualizar_progresso(
        50,
        f"Fontes coletadas | Restante: {formatar_tempo(restante)}"
    )

    # ========================================================
    # 14. RESUMO DA COLETA
    # ========================================================

    print()
    print("==============================")
    print("TESTE DE COLETA FINALIZADO")
    print("==============================")

    print()

    print(
        "PÁGINAS COLETADAS:",
        len(paginas)
    )

    print()

    for i, item in enumerate(
        paginas,
        start=1
    ):

        print(
            f"{i} - {item['url']}"
        )

    print()

    print(
        "TEXTOS CAPTURADOS:",
        len(textos)
    )

    print()

    for i, texto in enumerate(
        textos,
        start=1
    ):

        print(
            f"{i} - {len(texto.get('texto', ''))} caracteres"
        )

    # ========================================================
    # 15. GARANTIR FONTES PARA MAPA
    # ========================================================

    if not textos:

        print()
        print("==============================")
        print("NENHUM TEXTO DISPONÍVEL")
        print("==============================")

        return

    # ========================================================
    # 16. NORMALIZAR TEXTOS PARA MAPA
    # ========================================================

    textos_para_mapa = []

    for item in textos:

        if isinstance(
            item,
            dict
        ):

            texto_item = item.get(
                "texto",
                ""
            )

            if texto_item:

                textos_para_mapa.append({

                    "url":
                        item.get(
                            "url",
                            ""
                        ),

                    "tipo":
                        item.get(
                            "tipo",
                            "texto"
                        ),

                    "texto":
                        texto_item

                })

        elif isinstance(
            item,
            str
        ):

            if item.strip():

                textos_para_mapa.append({

                    "url":
                        "",

                    "tipo":
                        "texto",

                    "texto":
                        item

                })

    # ========================================================
    # 17. LIMPAR REFERÊNCIAS
    # ========================================================

    try:

        textos_para_mapa = limpar_lista_referencias(
            textos_para_mapa,
            tema
        )

    except Exception as e:

        print()
        print("==============================")
        print("AVISO - LIMPEZA DO MAPA")
        print("==============================")

        print(
            repr(e)
        )

    # ========================================================
    # 18. CORRIGIR FORMATO FINAL
    # ========================================================

    textos_para_mapa_corrigidos = []

    for item in textos_para_mapa:

        if isinstance(
            item,
            dict
        ):

            texto_item = item.get(
                "texto",
                ""
            )

            if texto_item and texto_item.strip():

                textos_para_mapa_corrigidos.append({

                    "url":
                        item.get(
                            "url",
                            ""
                        ),

                    "tipo":
                        item.get(
                            "tipo",
                            "texto"
                        ),

                    "texto":
                        texto_item

                })

        elif isinstance(
            item,
            str
        ):

            if item.strip():

                textos_para_mapa_corrigidos.append({

                    "url":
                        "",

                    "tipo":
                        "texto",

                    "texto":
                        item

                })

    print()
    print("==============================")
    print("TEXTOS PARA MAPA MEAD")
    print("==============================")

    print(
        "QUANTIDADE:",
        len(
            textos_para_mapa_corrigidos
        )
    )

    # ========================================================
    # 19. GERAR MAPA MEAD
    # ========================================================

    print()
    print("==============================")
    print("GERANDO MAPA MEAD")
    print("==============================")

    mapa_mead = gerar_e_salvar_mapa_mead(
        tema,
        textos_para_mapa_corrigidos
    )

    # ========================================================
    # 20. VALIDAR MAPA
    # ========================================================

    print()
    print("==============================")
    print("RESULTADO MAPA MEAD")
    print("==============================")

    print(
        "TIPO:",
        type(mapa_mead)
    )

    print(
        "CARACTERES:",
        len(
            str(
                mapa_mead or ""
            )
        )
    )

    if mapa_mead:

        print()

        print(
            mapa_mead[:3000]
        )

    else:

        print(
            "MAPA MEAD VAZIO"
        )

        print()
        print("==============================")
        print("CONTEÚDO NÃO SERÁ GERADO")
        print("==============================")

        return

    # ========================================================
    # 21. GERAR CONTEÚDO
    # ========================================================

    print()
    print("==============================")
    print("GERANDO CONTEÚDO COM IA")
    print("==============================")

    print()
    print("==============================")
    print("GERANDO CONTEÚDO COMPLETO")
    print("==============================")

    etapa_atual += 1

    tempo = time.time() - inicio_geracao

    tempos_etapas.append(
        tempo
    )

    media = (
        sum(tempos_etapas)
        /
        len(tempos_etapas)
    )

    restante = media * (
        etapas_total
        -
        etapa_atual
    )

    atualizar_progresso(
        75,
        f"Gerando conteúdo IA | Restante: {formatar_tempo(restante)}"
    )

    # ========================================================
    # 22. CHAMAR GERADOR
    # ========================================================

    print()
    print("==============================")
    print("CHAMANDO GERADOR DE CONTEÚDO")
    print("==============================")

    print(
        "TEXTOS:",
        len(textos)
    )

    print(
        "TEMA:",
        tema
    )

    print(
        "MAPA:",
        len(
            str(
                mapa_mead
            )
        )
    )

    print(
        "ENTENDIMENTO:",
        len(
            str(
                entendimento or ""
            )
        ),
        "caracteres"
    )

    # ========================================================
    # 22.1 VALIDAR DADOS ANTES DO GERADOR
    # ========================================================

    if not textos:

        print()
        print("==============================")
        print("NENHUM TEXTO DISPONÍVEL")
        print("==============================")

        return None

    if not mapa_mead:

        print()
        print("==============================")
        print("MAPA MEAD NÃO DISPONÍVEL")
        print("==============================")

        return None

    print()
    print("==============================")
    print("DADOS ENVIADOS AO GERADOR")
    print("==============================")

    print(
        "TEXTOS BRUTOS:",
        len(textos)
    )

    print(
        "MAPA MEAD:",
        len(
            str(
                mapa_mead
            )
        ),
        "caracteres"
    )

    print(
        "ESTRUTURA EDITORIAL:",
        obter_estrutura_editorial()
    )

    # ========================================================
    # 22.2 CHAMAR GERADOR
    # ========================================================

    conteudo_completo = gerar_conteudo_completo(
        tema,
        textos,
        mapa_mead,
        obter_estrutura_editorial()
    )

    # ========================================================
    # 22.3 VALIDAR RETORNO DO GERADOR
    # ========================================================

    print()
    print("==============================")
    print("RETORNO DO GERADOR")
    print("==============================")

    if not conteudo_completo:

        print(
            "GERADOR NÃO RETORNOU CONTEÚDO"
        )

        return None

    conteudo_completo = str(
        conteudo_completo
    ).strip()

    print(
        "TIPO:",
        type(
            conteudo_completo
        ).__name__
    )

    print(
        "CARACTERES:",
        len(
            conteudo_completo
        )
    )

    # ========================================================
    # 23. RESULTADO CONTEÚDO
    # ========================================================

    print()
    print("==============================")
    print("RESULTADO CONTEÚDO")
    print("==============================")

    print(
        "TIPO:",
        type(conteudo_completo)
    )

    print(
        "CARACTERES:",
        len(
            str(
                conteudo_completo or ""
            )
        )
    )



    # ========================================================
    # 24. SALVAR
    # ========================================================

    if conteudo_completo:

        salvar_banco(
            tema,
            "conteudo_completo",
            conteudo_completo
        )

        print()
        print(
            "CONTEÚDO COMPLETO SALVO NO BANCO"
        )

    else:

        print()
        print(
            "FALHA AO GERAR CONTEÚDO COMPLETO"
        )

    # ========================================================
    # 25. FINAL
    # ========================================================

    print()
    print("==============================")
    print("CONTEÚDO FINALIZADO")
    print("==============================")

    print(
        "TEMA:",
        tema
    )

    print(
        "ENTENDIMENTO GERADO:",
        bool(entendimento)
    )

    print(
        "MAPA MEAD GERADO:",
        bool(mapa_mead)
    )

    print(
        "CONTEÚDO GERADO:",
        bool(conteudo_completo)
    )

    print()
    print(
        "TEMPO TOTAL:",
        formatar_tempo(
            time.time()
            -
            inicio_geracao
        )
    )

    # ========================================================
    # 26. EXECUÇÃO DO PROGRAMA
    # ========================================================

if __name__ == "__main__":

    abrir_interface()
